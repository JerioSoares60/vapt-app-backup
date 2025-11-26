from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request, Depends
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import sys
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from datetime import datetime
import traceback
import io
import shutil
import tempfile
import json
import matplotlib.pyplot as plt
import numpy as np
from collections import Counter
import urllib.parse
from zipfile import ZipFile
import glob
import hashlib
import re
from copy import deepcopy
from collections import defaultdict
import matplotlib
matplotlib.use('Agg')
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: PIL/Pillow library not available. Logo resizing will be disabled.")

# Add parent directory to path to import excel_parser
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../..'))
from excel_parser import parse_excel_data, format_for_template

# Version tracking for code changes
# Format: MAJOR.MINOR.PATCH
# Increment MAJOR for significant structural changes
# Increment MINOR for feature additions or improvements
# Increment PATCH for bug fixes
VERSION = "1.0.7"  # Updated PATCH version for splitting parsing logic

app = FastAPI()
from fastapi import APIRouter
router = APIRouter()
from sqlalchemy.orm import Session
from db import get_db, AuditLog, DashboardDataset

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to store uploaded files
UPLOAD_DIR = "uploads"
# Define the fixed path for the document control file
DOC_CONTROL_FILE = os.path.join(UPLOAD_DIR, "Comprehensive_Vulnerability_Template (2).xlsx")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(os.path.join(UPLOAD_DIR, "screenshots"), exist_ok=True)

# Define the order of severities
severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4, "Unknown": 5}

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        # Sanitize filename to avoid traversal/XSS
        original_name = file.filename or "upload.bin"
        safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", original_name)
        file_location = os.path.join(UPLOAD_DIR, safe_name)
        with open(file_location, "wb") as f:
            f.write(await file.read())
        return JSONResponse(content={"filename": file.filename}, status_code=200)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-screenshots/")
async def upload_screenshots(files: list[UploadFile] = File(...)):
    """
    Upload multiple screenshot files.
    Returns a dictionary of the saved filenames.
    """
    try:
        saved_files = []
        for file in files:
            # Only allow image files
            if not file.content_type.startswith('image/'):
                continue
                
            original_name = file.filename or "image.png"
            safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", original_name)
            screenshots_dir = os.path.join(UPLOAD_DIR, "screenshots")
            os.makedirs(screenshots_dir, exist_ok=True)
            file_location = os.path.join(screenshots_dir, safe_name)
            with open(file_location, "wb") as f:
                f.write(await file.read())
            saved_files.append(safe_name)
        
        return JSONResponse(
            content={"filenames": saved_files},
            status_code=200
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# Function to determine severity from CVSS score
def get_severity_from_cvss(cvss):
    print(f"Determining severity for CVSS: {cvss}, type: {type(cvss)}")  # Debug output
    
    if cvss is None:
        print("CVSS is None, defaulting to Informational")
        return "Informational"
    
    try:
        cvss_float = float(cvss)
        if cvss_float >= 9.0 and cvss_float <= 10.0:
            print(f"CVSS {cvss_float} classified as Critical")
            return "Critical"
        elif cvss_float >= 7.0 and cvss_float < 9.0:
            print(f"CVSS {cvss_float} classified as High")
            return "High"
        elif cvss_float >= 4.0 and cvss_float < 7.0:
            print(f"CVSS {cvss_float} classified as Medium")
            return "Medium"
        elif cvss_float > 0.0 and cvss_float < 4.0:
            print(f"CVSS {cvss_float} classified as Low")
            return "Low"
        else:
            print(f"CVSS {cvss_float} classified as Informational")
            return "Informational"
    except (ValueError, TypeError) as e:
        print(f"Error converting CVSS to float: {e}, defaulting to Informational")
        return "Informational"

# Function to get severity colors based on severity level
def get_severity_colors(severity):
    colors = {
        "Critical": {"row1_bg": "#990000", "row2_bg": "#FF3333", "font": "#990000"},
        "High": {"row1_bg": "#FF0000", "row2_bg": "#FF6666", "font": "#FF0000"},
        "Medium": {"row1_bg": "#FFCC00", "row2_bg": "#FFCC66", "font": "#FFCC00"},
        "Low": {"row1_bg": "#009933", "row2_bg": "#99CC33", "font": "#009933"},
        "Informational": {"row1_bg": "#3399CC", "row2_bg": "#66CCFF", "font": "#3399CC"}
    }
    return colors.get(severity, colors["Informational"])

# Function to generate more detailed vulnerability content based on vulnerability name
def get_vulnerability_description(vuln_name):
    """
    Generate a detailed description paragraph for the vulnerability based on OWASP guidance.
    """
    descriptions = {
        "SQL Injection": """SQL Injection is a code injection technique that exploits vulnerabilities in the interface between web applications and database servers. The vulnerability occurs when user input is incorrectly filtered and directly included in SQL statements, allowing attackers to manipulate the structure of the SQL query. This can lead to unauthorized access to sensitive data, bypassing authentication mechanisms, and in severe cases, compromising the underlying server or performing denial of service attacks. The primary cause is insufficient validation and sanitization of user-supplied input used in database queries.""",
        
        "Cross-Site Scripting": """Cross-Site Scripting (XSS) is a security vulnerability that allows attackers to inject malicious client-side scripts into web pages viewed by other users. These attacks occur when an application includes untrusted data in a new web page without proper validation or encoding. XSS vulnerabilities can be found in various forms: Reflected XSS where malicious script is reflected off the web server in an error message or search result; Stored XSS where the malicious script is permanently stored on target servers; and DOM-based XSS where vulnerability exists in client-side code. The impact of XSS can range from session theft to complete compromise of user accounts.""",
        
        "File Upload": """Unrestricted File Upload vulnerabilities allow attackers to upload files of dangerous types or content to the server. This vulnerability occurs when an application fails to properly validate the type, content, size, or naming convention of uploaded files. Attackers can exploit this to upload web shells, malware, or other malicious files that allow them to execute arbitrary code on the server. This may lead to complete system compromise, data theft, or defacement of websites. The vulnerability is particularly dangerous when uploaded files are stored in web-accessible locations, allowing direct execution through browser requests.""",
        
        "Insecure Direct Object References": """Insecure Direct Object References (IDOR) occur when an application provides direct access to objects based on user-supplied input without appropriate authorization checks. This vulnerability allows attackers to bypass authorization and directly access resources in the system, potentially exposing sensitive data that should be protected. IDOR vulnerabilities commonly manifest as manipulated parameters within requests that reveal sensitive objects like files, database records, or user accounts. This can lead to unauthorized access to sensitive information, modification of data belonging to other users, or privilege escalation allowing attackers to gain administrative access to the application.""",
        
        "Sensitive Data Exposure": """Sensitive Data Exposure occurs when an application fails to adequately protect sensitive information such as financial data, healthcare information, personally identifiable information (PII), or authentication credentials. This vulnerability can manifest through various channels including insufficient encryption of data at rest or in transit, improper handling of sensitive information in logs or error messages, and inadequate access controls. Attackers can exploit these weaknesses to access sensitive data through packet sniffing, man-in-the-middle attacks, or by directly accessing poorly protected storage. The impact can include identity theft, financial fraud, privacy violations, and regulatory compliance failures."""
    }
    
    # Try to find a match in our predefined descriptions
    for key in descriptions:
        if key.lower() in vuln_name.lower():
            return descriptions[key]
    
    # Default description for any vulnerability not specifically defined
    return f"""This vulnerability ({vuln_name}) allows attackers to potentially compromise the application security through improper input validation and processing mechanisms. The vulnerability exists due to insufficient security controls in the application code, which could allow malicious actors to manipulate system behavior in unintended ways. Such vulnerabilities often arise from coding practices that do not account for all possible edge cases or user inputs, creating potential attack vectors. When discovered, these issues should be addressed promptly to prevent potential exploitation."""

def get_vulnerability_impact(vuln_name):
    """
    Generate a detailed impact paragraph for the vulnerability based on OWASP guidance.
    """
    impacts = {
        "SQL Injection": """The impact of SQL Injection vulnerabilities can be severe and far-reaching for an organization. Successful exploitation can lead to unauthorized access to sensitive database information, including personal user data, financial records, and intellectual property. Attackers may bypass authentication mechanisms, gaining unauthorized administrative access to the application. In more severe cases, attackers might execute commands on the database server, potentially leading to complete server compromise. This can result in significant financial losses, regulatory penalties for data breaches, damage to company reputation, and loss of customer trust. Business operations may be disrupted if data is stolen, manipulated, or deleted entirely.""",
        
        "Cross-Site Scripting": """The exploitation of Cross-Site Scripting vulnerabilities can have significant consequences for both users and the organization. Attackers can hijack user sessions, steal authentication cookies, and impersonate legitimate users to access sensitive information or perform unauthorized actions. They may also inject malicious content that defaces websites, leads to redirection to malicious sites, or installs malware on user devices. For organizations, this can lead to reputational damage, loss of customer trust, potential data breaches, and regulatory compliance issues. XSS attacks targeting administrative users can result in complete application compromise, allowing attackers to access backend systems, modify server-side logic, or obtain sensitive business data.""",
        
        "File Upload": """Unrestricted file upload vulnerabilities can have devastating impacts on organization security. Successfully exploited vulnerabilities allow attackers to gain remote code execution capabilities on the server, potentially leading to complete system compromise. This can result in unauthorized access to databases, internal networks, and sensitive enterprise systems. Attackers may establish persistence through backdoors, maintain long-term access to compromised systems, and use the server as a platform to attack other internal systems. Organizations face significant risks including data theft, intellectual property loss, financial loss, service disruption due to regulatory penalties. The reputational damage from such breaches can lead to loss of customer confidence and business opportunities.""",
        
        "Insecure Direct Object References": """The business impact of Insecure Direct Object References can be substantial. Exploitation allows attackers to access unauthorized data across user boundaries, potentially exposing sensitive personal information, financial data, or proprietary business information. This can lead to privacy violations, identity theft, financial fraud, or industrial espionage. Organizations may face significant regulatory penalties for data protection failures, particularly under frameworks like GDPR, HIPAA, or PCI-DSS. The reputational damage from such data breaches can severely impact customer trust and result in customer churn. Additionally, competitors may gain access to confidential business information, leading to competitive disadvantage and potential revenue loss.""",
        
        "Sensitive Data Exposure": """Exposure of sensitive data can have catastrophic consequences for both individuals and organizations. When personal information is compromised, individuals may become victims of identity theft, financial fraud, or reputational damage. For organizations, the impacts include regulatory fines and penalties under data protection laws such as GDPR, HIPAA, or PCI-DSS, which can amount to millions of dollars. Legal costs may escalate through lawsuits from affected parties seeking damages. Business reputation suffers significantly, leading to loss of customer trust, decreased market share, and reduced stock value for public companies. Operational disruptions can occur while addressing the breach, and long-term business relationships may be damaged if partner data is compromised."""
    }
    
    # Try to find a match in our predefined impacts
    for key in impacts:
        if key.lower() in vuln_name.lower():
            return impacts[key]
    
    # Default impact for any vulnerability not specifically defined
    return f"""Exploitation of this vulnerability ({vuln_name}) could lead to unauthorized access, data breaches, reputation damage, and potential regulatory compliance violations. Customers' trust may be eroded, resulting in business loss. The organization may face financial penalties from regulatory bodies for failing to protect sensitive information adequately. Remediation costs can be substantial, including not only the technical fixes but also potential legal fees, customer notification expenses, and increased security monitoring. In severe cases, business operations might be disrupted, leading to revenue loss and long-term damage to the organization's brand and market position."""

def get_vulnerability_recommendation(vuln_name):
    """
    Generate detailed remediation recommendations based on OWASP guidance.
    """
    recommendations = {
        "SQL Injection": """To remediate SQL Injection vulnerabilities, implement parameterized queries (prepared statements) for all database operations, which ensures that user input is treated as data rather than executable code. Apply the principle of least privilege to database accounts used by applications, limiting their capabilities to only necessary operations. Implement input validation on both client and server sides, with a whitelist approach that accepts only known-good input. Utilize stored procedures and ORM (Object-Relational Mapping) frameworks that handle SQL escaping automatically. Implement a WAF (Web Application Firewall) as an additional layer of protection to filter malicious SQL injection attempts. Regularly scan application code for SQL injection vulnerabilities and conduct penetration testing to verify the effectiveness of implemented controls.""",
        
        "Cross-Site Scripting": """To effectively mitigate Cross-Site Scripting vulnerabilities, implement context-sensitive output encoding for all user-supplied data before rendering it in HTML, JavaScript, CSS, or URL contexts. Use Content Security Policy (CSP) headers to restrict sources of executable scripts and prevent execution of inline JavaScript. Apply input validation with a whitelist approach to filter potentially dangerous inputs before processing. Utilize modern web frameworks that automatically handle XSS protection through built-in template escaping mechanisms. For cookie security, implement the HttpOnly flag to prevent JavaScript access and use the Secure flag for transmission over HTTPS only. Keep all frameworks and libraries updated to ensure known XSS vulnerabilities are patched. Regularly conduct security testing including both automated scanning and manual code review to identify potential XSS vulnerabilities.""",
        
        "File Upload": """To securely handle file uploads, implement strict file type validation using content inspection (checking file headers) rather than relying only on file extensions. Enforce file size limits to prevent denial of service attacks through oversized file uploads. Store uploaded files outside the web root or application directory to prevent direct execution, and if files must be publicly accessible, store them in a different domain or subdomain. Generate random file names to prevent directory traversal attempts and consistent naming patterns that could be guessed. Scan uploaded files with antivirus software or integrate with security services that can detect malicious content. Implement proper access controls to ensure uploaded files are only accessible by authorized users. Consider using a CDN or dedicated file storage service that specializes in secure file handling for high-risk environments.""",
        
        "Insecure Direct Object References": """To remediate Insecure Direct Object References, implement proper access control checks for every request that accesses data objects, verifying that the authenticated user has appropriate permissions for the requested resource. Replace direct references to objects with indirect references or tokens that are mapped server-side to actual resource identifiers. Use server-side session objects to store and verify user authorization to access specific resources. Create a centralized authorization mechanism that provides consistent enforcement across the application. Implement the principle of least privilege, ensuring users only have access to resources necessary for their legitimate functions. Log access attempts to sensitive resources and configure alerts for potential access control violations. Regularly review and test access control mechanisms through security assessments and penetration testing.""",
        
        "Sensitive Data Exposure": """To protect sensitive data effectively, implement strong encryption for data both at rest and in transit using industry-standard algorithms and appropriate key management. Use TLS (at least version 1.2 or higher) for all connections that transmit sensitive information, and enforce this through HSTS headers. Store passwords using strong adaptive hashing algorithms with salts (such as bcrypt, Argon2, or PBKDF2). Disable caching for responses containing sensitive data using appropriate cache-control headers. Minimize the collection and retention of sensitive data – only collect what is absolutely necessary for business functions and purge unnecessary data regularly. Implement proper key management practices including secure key generation, storage, distribution, and rotation. Classify data based on sensitivity and apply appropriate controls based on classification. Regularly audit systems for sensitive data handling compliance and conduct penetration testing focused on data protection controls."""
    }
    
    # Try to find a match in our predefined recommendations
    for key in recommendations:
        if key.lower() in vuln_name.lower():
            return recommendations[key]
    
    # Default recommendation for any vulnerability not specifically defined
    return f"""To remediate this {vuln_name} vulnerability, implement proper input validation with a strong allowlist approach that verifies all user inputs against expected formats and values. Follow the principle of least privilege throughout the application, ensuring components only have access to resources necessary for their legitimate function. Develop and maintain secure coding practices including regular code reviews and security-focused testing. Keep all software components, including frameworks and libraries, updated with security patches. Implement defense in depth with multiple layers of security controls, so that if one layer fails, others will still provide protection. Conduct regular security assessments and penetration testing to identify vulnerabilities before they can be exploited. Create a security-focused development lifecycle that incorporates security at every stage from design to deployment and maintenance."""

def normalize_filename(fname):
    """Normalize filenames for robust matching: strip, lower, keep extension."""
    return fname.strip().lower() if fname else ''

# Function to parse the uploaded Excel file for vulnerability data
def parse_vulnerabilities_excel(file_path):
    """
    Parse the uploaded Excel file using standardized parser and convert to legacy format.
    This function focuses only on vulnerability details and excludes document control info.
    """
    print(f"Parsing vulnerability data from: {file_path}")
    try:
        # Try to use standardized parser first
        try:
            # type2 is IP report - read from "IP" sheet
            parsed_data = parse_excel_data(file_path, sheet_name="IP")
            print(f"✅ Using standardized Excel parser")
            print(f"   Metadata: {parsed_data['metadata']}")
            print(f"   Vulnerabilities: {len(parsed_data['vulnerabilities'])}")
            
            vulnerabilities = []
            for index, vuln in enumerate(parsed_data['vulnerabilities']):
                # Convert steps to steps_with_screenshots format
                steps_with_screenshots = []
                if vuln.get('steps'):
                    for step in vuln['steps']:
                        steps_with_screenshots.append({
                            'text': step['content'],
                            'screenshot': ''  # Screenshots handled separately via PoC mapping
                        })
                
                # Extract CVE/CWE list
                associated_cves = []
                cve_cwe_str = vuln.get('cve_cwe', '')
                if cve_cwe_str:
                    associated_cves = [c.strip() for c in re.split(r'[\n,;]', cve_cwe_str) if c.strip()]
                
                # Convert to type2 format
                vulnerability = {
                    'name': vuln.get('observation', ''),
                    'description': vuln.get('detailed_observation', '') or vuln.get('observation_summary', ''),
                    'impact': '',  # Not in standardized format, will be generated
                    'severity': vuln.get('severity', 'Medium'),
                    'cvss': vuln.get('cvss', None),
                    'ip': vuln.get('affected_asset', '') or vuln.get('ip_url_app', ''),
                    'vulnerable_parameter': '',  # Not in standardized format
                    'remediation': vuln.get('recommendation', ''),
                    'steps_with_screenshots': steps_with_screenshots,
                    'sr_no': str(vuln.get('sr_no', f"VUL-{index+1:03d}")).replace('VULN-', 'VUL-').strip(),
                    'associated_cves': associated_cves,
                    'reference_link': vuln.get('reference', ''),
                    'reported_by': vuln.get('tester', parsed_data['metadata']['tester']),
                    'project': vuln.get('project', parsed_data['metadata']['project']),
                    'severity_level': vuln.get('severity', 'Medium'),
                    'is_no_vuln_box': False,
                    'row_order': int(index),
                }
                
                vulnerabilities.append(vulnerability)
                print(f"Parsed vulnerability Sr No: {vulnerability['sr_no']} at index {index}")
            
            print(f"Total vulnerabilities extracted (standardized parser): {len(vulnerabilities)}")
            return vulnerabilities
            
        except Exception as std_parse_error:
            print(f"⚠️  Standardized parser failed, falling back to legacy parser: {std_parse_error}")
            traceback.print_exc()
        
        # Fallback to original parsing logic
        df = pd.read_excel(file_path)
        print(f"Vulnerability Excel file read successfully. Columns: {df.columns.tolist()}")
        print(f"Excel file has {len(df)} rows")
        
        # Clean column names by stripping whitespace
        df.columns = df.columns.str.strip()

        # Forward-fill identification/summary columns so multiple rows under the
        # same IP can omit repeated metadata (Hostname, counts etc.)
        ffill_cols = [
            'Hostname', 'IP Type', 'VAPT Status', 'Critical', 'High',
            'Medium', 'Low', 'Informational', 'Total', 'ReportedBy', 'Project', 'Severity'
        ]
        for col in ffill_cols:
            if col in df.columns:
                df[col] = df[col].ffill()
        
        vulnerabilities = []
        
        for index, row in df.iterrows():
            try:
                # Skip completely blank lines that do not carry a vulnerability entry
                vn_raw = str(row.get('Vulnerability Name', '')).strip()
                sr_raw = str(row.get('Sr No', '')).strip()
                if vn_raw == '' and sr_raw == '':
                    continue
                # Extract CVSS score and convert to float
                cvss_score = None
                if 'CVSS Score' in row:
                    cvss_score = row['CVSS Score']
                elif 'CVSS' in row:  # Also check for 'CVSS' column for backwards compatibility
                    cvss_score = row['CVSS']
                
                # Get Vulnerable URL
                vulnerable_url = None
                if 'Vulnerable URL' in row:
                    vulnerable_url = row['Vulnerable URL']
                elif 'Vulnerable Parameter' in row:  # Check for 'Vulnerable Parameter' for backwards compatibility
                    vulnerable_url = row['Vulnerable Parameter']
                
                # Get Vulnerable Parameter (specific field)
                vulnerable_parameter = None
                if 'Vulnerable Parameter' in row:
                    parameter_value = row['Vulnerable Parameter']
                    # Check if it's NaN and handle it
                    if isinstance(parameter_value, float) and pd.isna(parameter_value):
                        vulnerable_parameter = None
                    else:
                        vulnerable_parameter = str(parameter_value)
                
                # Get Steps with screenshots if available
                steps_with_screenshots = []
                steps_text = str(row.get('Steps', '')).splitlines()
                screenshot_cols = [col for col in df.columns if col.strip().lower().startswith('screenshot')]
                for i, step_text in enumerate(steps_text):
                    screenshot_filename = ''
                    if i < len(screenshot_cols):
                        screenshot_filename = str(row.get(screenshot_cols[i], '')).strip().lower()
                    steps_with_screenshots.append({
                        'text': step_text.strip(),
                        'screenshot': screenshot_filename
                    })
                
                # Determine severity based on CVSS score
                severity = get_severity_from_cvss(cvss_score)

                # Get vulnerability name
                vuln_name = row.get('Vulnerability Name', '')
                def clean_str(s):
                    return str(s).strip().replace('\u200b', '').replace('\xa0', '').replace(' ', '').lower()
                is_no_vuln_box = clean_str(vuln_name) in ['novulnerabilityfound', 'novulnerability']
                
                # For 'No Vulnerability' rows, allow missing/empty fields
                if is_no_vuln_box:
                    description = row.get('Description', '') or row.get('Vulnerability Description', '') or ''
                    ip = row.get('Hostname', '')
                    steps_with_screenshots = []
                    steps_text = str(row.get('Steps', '')).splitlines()
                    screenshot_cols = [col for col in df.columns if col.strip().lower().startswith('screenshot')]
                    for i, step_text in enumerate(steps_text):
                        screenshot_filename = ''
                        if i < len(screenshot_cols):
                            screenshot_filename = str(row.get(screenshot_cols[i], '')).strip().lower()
                        steps_with_screenshots.append({
                            'text': step_text.strip(),
                            'screenshot': screenshot_filename
                        })
                    vulnerability = {
                        'name': vuln_name,
                        'description': description,
                        'impact': '',
                        'severity': '',
                        'cvss': None,
                        'ip': ip,
                        'vulnerable_parameter': '',
                        'remediation': '',
                        'steps_with_screenshots': steps_with_screenshots,
                        'sr_no': str(row.get('Sr No', f"VUL-{index+1:03d}")).replace('VULN-', 'VUL-').strip(),
                        'associated_cves': [],
                        'is_no_vuln_box': True,
                    }
                    vulnerabilities.append(vulnerability)
                    print(f"Parsed NO VULNERABILITY box Sr No: {vulnerability['sr_no']} at index {index}")
                    continue  # Skip the rest of the parsing for this row
                
                # Use custom description or generate one based on vulnerability name
                description = row.get('Description', '')
                if not description:
                    description = get_vulnerability_description(vuln_name)
                
                # Use custom remediation or generate one based on vulnerability name
                remediation = row.get('Remediation', '')
                if not remediation:
                    remediation = get_vulnerability_recommendation(vuln_name)
                
                # Get impact information (either from Excel or generate)
                impact = row.get('Impact', '')
                if not impact:
                    impact = get_vulnerability_impact(vuln_name)
                
                # Get Associated CVEs (as a list)
                associated_cves = []
                # Robustly find the column for Associated CVEs (case-insensitive, ignore spaces and apostrophes)
                cve_col = None
                for col in df.columns:
                    if col.strip().replace("'", "").replace(' ', '').lower() in [
                        'associatedcves', 'associatedcve', 'associatedcvees', 'associatedcvees']:
                        cve_col = col
                        break
                if cve_col and pd.notna(row.get(cve_col, None)):
                    cve_str = str(row.get(cve_col, ''))
                    print(f"[DEBUG] Row {index} Associated CVEs raw: {cve_str}")
                    # Split by comma, semicolon, or newline
                    associated_cves = [c.strip() for c in re.split(r'[\n,;]', cve_str) if c.strip()]
                
                # Get Reference Link
                reference_link = ""
                # Robustly find the column for Reference Link (case-insensitive, ignore spaces and apostrophes)
                ref_link_col = None
                for col in df.columns:
                    if col.strip().replace("'", "").replace(' ', '').lower() in [
                        'referencelink', 'referencelinks', 'refrencelink', 'refrencelinks', 'referencelink']:
                        ref_link_col = col
                        break
                if ref_link_col and pd.notna(row.get(ref_link_col, None)):
                    ref_link_str = str(row.get(ref_link_col, '')).strip()
                    if ref_link_str and ref_link_str.lower() not in ['', 'nan', 'none', 'null']:
                        reference_link = ref_link_str
                        print(f"[DEBUG] Row {index} Reference Link: {reference_link}")
                
                # Get new dashboard columns
                reported_by = str(row.get('ReportedBy', '')).strip() if pd.notna(row.get('ReportedBy', None)) else ''
                project = str(row.get('Project', '')).strip() if pd.notna(row.get('Project', None)) else ''
                severity_level = str(row.get('Severity', '')).strip() if pd.notna(row.get('Severity', None)) else ''
                
                vulnerability = {
                    'name': vuln_name,
                    'description': description,
                    'impact': impact,
                    'severity': severity,
                    'cvss': cvss_score,
                    'ip': row.get('Hostname', ''),
                    'vulnerable_parameter': vulnerable_parameter,
                    'remediation': remediation,
                    'steps_with_screenshots': steps_with_screenshots,
                    'sr_no': str(row.get('Sr No', f"VUL-{index+1:03d}")).replace('VULN-', 'VUL-').strip(),
                    'associated_cves': associated_cves,
                    'reference_link': reference_link,
                    'reported_by': reported_by,
                    'project': project,
                    'severity_level': severity_level,
                    'is_no_vuln_box': is_no_vuln_box,
                    'row_order': int(index),
                }
                
                vulnerabilities.append(vulnerability)
                
                print(f"Parsed vulnerability Sr No: {vulnerability['sr_no']} at index {index}") # Debug print
                
            except Exception as e:
                print(f"Error processing row {index}: {e}")
                continue
        
        print(f"Total vulnerabilities extracted: {len(vulnerabilities)}")
        print("After parsing:")
        for i, v in enumerate(vulnerabilities):
            print(f"  {i+1}. SR: {v['sr_no']}, Name: {v['name']}, is_no_vuln_box: {v['is_no_vuln_box']}, IP: {v['ip']}")
        # Print all unique IPs parsed
        all_ips = set(v['ip'] for v in vulnerabilities if v['ip'])
        print(f"All unique IPs parsed: {sorted(all_ips)}")
        
        # Count by type
        no_vuln_count = sum(1 for v in vulnerabilities if v['is_no_vuln_box'])
        regular_vuln_count = sum(1 for v in vulnerabilities if not v['is_no_vuln_box'])
        print(f"Breakdown: {regular_vuln_count} regular vulnerabilities, {no_vuln_count} 'No Vulnerability' entries")
        
        # The vulnerabilities list is now in the order they appeared in the Excel file.
        # No further sorting is needed to maintain the original order.
        
        return vulnerabilities
        
    except Exception as e:
        print(f"Error parsing vulnerability Excel: {e}")
        return []

# Function to parse the fixed document control file
def parse_doc_control_excel(file_path):
    """
    Parse the fixed document control Excel file and return document control details
    including assessment scope, dates, note, and amendment log.
    """
    print(f"Parsing document control data from: {file_path}")
    doc_control_data = {
        'assessment_scope': '',
        'assessment_start_date': '',
        'assessment_end_date': '',
        'assessment_note': '',  # Add assessment_note
        'amendment_log': [],
        'reviewed_by': '',
        'authorized_by': ''
    }
    try:
        df = pd.read_excel(file_path)
        print(f"Document Control Excel file read successfully. Columns: {df.columns.tolist()}")

        # Clean column names by stripping whitespace
        df.columns = df.columns.str.strip()
        
        # Find the header row containing "Version", "Date", "Reviewed By", "Brief description of the change"
        header_row_index = None
        for index, row in df.iterrows():
            # Check if key headers are present in this row
            if all(col in row.index for col in ["Version", "Date", "Reviewed By", "Brief description of the change"]):
                header_row_index = index
                print(f"Found document control header row at index: {header_row_index}")
                break
        
        if header_row_index is None:
            print("Document control header row not found. Cannot extract document control details or amendment log.")
            return doc_control_data # Return empty data

        # Assuming assessment details are in the row *after* the header row (first data row)
        first_data_row_index = header_row_index + 1
        if first_data_row_index < len(df):
            first_data_row = df.iloc[first_data_row_index]
            
            # Handle Assessment Scope
            if 'Assessment Scope' in first_data_row:
                doc_control_data['assessment_scope'] = str(first_data_row['Assessment Scope']).strip()
                print(f"Found Assessment Scope: {doc_control_data['assessment_scope']}")
            # Handle Assessment Note
            if 'Assessment Note' in first_data_row:
                doc_control_data['assessment_note'] = str(first_data_row['Assessment Note']).strip()
                print(f"Found Assessment Note: {doc_control_data['assessment_note']}")
            # Handle Assessment Start Date
            if 'Assessment Start Date' in first_data_row:
                start_date = first_data_row['Assessment Start Date']
                print(f"Raw Start Date (Doc Control): {start_date}, Type: {type(start_date)}")
                if isinstance(start_date, pd.Timestamp):
                    doc_control_data['assessment_start_date'] = start_date.strftime('%dth %B %Y')
                else:
                    # Try to parse the date string if it's in the format "10th March 2025"
                    try:
                         # Remove 'th', 'st', 'nd', 'rd' from the day
                        date_str = str(start_date).strip()
                        date_str = date_str.replace('th', '').replace('st', '').replace('nd', '').replace('rd', '')
                        # Parse the date
                        parsed_date = pd.to_datetime(date_str, format='%d %B %Y')
                        doc_control_data['assessment_start_date'] = parsed_date.strftime('%dth %B %Y')
                    except:
                        doc_control_data['assessment_start_date'] = str(start_date).strip()
                print(f"Formatted Start Date (Doc Control): {doc_control_data['assessment_start_date']}")

            # Handle Assessment End Date
            if 'Assessment End Date' in first_data_row:
                end_date = first_data_row['Assessment End Date']
                print(f"Raw End Date (Doc Control): {end_date}, Type: {type(end_date)}")
                if isinstance(end_date, pd.Timestamp):
                    doc_control_data['assessment_end_date'] = end_date.strftime('%dth %B %Y')
                else:
                     # Try to parse the date string if it's in the format "30th March 2025"
                    try:
                         # Remove 'th', 'st', 'nd', 'rd' from the day
                        date_str = str(end_date).strip()
                        date_str = date_str.replace('th', '').replace('st', '').replace('nd', '').replace('rd', '')
                        # Parse the date
                        parsed_date = pd.to_datetime(date_str, format='%d %B %Y')
                        doc_control_data['assessment_end_date'] = parsed_date.strftime('%dth %B %Y')
                    except:
                        doc_control_data['assessment_end_date'] = str(end_date).strip()
                print(f"Formatted End Date (Doc Control): {doc_control_data['assessment_end_date']}")
            
            # Extract "Reviewed By" and "Authorized By" from the first data row
            if first_data_row_index < len(df):
                first_data_row = df.iloc[first_data_row_index]
                doc_control_data['reviewed_by'] = str(first_data_row.get('Reviewed By', '')).strip()
                doc_control_data['authorized_by'] = str(first_data_row.get('Authorized By', '')).strip()
                print(f"Found Reviewed By: {doc_control_data['reviewed_by']}")
                print(f"Found Authorized By: {doc_control_data['authorized_by']}")

            print(f"Final assessment details from doc control:")
            print(f"Scope: {doc_control_data['assessment_scope']}")
            print(f"Start Date: {doc_control_data['assessment_start_date']}")
            print(f"End Date: {doc_control_data['assessment_end_date']}")
            print(f"Note: {doc_control_data['assessment_note']}")

        # Extract amendment log entries from rows below the header row
        amendment_log_start_index = header_row_index + 1
        for index in range(amendment_log_start_index, len(df)):
            row = df.iloc[index]
            # Stop if the "Version" cell is empty/NaN
            if pd.isna(row.get('Version', None)) or str(row.get('Version', '')).strip() == '':
                break # Stop reading amendment log
            
            # Handle Date formatting to remove time
            log_date = row.get('Date', '')
            formatted_date = str(log_date).strip()
            if isinstance(log_date, pd.Timestamp):
                formatted_date = log_date.strftime('%Y-%m-%d') # Format as YYYY-MM-DD
            
            log_entry = {
                'Version': str(row.get('Version', '')).strip(),
                'Date': formatted_date,
                'Reviewed By': str(row.get('Reviewed By', '')).strip(),
                'Brief description of the change': str(row.get('Brief description of the change', '')).strip()
            }
            doc_control_data['amendment_log'].append(log_entry)
            print(f"Added amendment log entry: {log_entry}")

        print(f"Total amendment log entries extracted: {len(doc_control_data['amendment_log'])}")
        
        return doc_control_data

    except FileNotFoundError:
        print(f"Error: Document control file not found at {file_path}")
        # Continue without document control data if file is missing
        return doc_control_data
    except Exception as e:
        print(f"Error parsing document control Excel file: {e}")
        # Continue without document control data if parsing fails
        return doc_control_data

def create_summary_table(doc, vulnerabilities):
    """
    Creates and inserts a color-coded vulnerability summary table.
    This function replaces a placeholder in the document.
    """
    placeholder_text = '_VULNERABILITY_SUMMARY_TABLE_PLACEHOLDER_'
    
    # Find the placeholder paragraph
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if placeholder_text in p.text:
            placeholder_paragraph = p
            break
            
    if not placeholder_paragraph:
        print("Warning: Placeholder '_VULNERABILITY_SUMMARY_TABLE_PLACEHOLDER_' not found in the template.")
        return

    # Headers and table creation
    headers = ['Sl.no', 'Vulnerability Name', 'Vulnerability Risk Type']
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set column widths for better layout
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(1.5)

    # Style header row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        # Clear the default paragraph and add new text
        p = cell.paragraphs[0]
        p.text = header_text
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White font
        # Set header background color (purple)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Add data rows with color coding
    for idx, vuln in enumerate(vulnerabilities, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = vuln.get('name', 'N/A')
        
        severity = vuln.get('severity', 'Unknown')
        severity_cell = row_cells[2]
        severity_cell.text = severity
        
        # Get severity colors and apply them
        colors = get_severity_colors(severity)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
        severity_cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Set font to white for dark backgrounds for better readability
        if severity in ["Critical", "High"]:
            p = severity_cell.paragraphs[0]
            if p.runs:
                run = p.runs[0]
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Replace placeholder paragraph with the new table
    parent = placeholder_paragraph._p.getparent()
    parent.insert(parent.index(placeholder_paragraph._p), table._element)
    parent.remove(placeholder_paragraph._p)

# Function to create a styled vulnerability table in Word document
def create_vulnerability_table(doc, vulnerability, display_sr_no=None, image_map=None):
    print(f"Creating table for vulnerability: {vulnerability['name']}, Severity: {vulnerability['severity']}")  # Debug output
    
    # Row titles in exact order requested
    row_labels = [
        "",  # Row 1: Vulnerability Name (handled specially)
        "",  # Row 2: IP (handled specially)
        "Vulnerable Parameter",  # Row 3
        "CVSS Score",  # Row 4
        "Severity",  # Row 5
        "Vulnerability Description",  # Row 6
        "Vulnerability Impact",  # Row 7
        "Recommendation",  # Row 8
        "Associated CVEs",  # Row 9 (new)
        "Proof of Concept / Steps to Reproduce"  # Row 10
    ]
    
    # Create a table with the appropriate number of rows
    # We need 10 rows in total as per the specified order
    table = doc.add_table(rows=len(row_labels), cols=1)
    table.style = 'Table Grid'  # Add borders
    
    # Get severity and colors
    severity = vulnerability["severity"]
    colors = get_severity_colors(severity)
    
    print(f"Using colors: {colors}")  # Debug output
    
    # Extract RGB components from hex colors
    row1_bg_hex = colors["row1_bg"].lstrip('#')
    row2_bg_hex = colors["row2_bg"].lstrip('#')
    font_hex = colors["font"].lstrip('#')
    
    row1_r, row1_g, row1_b = tuple(int(row1_bg_hex[i:i+2], 16) for i in (0, 2, 4))
    row2_r, row2_g, row2_b = tuple(int(row2_bg_hex[i:i+2], 16) for i in (0, 2, 4))
    font_r, font_g, font_b = tuple(int(font_hex[i:i+2], 16) for i in (0, 2, 4))
    
    # Row 1: Vulnerability Name
    row = table.rows[0]
    cell = row.cells[0]
    # Apply background color based on severity
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add text with formatting - use the display_sr_no if provided
    sr_display = display_sr_no if display_sr_no else vulnerability.get('sr_no', '')
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(f"{sr_display}: {vulnerability['name']}")
    run.font.name = 'Altone Trial'  # Change to Altone Trial font
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Row 2: IP
    row = table.rows[1]
    cell = row.cells[0]
    # Apply background color based on severity
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row2_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add text with formatting
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("IP: ")
    run.font.name = 'Altone Trial'  # Change to Altone Trial font
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    run = paragraph.add_run(vulnerability.get('ip', '') or "N/A")
    run.font.name = 'Altone Trial'  # Change to Altone Trial font
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Rows 3-10: Content rows with white background and proper formatting
    # Prepare content for each row - ENSURE ALL VALUES ARE STRINGS
    vuln_param = vulnerability.get('vulnerable_parameter', '')
    # If vuln_param is a float/nan, handle it
    if isinstance(vuln_param, float) and pd.isna(vuln_param):
        vuln_param = "N/A"
        
    row_contents = [
        str(vuln_param),  # Vulnerable Parameter
        str(vulnerability['cvss']) if vulnerability['cvss'] is not None else "N/A",  # CVSS Score
        str(vulnerability['severity']) if vulnerability['severity'] else "N/A",  # Severity
        str(vulnerability['description']) if vulnerability['description'] else "N/A",  # Vulnerability Description
        str(vulnerability.get('impact', "N/A")),  # Vulnerability Impact
        str(vulnerability.get('remediation', "N/A")),  # Recommendation
        "",  # Associated CVEs (handled below)
        ""   # Steps (handled below)
    ]
    
    # Process content rows with white background and proper formatting
    for i in range(2, len(row_labels) - 2):  # Start from row index 2 (third row)
        row_idx = i  # Already 0-indexed
        row = table.rows[row_idx]
        cell = row.cells[0]
        
        # Apply white background
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Add text with formatting
        paragraph = cell.paragraphs[0]
        
        # Label with color based on severity
        run = paragraph.add_run(f"{row_labels[i]}: ")
        run.font.name = 'Altone Trial'  # Change to Altone Trial font
        run.font.size = Pt(12)
        run.font.bold = True
        # Set font color using the RGB values from the severity color
        run.font.color.rgb = RGBColor(font_r, font_g, font_b)
        
        # Content in normal formatting (except for steps which we'll handle specially)
        if i < len(row_labels) - 2:  # Not the steps row
            run = paragraph.add_run(row_contents[i-2])
            run.font.name = 'Altone Trial'  # Change to Altone Trial font
            run.font.size = Pt(11)
            run.font.bold = False
        else:  # This is the Steps row - handle steps with screenshots
            steps = vulnerability.get('steps_with_screenshots', [])
            
            # Remove duplicate steps based on text content
            unique_steps = []
            seen_texts = set()
            for step in steps:
                step_text = step.get('text', '').strip()
                if step_text and step_text not in seen_texts:
                    unique_steps.append(step)
                    seen_texts.add(step_text)
            
            if not unique_steps:
                run = paragraph.add_run("No steps provided.")
                run.font.name = 'Altone Trial'
                run.font.size = Pt(11)
            else:
                for step_idx, step in enumerate(unique_steps):
                    step_para = cell.add_paragraph()
                    step_para.paragraph_format.left_indent = Pt(10)
                    step_text = step['text']
                    run = step_para.add_run(f"{step_text}")
                    run.font.name = 'Altone Trial'
                    run.font.size = Pt(11)
                    run.font.bold = False
                    screenshot_path = None
                    if image_map and step.get('screenshot'):
                        screenshot_path = find_image_for_step(step['screenshot'], image_map, display_sr_no, step_idx+1)
                    if screenshot_path and os.path.exists(screenshot_path):
                        print(f"[DEBUG] Adding image from ZIP: {screenshot_path}")
                        screenshot_para = cell.add_paragraph()
                        screenshot_para.paragraph_format.left_indent = Pt(20)
                        screenshot_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        screenshot_para.add_run().add_picture(screenshot_path, width=Inches(5))
                    else:
                        print(f"  -> No image found for step {step_idx+1}, inserting missing message.")
                        missing_para = cell.add_paragraph()
                        missing_para.paragraph_format.left_indent = Pt(20)
                        missing_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = missing_para.add_run(f"[Screenshot missing: {step.get('screenshot')}]")
                        run.font.name = 'Altone Trial'
                        run.font.size = Pt(10)
                        run.font.italic = True
                    if step_idx < len(steps) - 1:
                        spacer = cell.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(6)
    
    # Row for Associated CVEs
    row = table.rows[len(row_labels) - 2]
    cell = row.cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Associated CVEs: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(font_r, font_g, font_b)
    cves = vulnerability.get('associated_cves', [])
    if cves:
        for cve in cves:
            para = cell.add_paragraph(f"• {cve}")
            para.paragraph_format.left_indent = Pt(20)
            for run in para.runs:
                run.font.name = 'Altone Trial'
                run.font.size = Pt(11)
                run.font.bold = False
    else:
        run = paragraph.add_run("N/A")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
    
    # Add Reference Link after Associated CVEs (only if it exists and is not empty)
    reference_link = vulnerability.get('reference_link', '')
    if reference_link and reference_link.strip():
        # Add a new paragraph for the reference link
        ref_para = cell.add_paragraph()
        ref_para.paragraph_format.space_before = Pt(6)
        ref_run = ref_para.add_run("Reference Link: ")
        ref_run.font.name = 'Altone Trial'
        ref_run.font.size = Pt(12)
        ref_run.font.bold = True
        ref_run.font.color.rgb = RGBColor(font_r, font_g, font_b)
        
        # Add the actual link
        link_para = cell.add_paragraph(reference_link)
        link_para.paragraph_format.left_indent = Pt(20)
        link_para.paragraph_format.space_before = Pt(3)
        for run in link_para.runs:
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color for links
    
    # Row for PoC steps
    row = table.rows[len(row_labels) - 1]
    cell = row.cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Proof of Concept / Steps to Reproduce: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(font_r, font_g, font_b)
    steps = vulnerability.get('steps_with_screenshots', [])
    if not steps:
        run = paragraph.add_run("No steps provided.")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
    else:
        for step_idx, step in enumerate(steps):
            step_para = cell.add_paragraph()
            step_para.paragraph_format.left_indent = Pt(10)
            step_text = step['text']
            run = step_para.add_run(f"{step_text}")
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = False
            screenshot_path = None
            if image_map and step.get('screenshot'):
                screenshot_path = find_image_for_step(step['screenshot'], image_map, display_sr_no, step_idx+1)
            if screenshot_path and os.path.exists(screenshot_path):
                screenshot_para = cell.add_paragraph()
                screenshot_para.paragraph_format.left_indent = Pt(20)
                screenshot_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                screenshot_para.add_run().add_picture(screenshot_path, width=Inches(5))
            else:
                missing_para = cell.add_paragraph()
                missing_para.paragraph_format.left_indent = Pt(20)
                missing_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = missing_para.add_run(f"[Screenshot missing: {step.get('screenshot')}]")
                run.font.name = 'Altone Trial'
                run.font.size = Pt(10)
                run.font.italic = True
            if step_idx < len(steps) - 1:
                spacer = cell.add_paragraph()
                spacer.paragraph_format.space_after = Pt(6)
    
    return table

def create_vulnerability_table_with_poc(doc, vulnerability, display_sr_no=None, image_map=None):
    # Create the main vulnerability table as before
    table = create_vulnerability_table(doc, vulnerability, display_sr_no, image_map)
    # Avoid printing PoC steps twice: they are already inside the table under the PoC row
    return table

# Updated function signature to accept amendment_log and assessment details
def generate_word_report(
    vulnerabilities, doc_control_data, template_path, output_path, client_name="Client", image_map=None,
    vulnerability_file_location=None,
    company_logo=None,
    tpl_instance=None
):
    try:
        tpl = tpl_instance if tpl_instance is not None else DocxTemplate(template_path)
        # Generate document ID in the format C-XX-IP-Q1-2025
        current_year = datetime.now().year
        current_quarter = (datetime.now().month - 1) // 3 + 1
        client_code = client_name.upper()[:2]  # Get first two letters of client name
        doc_id = f"C-{client_code}-IP-Q{current_quarter}-{current_year}"
        # Get assessment details and amendment log from parsed doc_control_data
        assessment_scope = doc_control_data.get('assessment_scope', '')
        assessment_start_date = doc_control_data.get('assessment_start_date', '')
        assessment_end_date = doc_control_data.get('assessment_end_date', '')
        assessment_note = doc_control_data.get('assessment_note', '')
        amendment_log = doc_control_data.get('amendment_log', [])
        # Format the date range string
        date_range = f"{assessment_start_date} to {assessment_end_date}"
        print("Template variables being passed:")
        print(f"assessment_scope: {assessment_scope}")
        print(f"date_range: {date_range}")
        print(f"document_id: {doc_id}")
        print(f"amendment_log (count): {len(amendment_log)}") # Debug print
        # Build steps for each vulnerability (text only, not InlineImage)
        for vuln in vulnerabilities:
            vuln['steps'] = [
                {'text': step['text'], 'image': f"[Screenshot missing: {step['screenshot']}]"}
                for step in vuln['steps_with_screenshots']
            ]
        # Remove ip_list and ip_totals from context (no docxtpl table for IP summary)
        context = {
            'client_name': client_name,
            'assessment_scope': assessment_scope,
            'date_range': date_range,
            'document_id': doc_id,
            'amendment_log': amendment_log,
            'reviewed_by': doc_control_data.get('reviewed_by', ''),
            'authorized_by': doc_control_data.get('authorized_by', ''),
            'report_date': datetime.now().strftime('%d %B %Y'),
            'assessment_note': assessment_note,
            'vulnerabilities': vulnerabilities,
            'company_logo': company_logo
        }
        print("Full template context:")
        print(context)
        tpl.render(context)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tpl.save(tmp.name)
            tmp_path = tmp.name
        # Open with python-docx for dynamic PoC steps/images
        doc = Document(tmp_path)
        # Generate IP summary table from actual vulnerability data
        ip_summary_list = generate_ip_summary_from_vulnerabilities(vulnerabilities)
        insert_ip_summary_table_exact(doc, ip_summary_list)
        
        # Remove extra blank paragraphs or page breaks after the IP summary table
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            p = doc.paragraphs[-1]
            p._element.getparent().remove(p._element)
        # Add a single page break if the last paragraph is not already a page break
        if not (doc.paragraphs and doc.paragraphs[-1].text.strip() == ''):
            doc.add_page_break()

        # Sort vulnerabilities by severity order: Critical, High, Medium, Low, Informational, No Vulnerability
        def get_severity_order(vuln):
            severity = vuln.get('severity', '').strip()
            # Map severities to order numbers
            severity_map = {
                'Critical': 0,
                'High': 1, 
                'Medium': 2,
                'Low': 3,
                'Informational': 4,
                'No Vulnerability': 5
            }
            return severity_map.get(severity, 6)  # Unknown gets highest number
        
        try:
            vulnerabilities.sort(key=get_severity_order)
        except Exception as e:
            print(f"Error sorting vulnerabilities: {e}")
            pass

        # Define chart variables before try block
        severity_counts = count_vulnerabilities_by_severity(vulnerabilities)
        # Change 'Unknown' to 'No Vulnerability' in chart code
        bar_chart_path = os.path.join(tempfile.gettempdir(), "bar_chart.png")
        donut_chart_path = os.path.join(tempfile.gettempdir(), "donut_chart.png")

        # --- Insert severity charts after cleanup, before vulnerabilities ---
        try:
            # Use larger figure size for both charts
            def patched_create_severity_bar_chart(severity_counts, output_path):
                cleaned_counts = {}
                for k, v in severity_counts.items():
                    key = k if k in severity_order else 'No Vulnerability'
                    if not key or str(key).strip() == '':
                        key = 'No Vulnerability'
                    cleaned_counts[key] = cleaned_counts.get(key, 0) + v
                severities = list(cleaned_counts.keys())
                counts = list(cleaned_counts.values())
                colors = [
                    '#990000' if s == 'Critical' else
                    '#FF0000' if s == 'High' else
                    '#FFCC00' if s == 'Medium' else
                    '#009933' if s == 'Low' else
                    '#3399CC' if s == 'Informational' else
                    '#808080'  # No Vulnerability
                    for s in severities
                ]
                if len(severities) == 0:
                    # Nothing to plot; create an empty placeholder image
                    plt.figure(figsize=(8.5, 4.5))
                    plt.text(0.5, 0.5, 'No data', ha='center', va='center')
                    plt.axis('off')
                    plt.savefig(output_path)
                    plt.close()
                    return output_path
                sorted_severities, sorted_counts, sorted_colors = zip(*sorted(
                    zip(severities, counts, colors),
                    key=lambda x: severity_order.get(x[0], 5)
                ))
                x_pos = np.arange(len(sorted_severities))
                plt.figure(figsize=(8.5, 4.5))
                bars = plt.bar(x_pos, sorted_counts, color=sorted_colors, width=0.6)
                plt.xticks(x_pos, sorted_severities, rotation=10, ha='center', fontsize=12)
                plt.ylabel('Number of Vulnerabilities', fontsize=14)
                plt.title('Total Vulnerabilities Found', fontsize=16)
                for i, bar in enumerate(bars):
                    height = bar.get_height()
                    plt.text(bar.get_x() + bar.get_width()/2, height/2, str(int(height)),
                             ha='center', va='center', fontsize=14, fontweight='bold', color='white')
                plt.tight_layout()
                plt.savefig(output_path)
                plt.close()
                return output_path
            def patched_create_severity_donut_chart(severity_counts, output_path):
                cleaned_counts = {}
                for k, v in severity_counts.items():
                    key = k if k in severity_order else 'No Vulnerability'
                    if not key or str(key).strip() == '':
                        key = 'No Vulnerability'
                    cleaned_counts[key] = cleaned_counts.get(key, 0) + v
                severities = list(cleaned_counts.keys())
                counts = list(cleaned_counts.values())
                total_findings = sum(counts)
                color_map = {severity: get_severity_colors(severity)['row1_bg'] for severity in severity_order.keys()}
                color_map['No Vulnerability'] = '#808080'
                if len(severities) == 0:
                    plt.figure(figsize=(8.5, 4.5))
                    plt.text(0.5, 0.5, 'No data', ha='center', va='center')
                    plt.axis('off')
                    plt.savefig(output_path)
                    plt.close()
                    return output_path
                sorted_severities, sorted_counts = zip(*sorted(cleaned_counts.items(), key=lambda item: severity_order.get(item[0], 5)))
                sorted_colors = [color_map.get(s, '#808080') for s in sorted_severities]
                fig, ax = plt.subplots(figsize=(8.5, 4.5))
                wedges, texts, autotexts = ax.pie(
                    sorted_counts,
                    colors=sorted_colors, 
                    wedgeprops=dict(width=0.4, edgecolor='white'),
                    startangle=90,
                    autopct=lambda pct: f'{pct:.0f}%',
                    pctdistance=0.75,
                    textprops={'fontsize': 10, 'weight': 'bold', 'color': 'white'}
                )
                centre_circle = plt.Circle((0, 0), 0.60, fc='white')
                fig.gca().add_artist(centre_circle)
                ax.text(
                    0, 0,
                    f'Total\nFindings\n{total_findings}',
                    ha='center',
                    va='center',
                    fontsize=12,
                    weight='bold',
                    color='black'
                )
                legend_labels = [f'{s} ({c})' for s, c in zip(sorted_severities, sorted_counts)]
                ax.legend(
                    wedges,
                    legend_labels,
                    title="Severity Levels",
                    loc="center",
                    bbox_to_anchor=(0.5, -0.1),
                    ncol=len(sorted_severities),
                    fontsize=10,
                    title_fontsize=12,
                    frameon=False
                )
                plt.title(
                    'Overall Vulnerabilities Identified',
                    fontsize=14,
                    weight='bold',
                    pad=20,
                    color='#660099'
                )
                ax.axis('equal')
                plt.tight_layout()
                plt.savefig(output_path, bbox_inches='tight')
                plt.close()
                return output_path
            patched_create_severity_bar_chart(severity_counts, bar_chart_path)
            print(f"Bar chart created at: {bar_chart_path}, exists: {os.path.exists(bar_chart_path)}")
            patched_create_severity_donut_chart(severity_counts, donut_chart_path)
            print(f"Donut chart created at: {donut_chart_path}, exists: {os.path.exists(donut_chart_path)}")
            doc.add_picture(bar_chart_path, width=Inches(6))
            doc.add_paragraph()  # Add a blank line for spacing
            doc.add_picture(donut_chart_path, width=Inches(6))
            doc.add_page_break()
        except Exception as e:
            print(f"Error inserting charts: {e}")
            traceback.print_exc()
            raise HTTPException(status_code=500, detail=f"Chart error: {e}")

        # Sort vulnerabilities by severity order: Critical, High, Medium, Low, Informational, No Vulnerability
        def get_severity_order(vuln):
            if vuln.get('is_no_vuln_box'):
                return 5  # No Vulnerability at the end
            
            # Get severity from the vulnerability data
            severity = vuln.get('severity', '').strip()
            if not severity:
                # Try to get severity from CVSS score
                cvss = vuln.get('cvss')
                if cvss is not None and cvss != '' and str(cvss).lower() != 'nan':
                    try:
                        cvss_float = float(cvss)
                        if cvss_float >= 9.0:
                            severity = 'Critical'
                        elif cvss_float >= 7.0:
                            severity = 'High'
                        elif cvss_float >= 4.0:
                            severity = 'Medium'
                        elif cvss_float >= 0.1:
                            severity = 'Low'
                        else:
                            severity = 'Informational'
                    except (ValueError, TypeError):
                        severity = 'Informational'
                else:
                    severity = 'Informational'
            
            # Map severities to order numbers
            severity_map = {
                'Critical': 0,
                'High': 1, 
                'Medium': 2,
                'Low': 3,
                'Informational': 4,
                'No Vulnerability': 5
            }
            return severity_map.get(severity, 6)  # Unknown gets highest number
        
        # Sort vulnerabilities by severity order
        try:
            vulnerabilities.sort(key=get_severity_order)
        except Exception as e:
            print(f"Error sorting vulnerabilities by severity: {e}")
            # Fallback to original sorting
            vulnerabilities.sort(key=lambda v: (str(v.get('ip','')), v.get('row_order', 0)))
        
        # Debug: Print vulnerabilities after sorting
        print(f"After sorting - Total vulnerabilities: {len(vulnerabilities)}")
        for i, vuln in enumerate(vulnerabilities):
            print(f"  {i+1}. SR: {vuln['sr_no']}, Name: {vuln['name']}, CVSS: {vuln.get('cvss')}, is_no_vuln_box: {vuln.get('is_no_vuln_box')}")
        # Programmatically create and insert the summary table
        create_summary_table(doc, vulnerabilities)
        # Attempt to set default font for the document body
        try:
            style = doc.styles['Normal']
            style.font.name = 'Altone Trial'
        except KeyError:
            print("Warning: 'Normal' style not found in template. Cannot set default font.")
        except Exception as e:
            print(f"Warning: Could not set default font on 'Normal' style: {e}")
        # Apply Altone Trial font to all paragraphs and runs in the document body (additional pass)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Altone Trial'
        for table in doc.tables:
             for row in table.rows:
                 for cell in row.cells:
                     for paragraph in cell.paragraphs:
                         for run in paragraph.runs:
                             run.font.name = 'Altone Trial'
        # Assign display_sr_no to each vulnerability BEFORE generating tables/PoC steps
        for idx, vuln in enumerate(vulnerabilities, 1):
            vuln['display_sr_no'] = f"VUL-{idx:03d}"

        # Debug: Print associated_cves for each vulnerability after grouping
        for vuln in vulnerabilities:
            print(f"[DEBUG] {vuln['display_sr_no']} - Associated CVEs: {vuln.get('associated_cves')}")

        # Insert vulnerability table and PoC steps/images for each vulnerability using python-docx only
        print(f"Processing {len(vulnerabilities)} vulnerabilities for document insertion:")
        for i, vuln in enumerate(vulnerabilities):
            print(f"  Processing vulnerability {i+1}/{len(vulnerabilities)}: SR: {vuln['sr_no']}, Name: {vuln['name']}, is_no_vuln_box: {vuln.get('is_no_vuln_box')}")
            if vuln.get('is_no_vuln_box'):
                print(f"    Creating NO VULNERABILITY box for {vuln['display_sr_no']}")
                create_no_vuln_box(doc, vuln, vuln['display_sr_no'], image_map)
            else:
                print(f"    Creating vulnerability table for {vuln['display_sr_no']}")
                create_vulnerability_table(doc, vuln, vuln['display_sr_no'], image_map)
            # Only add a page break if not the last vulnerability
            if i < len(vulnerabilities) - 1:
                doc.add_page_break()
        # Add a page break and the SUMMARY OF FINDINGS & CONCLUSION section as the last page
        doc.add_page_break()
        # Add two blank lines after the SUMMARY OF FINDINGS & CONCLUSION: heading
        doc.add_paragraph()
        doc.add_paragraph()
        # Heading as a single paragraph with a line break for TOC compatibility
        summary_heading = doc.add_paragraph()
        run = summary_heading.add_run("SUMMARY OF FINDINGS &\nCONCLUSION:")
        doc.add_paragraph()
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = 'Altone Trial'
        run.font.color.rgb = RGBColor(106, 68, 154)
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set style to Heading 1 if available for TOC
        try:
            summary_heading.style = doc.styles['Heading 1']
        except Exception:
            pass
        # Paragraph 1
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_after = Pt(24)  # Double spacing after
        run = p1.add_run(
            "Finally, it must be remembered that security is an ongoing process, and that this report will provide an idea of the current vulnerabilities we were able to detect. There is no guarantee that new vulnerabilities will not be found and exploited in the future."
        )
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Paragraph 2
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.space_after = Pt(24)
        run = p2.add_run("The assessment was only possible because ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("joint support & coordination")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" from the information security team of organization for ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("sharing & coordinating")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" during the assessment period. It is advised to refer the ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("Technical Report")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" for understanding in-depth of vulnerabilities that were discovered by technical team of ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("CyberSmithSECURE Pvt. Ltd.")
        run.bold = True
        run.underline = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Paragraph 3
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p3.paragraph_format.space_after = Pt(24)
        run = p3.add_run("The Security Researchers of the ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p3.add_run("CyberSmithSECURE")
        run.bold = True
        run.underline = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p3.add_run(" performed Vulnerability Testing. We jointly recommend that all suggested measures in this document be performed to ensure the overall security of the target device. The following targeted sectors were identified by the security researchers for the scope of this testing.")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Paragraph 4 (last line, bold and purple)
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p4.paragraph_format.space_after = Pt(24)
        run = p4.add_run("We thank internal Information Security team for their support & cooperation during the time of assessment.")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(106, 68, 154)
        # Copy header/footer from first section to all sections to ensure document ID appears on all pages
        if len(doc.sections) > 1:
            first_section = doc.sections[0]
            for section in doc.sections[1:]:
                section.header.is_linked_to_previous = False
                section.footer.is_linked_to_previous = False
                # Clear and copy header
                section.header._element.clear_content()
                for element in first_section.header._element:
                    section.header._element.append(element)
                # Clear and copy footer
                section.footer._element.clear_content()
                for element in first_section.footer._element:
                    section.footer._element.append(element)
        doc.save(output_path)
        return output_path
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating Word report: {str(e)}")

# Function to check if file exists in uploads directory
def check_screenshots_directory(vulnerability, screenshots_dir="uploads/screenshots"):
    """
    Check if screenshot directory exists for a vulnerability, create it if not
    Returns the path to the screenshots directory
    """
    try:
        # Create directory structure based on severity and vulnerability name
        severity = vulnerability.get('severity', 'Unknown')
        vuln_name = vulnerability.get('name', 'Unknown').replace(' ', '_').replace('/', '_')
        
        # Create path: uploads/screenshots/[Severity]/[Vulnerability_Name]
        screenshot_path = os.path.join(screenshots_dir, severity, vuln_name)
        os.makedirs(screenshot_path, exist_ok=True)
        
        return screenshot_path
    except Exception as e:
        print(f"Error creating screenshots directory: {e}")
        return None

def find_image_case_insensitive(directory, filename):
    """
    Search for an image file in the directory, case-insensitive, and try common extensions.
    """
    base, ext = os.path.splitext(filename)
    possible_exts = ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']
    candidates = [base + e for e in possible_exts]
    files = os.listdir(directory)
    for candidate in candidates:
        for f in files:
            if f.lower() == candidate.lower():
                return os.path.join(directory, f)
    return None

def index_images_recursively(root_dir):
    image_map = {}
    for dirpath, _, filenames in os.walk(root_dir):
        for fname in filenames:
            if fname.lower().endswith((".png", ".jpg", ".jpeg")):
                norm = normalize_filename(fname)
                full_path = os.path.join(dirpath, fname)
                with open(full_path, 'rb') as f:
                    data = f.read()
                    h = hashlib.md5(data).hexdigest()
                    print(f"[IMAGE MAP] {fname} => {norm} | Size: {len(data)} | Hash: {h}")
                    f.seek(0)
                image_map[norm] = full_path
    print(f"[DEBUG] Final image map: {list(image_map.keys())}")
    return image_map

def insert_steps_with_images(doc, vulnerability, image_map):
    doc.add_paragraph(f"Proof of Concept / Steps to Reproduce for {vulnerability['name']}").runs[0].bold = True
    for idx, step in enumerate(vulnerability.get("steps_with_screenshots", []), 1):
        doc.add_paragraph(f"Step {idx}: {step['text']}")
        norm = normalize_filename(step['screenshot'])
        img_path = image_map.get(norm)
        print(f"[PoC] Step {idx} for '{vulnerability['name']}' → Looking for '{norm}' → {img_path}")
        if img_path and os.path.exists(img_path):
            with open(img_path, 'rb') as f:
                img_data = f.read()
                print(f"[PoC] FOUND: {img_path} | Hash: {hashlib.md5(img_data).hexdigest()} | Size: {len(img_data)}")
                f.seek(0)
                doc.add_paragraph().add_run().add_picture(f, width=Inches(4))
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]")

def get_script_dir():
    return os.path.dirname(os.path.realpath(__file__))

async def auto_upload_to_dashboard(vulnerabilities, client_name, db, request):
    """
    Automatically upload parsed vulnerability data to dashboard.
    Creates a dashboard dataset with the parsed data for monitoring.
    """
    try:
        # Import dashboard models
        from db import DashboardDataset, AuditLog
        
        # Create a CSV-like data structure for dashboard
        dashboard_data = []
        for vuln in vulnerabilities:
            if not vuln.get('is_no_vuln_box'):  # Skip "No Vulnerability" entries
                dashboard_data.append({
                    'IP': vuln.get('ip', ''),
                    'Vulnerability': vuln.get('name', ''),
                    'Severity': vuln.get('severity', ''),
                    'CVSS': vuln.get('cvss', ''),
                    'ReportedBy': vuln.get('reported_by', ''),
                    'Project': vuln.get('project', ''),
                    'SeverityLevel': vuln.get('severity_level', ''),
                    'Description': vuln.get('description', ''),
                    'Impact': vuln.get('impact', ''),
                    'Remediation': vuln.get('remediation', ''),
                    'SrNo': vuln.get('sr_no', ''),
                    'AssociatedCVEs': ', '.join(vuln.get('associated_cves', [])),
                    'ReferenceLink': vuln.get('reference_link', '')
                })
        
        if not dashboard_data:
            print("No vulnerability data to upload to dashboard")
            return
        
        # Create a temporary CSV file
        import csv
        import tempfile
        import pandas as pd
        
        # Create DataFrame and save as CSV
        df = pd.DataFrame(dashboard_data)
        temp_csv_path = os.path.join(tempfile.gettempdir(), f"dashboard_upload_{datetime.now().strftime('%Y%m%d%H%M%S')}.csv")
        df.to_csv(temp_csv_path, index=False)
        
        # Create dashboard dataset entry
        user = request.session.get('user') or {} if request else {}
        ds = DashboardDataset(
            title=f"{client_name} VAPT Data - {datetime.now().strftime('%Y-%m-%d')}",
            project_name=client_name,
            file_path=temp_csv_path,
            uploaded_by_email=user.get('email', 'system'),
            uploaded_by_name=user.get('name', 'System Auto-Upload')
        )
        db.add(ds)
        
        # Add audit log
        db.add(AuditLog(
            user_email=user.get('email', 'system'),
            user_name=user.get('name', 'System Auto-Upload'),
            action='auto-dashboard-upload',
            metadata_json=json.dumps({
                'client_name': client_name,
                'vulnerability_count': len(dashboard_data),
                'source': 'type2-report-generation'
            }),
            ip_address=request.client.host if request and request.client else None,
            user_agent=request.headers.get('user-agent') if request else None
        ))
        
        db.commit()
        db.refresh(ds)
        
        print(f"Successfully auto-uploaded {len(dashboard_data)} vulnerabilities to dashboard (ID: {ds.id})")
        
    except Exception as e:
        print(f"Error in auto_upload_to_dashboard: {e}")
        raise

@app.post("/generate-report/")
async def generate_report(
    file: UploadFile = File(...),  # Vulnerability file
    assessment_file: UploadFile = File(...),  # Assessment details file
    poc_images: UploadFile = File(None),  # POC screenshots ZIP file (optional)
    client_name: str = Form("Client"),
    company_logo: UploadFile = File(None),  # <-- Add company_logo as optional
    request: Request = None,
    db: Session = Depends(get_db)
):
    try:
        print(f"Starting report generation - Version {VERSION}")
        # Create directories if they don't exist
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)
        screenshots_dir = os.path.join(UPLOAD_DIR, "screenshots")
        if not os.path.exists(screenshots_dir):
            os.makedirs(screenshots_dir)
        # Save the uploaded vulnerability file
        vulnerability_file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(vulnerability_file_location, "wb") as f:
            f.write(await file.read())
        print(f"Saved uploaded vulnerability file to: {vulnerability_file_location}")
        # Save the uploaded assessment file
        assessment_file_location = os.path.join(UPLOAD_DIR, assessment_file.filename)
        with open(assessment_file_location, "wb") as f:
            f.write(await assessment_file.read())
        print(f"Saved uploaded assessment file to: {assessment_file_location}")
        # No dashboard ingestion here; dashboard is a separate feature
        # Handle POC images ZIP (accept only .zip)
        poc_images_dir = None
        if poc_images is not None:
            poc_name = (poc_images.filename or '').lower()
            if not poc_name.endswith('.zip'):
                raise HTTPException(status_code=400, detail="POC screenshots must be uploaded as a .zip file")
            tmpdirname = tempfile.mkdtemp()
            zip_path = os.path.join(tmpdirname, poc_images.filename)
            with open(zip_path, 'wb') as f:
                f.write(await poc_images.read())
            with ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(tmpdirname)
            print('--- Extracted ZIP Directory Tree ---')
            for dirpath, dirnames, filenames in os.walk(tmpdirname):
                print(f"DIR: {dirpath}")
                print(f"  DIRNAMES: {dirnames}")
                print(f"  FILENAMES: {filenames}")
            print('--- End of Directory Tree ---')
            for entry in os.listdir(tmpdirname):
                entry_path = os.path.join(tmpdirname, entry)
                if os.path.isdir(entry_path):
                    poc_images_dir = entry_path
                    break
            if poc_images_dir is None:
                poc_images_dir = tmpdirname
        # --- Handle company logo upload or fallback ---
        logo_path = None
        logo_filename = None
        if company_logo is not None:
            logo_filename = f"uploaded_logo.png"
            logo_path = os.path.join(UPLOAD_DIR, logo_filename)
            with open(logo_path, "wb") as f:
                f.write(await company_logo.read())
            print(f"Saved uploaded company logo to: {logo_path}")
        else:
            # Use default logo
            logo_path = os.path.join(os.path.dirname(__file__), "default_logo.jpg")
            if not os.path.exists(logo_path):
                # fallback to uploads dir if not found
                logo_path = os.path.join(UPLOAD_DIR, "default_logo.jpg")
            print(f"Using default logo: {logo_path}")
        # Parse the uploaded vulnerability file
        vulnerabilities = parse_vulnerabilities_excel(vulnerability_file_location)
        vulnerabilities = group_vulnerabilities(vulnerabilities)
        doc_control_data = parse_doc_control_excel(assessment_file_location)
        
        # Make template path absolute
        template_path = os.path.join(get_script_dir(), "CSS-NEW_Technical_Sample_Report_Automated_v2.1.docx")

        report_filename = f"{client_name} VAPT Report {datetime.now().strftime('%Y-%d-%m')}.docx"
        output_path = os.path.join(UPLOAD_DIR, report_filename)
        image_map = None
        if poc_images_dir:
            image_map = index_images_from_poc_zip(poc_images_dir)
            if image_map:
                print('Available images in ZIP:')
                for k in image_map.keys():
                    print('  ', k)
        print("After grouping:")
        for v in vulnerabilities:
            print(f"SR: {v['sr_no']}, Name: {v['name']}, is_no_vuln_box: {v['is_no_vuln_box']}")
        # --- Prepare InlineImage for docxtpl ---
        from docxtpl import DocxTemplate, InlineImage
        from docx.shared import Inches
        tpl = DocxTemplate(template_path)
        tpl_logo = None
        try:
            if logo_path and os.path.exists(logo_path):
                optimal_width = calculate_optimal_logo_size(logo_path, max_width=2.3, max_height=2.0)
                tpl_logo = InlineImage(tpl, logo_path, width=Inches(optimal_width))
            else:
                default_logo_path = os.path.join(os.path.dirname(__file__), "default_logo.png")
                if os.path.exists(default_logo_path):
                    optimal_width = calculate_optimal_logo_size(default_logo_path, max_width=2.3, max_height=2.0)
                    tpl_logo = InlineImage(tpl, default_logo_path, width=Inches(optimal_width))
                else:
                    tpl_logo = None
        except Exception as e:
            print(f"Error loading logo as InlineImage: {e}")
            tpl_logo = None
        # --- Pass logo to report generator ---
        output_path = generate_word_report(
            vulnerabilities, doc_control_data, template_path, output_path, client_name, image_map,
            vulnerability_file_location,
            company_logo=tpl_logo,
            tpl_instance=tpl
        )
        print(f"Report generation completed successfully - Version {VERSION}")
        # Audit log
        try:
            user = request.session.get('user') if request else None
            db.add(AuditLog(
                user_email=(user or {}).get('email'),
                user_name=(user or {}).get('name'),
                action='generate-report-type2',
                metadata_json=json.dumps({'vuln_file': file.filename, 'assessment_file': assessment_file.filename}),
                ip_address=request.client.host if request and request.client else None,
                user_agent=request.headers.get('user-agent') if request else None
            ))
            db.commit()
        except Exception:
            pass
        
        # Automatically upload parsed data to dashboard
        try:
            await auto_upload_to_dashboard(vulnerabilities, client_name, db, request)
        except Exception as e:
            print(f"Warning: Failed to auto-upload to dashboard: {e}")
            # Don't fail the report generation if dashboard upload fails
        
        return FileResponse(
            path=output_path,
            filename=report_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

def count_vulnerabilities_by_severity(vulnerabilities):
    """
    Counts the number of vulnerabilities for each severity level.
    """
    severity_counts = Counter()
    for vuln in vulnerabilities:
        severity_counts[vuln.get('severity', 'Unknown')] += 1
        
    return severity_counts

def create_severity_bar_chart(severity_counts, output_path):
    # Map empty or unknown severities to 'Unknown'
    cleaned_counts = {}
    for k, v in severity_counts.items():
        key = k if k in severity_order else 'No Vulnerability'
        if not key or str(key).strip() == '':
            key = 'No Vulnerability'
        cleaned_counts[key] = cleaned_counts.get(key, 0) + v
    
    # Ensure we have data to plot
    if not cleaned_counts or sum(cleaned_counts.values()) == 0:
        # Create a placeholder chart
        plt.figure(figsize=(10, 7))
        plt.text(0.5, 0.5, 'No vulnerability data available', 
                ha='center', va='center', fontsize=16, transform=plt.gca().transAxes)
        plt.title('Total Vulnerabilities Found', fontsize=16)
        plt.axis('off')
        plt.tight_layout()
        plt.savefig(output_path)
        plt.close()
        return output_path
    
    severities = list(cleaned_counts.keys())
    counts = list(cleaned_counts.values())
    colors = [
        '#990000' if s == 'Critical' else
        '#FF0000' if s == 'High' else
        '#FFCC00' if s == 'Medium' else
        '#009933' if s == 'Low' else
        '#3399CC' if s == 'Informational' else
        '#808080'  # No Vulnerability
        for s in severities
    ]
    sorted_severities, sorted_counts, sorted_colors = zip(*sorted(
        zip(severities, counts, colors),
        key=lambda x: severity_order.get(x[0], 5)
    ))
    
    # Create 3D-like bar chart with enhanced styling
    fig = plt.figure(figsize=(12, 8))
    ax = fig.add_subplot(111)
    
    x_pos = np.arange(len(sorted_severities))
    
    # Create bars with 3D effect using gradient colors and shadows
    bars = ax.bar(x_pos, sorted_counts, color=sorted_colors, width=0.7, 
                  edgecolor='black', linewidth=1.5, alpha=0.8)
    
    # Add shadow effect for 3D appearance
    for i, bar in enumerate(bars):
        # Create shadow
        shadow = ax.bar(x_pos[i] + 0.02, sorted_counts[i], color='black', 
                       width=0.7, alpha=0.3, zorder=0)
        # Bring main bar to front
        bar.set_zorder(2)
    
    # Enhanced styling
    ax.set_xticks(x_pos)
    ax.set_xticklabels(sorted_severities, rotation=15, ha='center', fontsize=12, fontweight='bold')
    ax.set_ylabel('Number of Vulnerabilities', fontsize=14, fontweight='bold')
    ax.set_title('Total Vulnerabilities Found', fontsize=18, fontweight='bold', pad=20)
    
    # Add grid for better readability
    ax.grid(True, alpha=0.3, linestyle='--')
    ax.set_axisbelow(True)
    
    # Place count labels inside the bars with enhanced styling
    for i, bar in enumerate(bars):
        height = bar.get_height()
        if height > 0:
            ax.text(bar.get_x() + bar.get_width()/2, height/2, str(int(height)),
                   ha='center', va='center', fontsize=14, fontweight='bold', 
                   color='white', bbox=dict(boxstyle="round,pad=0.3", facecolor='black', alpha=0.7))
    
    # Add total count at the top
    total_count = sum(sorted_counts)
    ax.text(0.5, 0.95, f'Total: {total_count}', transform=ax.transAxes, 
            ha='center', va='top', fontsize=16, fontweight='bold',
            bbox=dict(boxstyle="round,pad=0.5", facecolor='lightblue', alpha=0.8))
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    return output_path

def create_severity_donut_chart(severity_counts, output_path):
    # Map empty or unknown severities to 'Unknown'
    cleaned_counts = {}
    for k, v in severity_counts.items():
        key = k if k in severity_order else 'No Vulnerability'
        if not key or str(key).strip() == '':
            key = 'No Vulnerability'
        cleaned_counts[key] = cleaned_counts.get(key, 0) + v

    # Ensure we have data to plot
    if not cleaned_counts or sum(cleaned_counts.values()) == 0:
        # Create a placeholder chart
        fig, ax = plt.subplots(figsize=(10, 8))
        ax.text(0.5, 0.5, 'No vulnerability data available', 
                ha='center', va='center', fontsize=16, transform=ax.transAxes)
        ax.set_title('Overall Vulnerabilities Identified', fontsize=18, fontweight='bold')
        ax.axis('off')
        plt.tight_layout()
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()
        return output_path

    severities = list(cleaned_counts.keys())
    counts = list(cleaned_counts.values())
    total_findings = sum(counts)

    # Enhanced color palette with gradients
    color_map = {
        'Critical': '#8B0000',
        'High': '#DC143C', 
        'Medium': '#FFD700',
        'Low': '#32CD32',
        'Informational': '#4169E1',
        'No Vulnerability': '#808080'
    }

    # Sort data by severity order
    sorted_severities, sorted_counts = zip(*sorted(cleaned_counts.items(), key=lambda item: severity_order.get(item[0], 5)))
    sorted_colors = [color_map.get(s, '#808080') for s in sorted_severities]

    # Create the enhanced donut chart
    fig, ax = plt.subplots(figsize=(10, 8))
    
    # Create gradient effect by using multiple pie slices with different shades
    wedges, texts, autotexts = ax.pie(
        sorted_counts,
        colors=sorted_colors, 
        wedgeprops=dict(width=0.5, edgecolor='white', linewidth=2),
        startangle=90,
        autopct='',  # Remove percentage labels for cleaner look
        pctdistance=0.8,
        explode=[0.05 if count > 0 else 0 for count in sorted_counts]  # Slight separation for active slices
    )

    # Add shadow effect for 3D appearance
    for wedge in wedges:
        wedge.set_alpha(0.8)
    
    # Draw center circle with gradient effect
    centre_circle = plt.Circle((0, 0), 0.55, fc='white', edgecolor='gray', linewidth=2)
    fig.gca().add_artist(centre_circle)

    # Add inner shadow circle
    shadow_circle = plt.Circle((0, 0), 0.53, fc='lightgray', alpha=0.3)
    fig.gca().add_artist(shadow_circle)

    # Enhanced center text with better styling
    ax.text(0, 0.1, 'Total', ha='center', va='center', 
            fontsize=14, weight='bold', color='gray')
    ax.text(0, -0.1, f'Findings\n{total_findings}', ha='center', va='center',
            fontsize=18, weight='bold', color='black')

    # Enhanced legend with better styling
    legend_labels = [f'{s} ({c})' for s, c in zip(sorted_severities, sorted_counts) if c > 0]
    legend_elements = [plt.Rectangle((0,0),1,1, facecolor=color_map.get(s, '#808080'), 
                                   edgecolor='white', linewidth=1) 
                      for s, c in zip(sorted_severities, sorted_counts) if c > 0]
    
    ax.legend(legend_elements, legend_labels, title="Severity Levels",
              loc="center", bbox_to_anchor=(0.5, -0.15), 
              ncol=min(len(legend_labels), 3), fontsize=11, title_fontsize=13,
              frameon=True, fancybox=True, shadow=True)

    # Set the title with consistent styling
    plt.title("Overall Vulnerabilities Identified", fontsize=18, fontweight='bold', pad=20)
    
    plt.tight_layout()
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()
    return output_path

def index_images_from_poc_zip(root_dir):
    import re
    image_map = {}
    screenshots_root = None

    print(f"[DEBUG] Starting image indexing from: {root_dir}")

    # If the current directory is 'POC Screenshots*', use it directly (accept suffix digits)
    base = os.path.basename(root_dir).strip().lower()
    if base.startswith('poc screenshots'):
        screenshots_root = root_dir
        print(f"[DEBUG] Using current directory as screenshots root: {screenshots_root}")
    else:
        # Look for a folder that starts with 'POC Screenshots' inside the extracted ZIP
        print(f"[DEBUG] Searching for POC Screenshots folder in: {root_dir}")
        for dirpath, dirnames, _ in os.walk(root_dir):
            print(f"[DEBUG] Checking directory: {dirpath}, subdirs: {dirnames}")
            for d in dirnames:
                if d.strip().lower().startswith('poc screenshots'):
                    screenshots_root = os.path.join(dirpath, d)
                    print(f"[DEBUG] Found screenshots root: {screenshots_root}")
                    break
            if screenshots_root:
                break

    if not screenshots_root:
        print("[ERROR] 'POC Screenshots' folder not found in ZIP.")
        print(f"[DEBUG] Available directories in {root_dir}:")
        try:
            for item in os.listdir(root_dir):
                print(f"  - {item}")
        except:
            print("  - Cannot list directory")
        return image_map

    print(f"[DEBUG] Processing screenshots from: {screenshots_root}")

    # Expected structure: POC Screenshots*/<IP>/<Severity>/<VUL-XXX>/stepN.*
    # Normalize to handle Windows-extracted zips and spaces in folder names
    try:
        for ip_folder in os.listdir(screenshots_root):
            ip_path = os.path.join(screenshots_root, ip_folder)
            if not os.path.isdir(ip_path):
                continue
            print(f"[DEBUG] Processing IP folder: {ip_folder}")
            
            for severity in os.listdir(ip_path):
                severity_path = os.path.join(ip_path, severity)
                if not os.path.isdir(severity_path):
                    continue
                print(f"[DEBUG] Processing severity: {severity}")
                
                for vuln_id in os.listdir(severity_path):
                    vuln_path = os.path.join(severity_path, vuln_id)
                    if os.path.isdir(vuln_path) and vuln_id.lower().startswith("vul-"):
                        print(f"[DEBUG] Processing vulnerability: {vuln_id}")
                        
                        for fname in os.listdir(vuln_path):
                            if fname.lower().endswith(('.png', '.jpg', '.jpeg')):
                                print(f"[DEBUG] Found image: {fname}")
                                
                                # Support names like 'step1.png', 'Step 1.PNG', 'STEP_1.jpg'
                                m = re.search(r'step\s*[_-]?(\d+)', fname.lower())
                                if m:
                                    step_num = int(m.group(1))
                                else:
                                    # Fallback: use file order
                                    step_num = len([f for f in os.listdir(vuln_path) if f.lower().endswith(('.png', '.jpg', '.jpeg'))])
                                
                                # Create multiple possible keys for better matching
                                keys = [
                                    f"{vuln_id.strip().lower()}_step{step_num}",
                                    f"{vuln_id.strip().lower()}step{step_num}",
                                    f"step{step_num}",
                                    fname.lower().replace('.png', '').replace('.jpg', '').replace('.jpeg', ''),
                                    f"{vuln_id.strip().lower()}_{fname.lower().replace('.png', '').replace('.jpg', '').replace('.jpeg', '')}",
                                    # Add IP-based keys for specific IPs mentioned
                                    f"{ip_folder.lower()}_{vuln_id.strip().lower()}_step{step_num}",
                                    f"{ip_folder.lower()}_{severity.lower()}_{vuln_id.strip().lower()}_step{step_num}",
                                    # Add just the filename as key
                                    fname.lower(),
                                    # Add step number variations
                                    f"step{step_num}.png",
                                    f"step{step_num}.jpg",
                                    f"step{step_num}.jpeg"
                                ]
                                
                                full_path = os.path.join(vuln_path, fname)
                                
                                # Add all possible keys
                                for key in keys:
                                    image_map[key] = full_path
                                    print(f"[MAP] {key} → {full_path}")
                                
                                # Also add a debug entry showing the full path structure
                                debug_key = f"DEBUG_{ip_folder}_{severity}_{vuln_id}_{fname}"
                                image_map[debug_key] = full_path
    except Exception as e:
        print(f"[ERROR] Error processing images: {e}")
    
    print('--- Final Image Map Keys ---')
    for k in image_map:
        print(f"  {k}")
    print('--- End Image Map Keys ---')
    return image_map

def build_steps_with_images(steps_with_screenshots, image_map, tpl, sr_no):
    steps = []
    sr_no_key = sr_no.strip().lower() if sr_no else ''
    
    print(f"[POC DEBUG] Building steps for SR: {sr_no}, Key: {sr_no_key}")
    print(f"[POC DEBUG] Available image map keys: {list(image_map.keys())}")
    
    # Remove duplicate steps based on text content
    unique_steps = []
    seen_texts = set()
    for step in steps_with_screenshots:
        step_text = step.get('text', '').strip()
        if step_text and step_text not in seen_texts:
            unique_steps.append(step)
            seen_texts.add(step_text)
    
    for idx, step in enumerate(unique_steps, 1):
        screenshot_name = step.get('screenshot', '')
        print(f"[POC DEBUG] Processing step {idx}: {screenshot_name}")
        
        match = re.search(r'step(\d+)', screenshot_name.strip().lower())
        if match:
            step_num = int(match.group(1))
        else:
            step_num = idx
        
        # Try multiple possible keys for the image
        possible_keys = [
            f"{sr_no_key}_step{step_num}",
            f"{sr_no_key}step{step_num}",
            f"step{step_num}",
            screenshot_name.lower().replace('.png', '').replace('.jpg', '').replace('.jpeg', ''),
            f"{sr_no_key}_{screenshot_name.lower().replace('.png', '').replace('.jpg', '').replace('.jpeg', '')}",
            screenshot_name.lower() if screenshot_name else f"step{step_num}.png",
            f"step{step_num}.png",
            f"step{step_num}.jpg",
            f"step{idx}.png",
            f"step{idx}.jpg"
        ]
        
        img_path = None
        for key in possible_keys:
            if key in image_map:
                img_path = image_map[key]
                print(f"[POC DEBUG] Found image with key: {key}")
                break
        
        # If still no match, try to find any image for this vulnerability
        if not img_path and image_map:
            print(f"[POC DEBUG] No direct match, searching for partial matches...")
            for map_key, map_path in image_map.items():
                print(f"[POC DEBUG] Checking map key: {map_key}")
                if sr_no_key in map_key.lower() and ('step' in map_key.lower() or '.png' in map_key.lower() or '.jpg' in map_key.lower()):
                    img_path = map_path
                    print(f"[POC DEBUG] Found partial match: {map_key}")
                    break
        
        # Additional fallback: try to match by vulnerability ID pattern
        if not img_path and image_map:
            print(f"[POC DEBUG] Trying vulnerability ID pattern matching...")
            vuln_pattern = sr_no_key.replace('vul-', '').replace('vul', '')
            for map_key, map_path in image_map.items():
                if vuln_pattern in map_key.lower():
                    img_path = map_path
                    print(f"[POC DEBUG] Found pattern match: {map_key}")
                    break
        
        # Final fallback: try to find any image that contains the vulnerability ID
        if not img_path and image_map:
            print(f"[POC DEBUG] Final fallback: searching for any image with vulnerability ID...")
            for map_key, map_path in image_map.items():
                if sr_no_key in map_key.lower() or vuln_pattern in map_key.lower():
                    img_path = map_path
                    print(f"[POC DEBUG] Found final fallback match: {map_key}")
                    break
        
        # Last resort: if we have any images at all, use the first one
        if not img_path and image_map:
            print(f"[POC DEBUG] Last resort: using first available image...")
            first_key = list(image_map.keys())[0]
            img_path = image_map[first_key]
            print(f"[POC DEBUG] Using first available image: {first_key}")
        
        print(f"[POC IMAGE DEBUG] Step {idx}, Screenshot: {screenshot_name}, Keys tried: {possible_keys}, Found: {img_path}")
        if img_path and os.path.exists(img_path):
            print(f"[POC IMAGE FOUND] Using image: {img_path}")
            image = InlineImage(tpl, img_path, width=Inches(5))
        else:
            print(f"[POC IMAGE MISSING] No image found for step {idx}")
            image = f"[Screenshot missing: {screenshot_name}]"
        steps.append({'text': step['text'], 'image': image})
    return steps

def find_image_for_step(screenshot_name, image_map, sr_no, step_idx):
    screenshot_name_key = screenshot_name.strip().lower()
    sr_no_key = sr_no.strip().lower() if sr_no else ''
    match = re.search(r'step(\d+)', screenshot_name_key)
    if match:
        step_num = int(match.group(1))
    else:
        step_num = step_idx
    key = f"{sr_no_key}_step{step_num}"
    print(f"Looking for key: {key} in image_map")
    if key in image_map:
        print(f"Found image for {key}: {image_map[key]}")
    else:
        print(f"Image for {key} NOT FOUND")
    return image_map.get(key)

# After parsing vulnerabilities, ensure each has a display_sr_no
def assign_display_sr_no(vulnerabilities):
    for idx, vuln in enumerate(vulnerabilities, 1):
        if not vuln.get('display_sr_no'):
            sr_no = str(vuln.get('sr_no', f"VUL-{idx:03d}")).replace('VULN-', 'VUL-').strip()
            vuln['display_sr_no'] = sr_no
    return vulnerabilities 

def insert_poc_steps_section(doc, vulnerability, image_map):
    # Deprecated: PoC steps are rendered inside the vulnerability table.
    return

def find_placeholder_paragraph(doc, placeholder_text):
    # Search in paragraphs
    for p in doc.paragraphs:
        print("Checking paragraph:", repr(p.text))
        if placeholder_text.strip() in p.text.strip():
            print("Found placeholder in main body paragraph.")
            return p
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    print("Checking table cell paragraph:", repr(p.text))
                    if placeholder_text.strip() in p.text.strip():
                        print("Found placeholder in table cell.")
                        return p
    print("Placeholder not found in any paragraph or table cell.")
    return None

# Update insert_ip_summary_table_colored to use the new search function

def insert_ip_summary_table_colored(doc, ip_summary_list):
    print("Inserting IP summary table. Number of IPs:", len(ip_summary_list))
    placeholder_text = "_IP_SUMMARY_TABLE_PLACEHOLDER_"
    placeholder_paragraph = find_placeholder_paragraph(doc, placeholder_text)
    if not placeholder_paragraph:
        print("IP summary placeholder not found in template.")
        return
    # Define headers and their colors
    headers = [
        ("Sr. No.", "#4A77A8"), ("Hostname", "#4A77A8"), ("IP Type", "#4A77A8"), ("VAPT Status", "#4A77A8"),
        ("Critical", "#990000"), ("High", "#FF0000"), ("Medium", "#FFCC00"),
        ("Low", "#009933"), ("Informational", "#3399CC"), ("Total", "#6A449A")
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hdr_cells = table.rows[0].cells
    for i, (header, color) in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)
    # Data rows
    for idx, row in enumerate(ip_summary_list, 1):
        cells = table.add_row().cells
        # Ensure numeric columns are integers (no decimals)
        def as_int_str(v):
            try:
                if isinstance(v, str) and v.strip() == '':
                    return '0'
                return str(int(float(v)))
            except Exception:
                return str(v)
        values = [
            str(idx),
            str(row.get("Hostname", "")),
            str(row.get("IP Type", "")),
            str(row.get("VAPT Status", "")),
            as_int_str(row.get("Critical", 0)),
            as_int_str(row.get("High", 0)),
            as_int_str(row.get("Medium", 0)),
            as_int_str(row.get("Low", 0)),
            as_int_str(row.get("Informational", row.get("Information", 0))),
            as_int_str(row.get("Total", 0))
        ]
        for i, value in enumerate(values):
            cells[i].text = value
            if i == 4:
                color = "#990000"
            elif i == 5:
                color = "#FF0000"
            elif i == 6:
                color = "#FFCC00"
            elif i == 7:
                color = "#009933"
            elif i == 8:
                color = "#3399CC"
            elif i == 9:
                color = "#6A449A"
            elif i in [0, 1, 2, 3]:
                color = "#B7D1F2"
            else:
                color = "#FFFFFF"
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
            cells[i]._element.get_or_add_tcPr().append(shading_elm)
            for run in cells[i].paragraphs[0].runs:
                if i in [4, 5, 9]:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif i == 6:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif i == 7:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif i == 8:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(11)
    parent = placeholder_paragraph._p.getparent()
    parent.insert(parent.index(placeholder_paragraph._p), table._element)
    parent.remove(placeholder_paragraph._p)



    



    

def generate_ip_summary_from_vulnerabilities(vulnerabilities):
    """
    Generate IP summary table from actual vulnerability data instead of Excel columns.
    This ensures accurate totals calculation.
    """
    ip_data = {}
    
    for vuln in vulnerabilities:
        ip = vuln.get('ip', '').strip()
        if not ip:
            continue
            
        # Split multiple IPs if comma-separated
        ips = [ip.strip() for ip in ip.split(',')]
        
        for single_ip in ips:
            if single_ip not in ip_data:
                ip_data[single_ip] = {
                    'Hostname': single_ip,
                    'IP Type': 'External',  # Default, can be updated from Excel if needed
                    'VAPT Status': 'Completed',
                    'Critical': 0,
                    'High': 0,
                    'Medium': 0,
                    'Low': 0,
                    'Informational': 0,
                    'Total': 0
                }
            
            # Count by severity
            severity = vuln.get('severity', '').strip()
            if severity == 'Critical':
                ip_data[single_ip]['Critical'] += 1
            elif severity == 'High':
                ip_data[single_ip]['High'] += 1
            elif severity == 'Medium':
                ip_data[single_ip]['Medium'] += 1
            elif severity == 'Low':
                ip_data[single_ip]['Low'] += 1
            elif severity == 'Informational':
                ip_data[single_ip]['Informational'] += 1
            elif severity == 'No Vulnerability':
                ip_data[single_ip]['Informational'] += 1  # Treat as informational
    
    # Calculate totals for each IP
    for ip in ip_data:
        ip_data[ip]['Total'] = (
            ip_data[ip]['Critical'] + 
            ip_data[ip]['High'] + 
            ip_data[ip]['Medium'] + 
            ip_data[ip]['Low'] + 
            ip_data[ip]['Informational']
        )
    
    summary_list = list(ip_data.values())
    print(f"Generated IP summary from vulnerabilities: {len(summary_list)} IPs")
    print(f"Sample IP data: {summary_list[0] if summary_list else 'No data'}")
    return summary_list

def parse_ip_summary_table(file_path):
    """
    Parse the Excel file and extract the IP summary table (columns A–I).
    Returns a list of dicts per IP using the ORIGINAL column names that
    insert_ip_summary_table_exact() expects, and filters blank rows.
    """
    df = pd.read_excel(file_path)
    df.columns = df.columns.str.strip()
    summary_cols = [
        "Hostname", "IP Type", "VAPT Status", "Critical", "High", "Medium", "Low", "Informational", "Total"
    ]
    # Some Excels may use "Information" instead of "Informational"
    if "Information" in df.columns:
        summary_cols = [c if c != "Informational" else "Information" for c in summary_cols]
    # Keep only the columns we need, coerce NaNs to empty strings for filtering
    summary_df = df[summary_cols].copy()
    summary_df = summary_df.fillna("")
    # Drop rows that are entirely empty (no hostname and all counts empty)
    def row_is_effectively_empty(r):
        if str(r.get(summary_cols[0], '')).strip() != '':
            return False
        # If hostname is empty and all severity totals are empty/zero → drop
        for col in summary_cols[3:]:
            val = str(r.get(col, '')).strip()
            if val not in ('', '0', '0.0'):
                return False
        return True
    mask_keep = ~summary_df.apply(row_is_effectively_empty, axis=1)
    summary_df = summary_df[mask_keep]
    # Build list of dicts with original keys
    summary_list = []
    for _, row in summary_df.iterrows():
        # Calculate total from individual severity counts
        critical = int(float(str(row.get(summary_cols[3], 0)).strip() or 0))
        high = int(float(str(row.get(summary_cols[4], 0)).strip() or 0))
        medium = int(float(str(row.get(summary_cols[5], 0)).strip() or 0))
        low = int(float(str(row.get(summary_cols[6], 0)).strip() or 0))
        informational = int(float(str(row.get(summary_cols[7], row.get('Information', 0))).strip() or 0))
        calculated_total = critical + high + medium + low + informational
        
        summary_list.append({
            "Hostname": row.get(summary_cols[0], ''),
            "IP Type": row.get(summary_cols[1], ''),
            "VAPT Status": row.get(summary_cols[2], ''),
            "Critical": critical,
            "High": high,
            "Medium": medium,
            "Low": low,
            "Informational": informational,
            "Total": calculated_total,
        })
    print(f"All IPs in summary table (filtered): {[r['Hostname'] for r in summary_list]}")
    print(f"Sample IP data: {summary_list[0] if summary_list else 'No data'}")
    return summary_list

def insert_ip_summary_table_colored_at_start(doc, ip_summary_list):
    print("Inserting IP summary table at start. Number of IPs:", len(ip_summary_list))
    headers = [
        ("Sr. No.", "#4A77A8"), ("Hostname", "#4A77A8"), ("IP Type", "#4A77A8"), ("VAPT Status", "#4A77A8"),
        ("Critical", "#990000"), ("High", "#FF0000"), ("Medium", "#FFCC00"),
        ("Low", "#009933"), ("Informational", "#3399CC"), ("Total", "#6A449A")
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, (header, color) in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(11)
    for idx, row in enumerate(ip_summary_list, 1):
        cells = table.add_row().cells
        values = [
            str(idx),
            str(row.get("Hostname", "")),
            str(row.get("IP Type", "")),
            str(row.get("VAPT Status", "")),
            str(row.get("Critical", "")),
            str(row.get("High", "")),
            str(row.get("Medium", "")),
            str(row.get("Low", "")),
            str(row.get("Informational", row.get("Information", ""))),
            str(row.get("Total", ""))
        ]
        for i, value in enumerate(values):
            cells[i].text = value
            if i == 4:
                color = "#990000"
            elif i == 5:
                color = "#FF0000"
            elif i == 6:
                color = "#FFCC00"
            elif i == 7:
                color = "#009933"
            elif i == 8:
                color = "#3399CC"
            elif i == 9:
                color = "#6A449A"
            elif i in [0, 1, 2, 3]:
                color = "#B7D1F2"
            else:
                color = "#FFFFFF"
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
            cells[i]._element.get_or_add_tcPr().append(shading_elm)
            for run in cells[i].paragraphs[0].runs:
                if i in [4, 5, 9]:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif i == 6:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                elif i == 7:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                elif i == 8:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.size = Pt(11)
    # Insert table at the start of the document
    doc._body._element.insert(0, table._element)

def compute_ip_summary_totals(ip_list):
    totals = {"crit": 0, "high": 0, "med": 0, "low": 0, "info": 0, "tot": 0}
    for row in ip_list:
        try:
            # Get individual severity counts
            critical = int(float(str(row.get('Critical', 0)).strip() or 0))
            high = int(float(str(row.get('High', 0)).strip() or 0))
            medium = int(float(str(row.get('Medium', 0)).strip() or 0))
            low = int(float(str(row.get('Low', 0)).strip() or 0))
            informational = int(float(str(row.get('Informational', 0)).strip() or 0))
            
            # Add to totals
            totals["crit"] += critical
            totals["high"] += high
            totals["med"] += medium
            totals["low"] += low
            totals["info"] += informational
            
            # Calculate total for this row
            row_total = critical + high + medium + low + informational
            totals["tot"] += row_total
            
        except Exception as e:
            print(f"Error calculating totals for row {row}: {e}")
            continue
    
    print(f"Calculated totals: {totals}")
    return totals

def insert_ip_summary_table_exact(doc, ip_summary_list):
    print("Inserting exact IP summary table. Number of IPs:", len(ip_summary_list))
    from docx.shared import Inches
    preferred_font = 'Altone Trial'  # Change to your preferred font if needed
    headers = [
        ("Sr. No.", "#B7D1F2", Inches(0.8)),
        ("Hostname", "#B7D1F2", Inches(1.7)),
        ("IP Type", "#B7D1F2", Inches(1.2)),
        ("VAPT Status", "#B7D1F2", Inches(1.3)),
        ("Critical", "#990000", Inches(1.0)),
        ("High", "#FF0000", Inches(1.0)),
        ("Medium", "#FFCC00", Inches(1.0)),
        ("Low", "#009933", Inches(1.0)),
        ("Informational", "#3399CC", Inches(1.2)),
        ("Total", "#6A449A", Inches(1.0)),
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = 'Table Grid'
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, (header, color, width) in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = header
        cell.width = width
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        for run in cell.paragraphs[0].runs:
            if i in [0, 1, 2, 3]:
                run.font.color.rgb = RGBColor(0, 0, 0)
            else:
                run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.name = preferred_font
        if i == 6:  # Medium column (yellow)
            for run in cell.paragraphs[0].runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Prevent text wrapping in header
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        no_wrap = parse_xml(r'<w:noWrap %s/>' % nsdecls('w'))
        tcPr.append(no_wrap)
    def as_int_str(v):
        try:
            if v is None:
                return '0'
            if isinstance(v, str):
                v = v.strip()
                if v == '':
                    return '0'
            return str(int(float(v)))
        except Exception:
            return str(v)
    for idx, row in enumerate(ip_summary_list, 1):
        cells = table.add_row().cells
        values = [
            str(idx),
            str(row.get("host", row.get("Hostname", ""))),
            str(row.get("type", row.get("IP Type", ""))),
            str(row.get("status", row.get("VAPT Status", ""))),
            as_int_str(row.get("crit", row.get("Critical", 0))),
            as_int_str(row.get("high", row.get("High", 0))),
            as_int_str(row.get("med", row.get("Medium", 0))),
            as_int_str(row.get("low", row.get("Low", 0))),
            as_int_str(row.get("info", row.get("Informational", row.get("Information", 0)))),
            as_int_str(row.get("tot", row.get("Total", 0)))
        ]
        for i, value in enumerate(values):
            cells[i].text = value
            # Set cell background and font color
            if i == 4:
                color = "#990000"; font_color = RGBColor(255,255,255)
            elif i == 5:
                color = "#FF0000"; font_color = RGBColor(255,255,255)
            elif i == 6:
                color = "#FFCC00"; font_color = RGBColor(255,255,255)
            elif i == 7:
                color = "#009933"; font_color = RGBColor(255,255,255)
            elif i == 8:
                color = "#3399CC"; font_color = RGBColor(255,255,255)
            elif i == 9:
                color = "#6A449A"; font_color = RGBColor(255,255,255)
            elif i in [0, 1, 2, 3]:
                color = "#B7D1F2"; font_color = RGBColor(0,0,0)
            else:
                color = "#FFFFFF"; font_color = RGBColor(0,0,0)
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
            cells[i]._element.get_or_add_tcPr().append(shading_elm)
            for run in cells[i].paragraphs[0].runs:
                run.font.color.rgb = font_color
                run.font.size = Pt(8)
                run.font.name = preferred_font
                if i in [2, 3]:
                    run.font.bold = False
                else:
                    run.font.bold = True
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[i].width = headers[i][2]
    # Add overall findings row (merged)
    totals = compute_ip_summary_totals(ip_summary_list)
    total_row = table.add_row().cells
    # Merge first four cells
    merged = total_row[0].merge(total_row[1]).merge(total_row[2]).merge(total_row[3])
    merged.text = "Overall Findings"
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#B7D1F2"/>')
    merged._element.get_or_add_tcPr().append(shading_elm)
    for run in merged.paragraphs[0].runs:
        run.font.color.rgb = RGBColor(0,0,0)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.name = preferred_font
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Set totals with correct colors
    color_map = ["#990000", "#FF0000", "#FFCC00", "#009933", "#3399CC", "#6A449A"]
    font_map = [RGBColor(255,255,255), RGBColor(255,255,255), RGBColor(255,255,255), RGBColor(255,255,255), RGBColor(255,255,255), RGBColor(255,255,255)]
    keys = ["crit", "high", "med", "low", "info", "tot"]
    for i, (key, color, font_color) in enumerate(zip(keys, color_map, font_map), start=4):
        total_row[i].text = str(totals[key])
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
        total_row[i]._element.get_or_add_tcPr().append(shading_elm)
        for run in total_row[i].paragraphs[0].runs:
            run.font.color.rgb = font_color
            run.font.bold = True
            run.font.size = Pt(8)
            run.font.name = preferred_font
        total_row[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_row[i].width = headers[i][2]
    # Insert table after heading
    heading_text = "Risk Assessment Analysis of the IPs"
    insert_idx = None
    for i, p in enumerate(doc.paragraphs):
        if heading_text.lower() in p.text.strip().lower():
            insert_idx = i
            break
    if insert_idx is not None:
        # Insert after the heading paragraph
        p = doc.paragraphs[insert_idx]
        parent = p._element.getparent()
        idx = parent.index(p._element)
        parent.insert(idx + 1, table._element)
    else:
        # Fallback: insert at the start
        doc._body._element.insert(0, table._element)

# In generate_word_report, after loading the doc, call this function:
# ip_summary_list = parse_ip_summary_table(vulnerability_file_location)
# insert_ip_summary_table_exact(doc, ip_summary_list)
# Remove any docxtpl table logic for the IP summary table.

# --- GROUPING FUNCTION ---
def group_vulnerabilities(vulnerabilities):
    """
    Groups vulnerabilities by all fields except IP, Sr No, and CVEs (for regular vulnerabilities), and by sr_no for 'No Vulnerability' entries.
    Aggregates all affected IPs and CVEs for each group.
    """
    print(f"Starting grouping with {len(vulnerabilities)} vulnerabilities")
    grouped = {}
    for i, vuln in enumerate(vulnerabilities):
        print(f"  Processing vulnerability {i+1}: SR: {vuln['sr_no']}, Name: {vuln['name']}, is_no_vuln_box: {vuln.get('is_no_vuln_box')}")
        if vuln.get('is_no_vuln_box'):
            key = (vuln['sr_no'],)
            print(f"    No vuln box key: {key}")
        else:
            key = (
                vuln['name'],
                vuln.get('vulnerable_parameter', ''),
                vuln.get('cvss', ''),
                vuln.get('description', ''),
                vuln.get('impact', ''),
                vuln.get('remediation', ''),
            )
            print(f"    Regular vuln key: {key}")
        if key not in grouped:
            grouped[key] = {**vuln, 'ips': [], 'associated_cves': [], 'reference_links': [], 'steps_with_screenshots': [], 'reported_by_list': [], 'project_list': [], 'severity_level_list': []}
            print(f"    Created new group for key")
        else:
            print(f"    Added to existing group for key")
        # Add all unique IPs
        if vuln.get('ip'):
            grouped[key]['ips'].extend([ip.strip() for ip in str(vuln.get('ip')).split(',') if ip.strip()])
        # Only add steps if they don't already exist to prevent duplicates
        existing_steps = grouped[key]['steps_with_screenshots']
        new_steps = vuln.get('steps_with_screenshots', [])
        for step in new_steps:
            if step not in existing_steps:
                existing_steps.append(step)
        grouped[key]['associated_cves'].extend(vuln.get('associated_cves', []))
        # Add reference link if it exists and is not empty
        ref_link = vuln.get('reference_link', '')
        if ref_link and ref_link.strip():
            grouped[key]['reference_links'].append(ref_link)
        # Add new dashboard fields
        if vuln.get('reported_by'):
            grouped[key]['reported_by_list'].append(vuln.get('reported_by'))
        if vuln.get('project'):
            grouped[key]['project_list'].append(vuln.get('project'))
        if vuln.get('severity_level'):
            grouped[key]['severity_level_list'].append(vuln.get('severity_level'))
    for g in grouped.values():
        g['ip'] = ', '.join(sorted(set(g['ips'])))
        g['associated_cves'] = sorted(set(g['associated_cves']))
        # Process reference links - keep unique ones
        g['reference_links'] = list(set(g['reference_links']))
        # Set the main reference_link field to the first one (for backward compatibility)
        if g['reference_links']:
            g['reference_link'] = g['reference_links'][0]
        else:
            g['reference_link'] = ''
        # Process new dashboard fields - keep unique values
        g['reported_by_list'] = list(set(g['reported_by_list']))
        g['project_list'] = list(set(g['project_list']))
        g['severity_level_list'] = list(set(g['severity_level_list']))
        # Set main fields to first values (for backward compatibility)
        g['reported_by'] = g['reported_by_list'][0] if g['reported_by_list'] else ''
        g['project'] = g['project_list'][0] if g['project_list'] else ''
        g['severity_level'] = g['severity_level_list'][0] if g['severity_level_list'] else ''
    print(f"After grouping - {len(grouped)} final vulnerabilities:")
    for g in grouped.values():
        print(f"SR: {g['sr_no']}, Name: {g['name']}, is_no_vuln_box: {g['is_no_vuln_box']}, IPs: {g['ip']}")
    return list(grouped.values())

def create_no_vuln_box(doc, vulnerability, display_sr_no=None, image_map=None):
    # Adjusted colors for a slightly darker header and mid purple IP row
    PURPLE = "#7C4FA3"  # Darker header purple
    MID_PURPLE = "#A98FD3"  # Mid purple for IP row
    table = doc.add_table(rows=5, cols=1)
    table.style = 'Table Grid'
    # Row 1: Title
    cell = table.rows[0].cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{PURPLE}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(f"{display_sr_no}: No Vulnerability Found")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Row 2: IP
    cell = table.rows[1].cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{MID_PURPLE}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("IP: ")
    run.font.bold = True
    run.font.color.rgb = RGBColor(80, 80, 120)
    run.font.name = 'Altone Trial'
    run = paragraph.add_run(vulnerability.get('ip', ''))
    run.font.bold = False
    run.font.color.rgb = RGBColor(80, 80, 120)
    run.font.name = 'Altone Trial'
    # Row 3: Ping
    cell = table.rows[2].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Ping: 0% Packet loss")
    run.font.bold = True
    run.font.color.rgb = RGBColor(124, 79, 163)  # Match header purple
    run.font.name = 'Altone Trial'
    # Row 4: Vulnerability Description
    cell = table.rows[3].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Vulnerability Description:\n")
    run.font.bold = True
    run.font.color.rgb = RGBColor(124, 79, 163)
    run.font.name = 'Altone Trial'
    run = paragraph.add_run(vulnerability.get('description', ''))
    run.font.bold = False
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.name = 'Altone Trial'
    # Row 5: Proof of Concept
    cell = table.rows[4].cells[0]
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Proof of Concept:\n")
    run.font.bold = True
    run.font.color.rgb = RGBColor(124, 79, 163)
    run.font.name = 'Altone Trial'
    steps = vulnerability.get('steps_with_screenshots', [])
    
    # Remove duplicate steps based on text content
    unique_steps = []
    seen_texts = set()
    for step in steps:
        step_text = step.get('text', '').strip()
        if step_text and step_text not in seen_texts:
            unique_steps.append(step)
            seen_texts.add(step_text)
    
    for step_idx, step in enumerate(unique_steps):
        step_para = cell.add_paragraph()
        step_para.paragraph_format.left_indent = Pt(10)
        # Clean the step text to remove any existing "Step X:" prefix
        step_text = step['text'].strip()
        if step_text.lower().startswith('step'):
            # Remove "Step X:" prefix if it exists
            import re
            step_text = re.sub(r'^step\s*\d*\s*:\s*', '', step_text, flags=re.IGNORECASE).strip()
        run = step_para.add_run(f"Step {step_idx+1}: {step_text}")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        # Insert PoC screenshot if available and valid
        screenshot_path = None
        screenshot_name = step.get('screenshot')
        if image_map and isinstance(screenshot_name, str) and screenshot_name and screenshot_name.lower() != 'nan':
            screenshot_path = find_image_for_step(screenshot_name, image_map, vulnerability.get('sr_no', display_sr_no), step_idx+1)
        if screenshot_path and os.path.exists(screenshot_path):
            img_para = cell.add_paragraph()
            img_para.paragraph_format.left_indent = Pt(20)
            img_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_run = img_para.add_run()
            img_run.font.name = 'Altone Trial'
            img_run.add_picture(screenshot_path, width=Inches(5))
    return table

        # Replace [[document_id]] and {{document_id}} in all headers/footers with the actual document ID
    def replace_header_footer_text(doc, placeholder, replacement):
        for section in doc.sections:
            # Header
            for paragraph in section.header.paragraphs:
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)
            # Footer
            for paragraph in section.footer.paragraphs:
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, replacement)

# Place this helper function at the top-level or just inside generate_word_report, but not nested inside another block

def replace_header_footer_text(doc, placeholder, replacement):
    for section in doc.sections:
        # Header
        for paragraph in section.header.paragraphs:
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)
        # Footer
        for paragraph in section.footer.paragraphs:
            if placeholder in paragraph.text:
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, replacement)

# ... inside generate_word_report, after doc = Document(tmp_path) and doc_id is available ...

def calculate_optimal_logo_size(image_path, max_width=2.3, max_height=2.0):
    """
    Calculate the optimal logo size to fit within the template while maintaining aspect ratio.
    Args:
        image_path (str): Path to the logo image
        max_width (float): Maximum width in inches (increase this value if logos appear too small)
        max_height (float): Maximum height in inches (increase this value for taller logos)
    Returns:
        float: Optimal width in inches
    Note:
        - The default parameters (2.3 x 2.0 inches) are suitable for most company logos
        - For wider logos, increase max_width
        - For taller logos, increase max_height
        - The minimum width is set to 1.5 inches to ensure logos are clearly visible
    """
    # Check if PIL is available
    if not PIL_AVAILABLE:
        print("PIL not available, using default logo width")
        return max_width * 0.9  # Use a slightly smaller width as a safe default
    try:
        # Open the image to get its dimensions
        with Image.open(image_path) as img:
            img_width, img_height = img.size
        # Convert pixels to inches (assuming 96 DPI which is standard for many screens)
        width_in_inches = img_width / 96
        height_in_inches = img_height / 96
        print(f"Original logo dimensions: {img_width}x{img_height}px ({width_in_inches:.2f}x{height_in_inches:.2f} inches)")
        # If the image is already within our size limits, use its actual size,
        # but ensure it's not too small (at least 1.5 inches wide for good visibility)
        if width_in_inches <= max_width and height_in_inches <= max_height:
            if width_in_inches < 1.5:
                print(f"Logo is too small, scaling up to minimum width: 1.5 inches")
                return 1.5
            print(f"Logo is already within size limits, using original size: {width_in_inches:.2f} inches wide")
            return width_in_inches
        # Calculate aspect ratio
        aspect_ratio = img_width / img_height
        # Calculate potential width and height based on max dimensions
        width_from_height = max_height * aspect_ratio
        height_from_width = max_width / aspect_ratio
        # Determine which dimension is the limiting factor
        if width_from_height <= max_width:
            # Height is the limiting factor
            optimal_width = width_from_height
        else:
            # Width is the limiting factor
            optimal_width = max_width
        # Ensure minimum size of 1.5 inches width
        if optimal_width < 1.5:
            optimal_width = 1.5
        print(f"Resizing logo to: {optimal_width:.2f} inches wide")
        print(f"Logo display size will be approximately {optimal_width:.2f} x {optimal_width/aspect_ratio:.2f} inches")
        return optimal_width
    except Exception as e:
        print(f"Error calculating optimal logo size: {e}")
        return max_width  # Default to max_width if there's an error
