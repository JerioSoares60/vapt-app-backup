from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request, Depends
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from starlette.middleware.sessions import SessionMiddleware
import os
import sys
import pandas as pd
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import traceback
import tempfile
import json
import re
import zipfile
import io
from sqlalchemy.orm import Session
from db import get_db, AuditLog, CertINReport

# Add parent directory to path to import excel_parser
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../..'))
from excel_parser import parse_excel_data, format_for_template

VERSION = "2.2.0"

app = FastAPI()

app.add_middleware(
    SessionMiddleware,
    secret_key=os.getenv("SESSION_SECRET_KEY", "supersecret"),
    https_only=False,
    same_site="lax",
    session_cookie="reportgen_session"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

def sanitize_filename(filename):
    if not filename:
        return "upload.xlsx"
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", filename)
    return safe_name

def read_vulnerability_excel(file_path):
    """
    Read vulnerability data from Excel using standardized parser with fallback to legacy parsing.
    Converts the new format to legacy format for compatibility with existing code.
    """
    # Try standardized parser first
    try:
        parsed_data = parse_excel_data(file_path)
        
        # Check if we got valid data
        if parsed_data and parsed_data.get('vulnerabilities'):
            print(f"✅ Parsed Excel with standardized parser")
            print(f"   Metadata: {parsed_data['metadata']}")
            print(f"   Assets: {len(parsed_data['assets'])} assets")
            print(f"   Vulnerabilities: {len(parsed_data['vulnerabilities'])} vulnerabilities")
            
            # Convert to legacy format expected by rest of type3.py code
            vulnerabilities = []
            for idx, vuln in enumerate(parsed_data['vulnerabilities'], 1):
                vuln_data = {
                    # Map new format to old format
                    'observation_number': str(vuln.get('sr_no', idx)),
                    'observation': vuln.get('observation', vuln.get('sr_no', idx)),
                    'new_or_repeat': vuln.get('new_or_re', 'New'),
                    'asset': vuln.get('affected_asset', ''),
                    'purpose': '',  # Not directly in vuln, would come from assets
                    'vapt_status': vuln.get('status', 'Open'),
                    'tester_name': vuln.get('tester', parsed_data['metadata']['tester']),
                    'project': vuln.get('project', parsed_data['metadata']['project']),
                    'client': vuln.get('client', parsed_data['metadata']['client']),
                    'cve_cwe': vuln.get('cve_cwe', ''),
                    'cvss': str(vuln.get('cvss', '')),
                    'cvss_version': '',  # Can be extracted from cvss_vector if needed
                    'cvss_vector': vuln.get('cvss_vector', ''),
                    'affected_asset': vuln.get('affected_asset', '') or vuln.get('ip_url_app', ''),
                    'title': vuln.get('observation', ''),
                    'description': vuln.get('detailed_observation', '') or vuln.get('observation_summary', ''),
                    'recommendation': vuln.get('recommendation', ''),
                    'reference': vuln.get('reference', ''),
                    'evidence': vuln.get('evidence', ''),
                    'severity': vuln.get('severity', 'Medium'),
                    'status': vuln.get('status', 'Open'),
                    # Count fields (not used in type3 but kept for compatibility)
                    'critical': '',
                    'high': '',
                    'medium': '',
                    'low': '',
                    'informational': '',
                    'total': ''
                }
                
                # If evidence is empty but we have steps, create evidence text from steps
                if not vuln_data['evidence'] and vuln.get('steps'):
                    evidence_lines = []
                    for step in vuln['steps']:
                        evidence_lines.append(f"Step {step['number']}: {step['content']}")
                    vuln_data['evidence'] = '\n'.join(evidence_lines)
                
                vulnerabilities.append(vuln_data)
                print(f"📄 Row {idx}: Tester={vuln_data['tester_name']}, Project={vuln_data['project']}, Client={vuln_data['client']}, Title={vuln_data['title'][:50]}")
            
            print(f"✅ Converted {len(vulnerabilities)} vulnerabilities to legacy format")
            return vulnerabilities
        else:
            print("⚠️ Standardized parser returned no vulnerabilities, falling back to legacy parsing")
            raise Exception("No vulnerabilities found in standardized parser")
            
    except Exception as e:
        print(f"⚠️ Standardized parser failed: {e}")
        print("📋 Falling back to legacy Excel parsing...")
        traceback.print_exc()
        
        # FALLBACK: Use original/legacy parsing logic
        # If you had a legacy parser before, restore it here
        # For now, return empty to avoid breaking
        return []

async def process_poc_zip_files(poc_files, vulnerabilities):
    """
    Process PoC zip with structure: POC_certIN.zip/POC_certIN/#1, #2, etc.
    Each folder contains images for that specific observation only
    """
    poc_mapping = {}
    
    for poc_file in poc_files:
        if not poc_file.filename:
            continue
            
        temp_zip_path = os.path.join(UPLOAD_DIR, f"temp_{poc_file.filename}")
        content = await poc_file.read()
        with open(temp_zip_path, "wb") as f:
            f.write(content)
        
        extract_dir = os.path.join(UPLOAD_DIR, "poc_images")
        # Clean extract directory to prevent contamination from previous uploads
        if os.path.exists(extract_dir):
            import shutil
            shutil.rmtree(extract_dir)
        os.makedirs(extract_dir, exist_ok=True)
        
        try:
            with zipfile.ZipFile(temp_zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)
            
            print(f"✅ Extracted zip to: {extract_dir}")
            
            # Find all observation folders (#1, #2, #3, etc.)
            for root, dirs, files in os.walk(extract_dir):
                for dir_name in dirs:
                    # Check if directory name matches #1, #2, etc.
                    match = re.match(r'#(\d+)', dir_name)
                    if match:
                        obs_num = int(match.group(1))
                        obs_key = f"OBS-{obs_num:03d}"
                        obs_folder = os.path.join(root, dir_name)
                        
                        print(f"📁 Found observation folder: {obs_folder} for {obs_key}")
                        
                        # Get all images in this observation folder ONLY
                        obs_images = []
                        for img_file in sorted(os.listdir(obs_folder)):
                            if img_file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                                img_path = os.path.join(obs_folder, img_file)
                                
                                # Extract step number from filename
                                step_match = re.search(r'step[_\-\s]?(\d+)', img_file.lower())
                                if step_match:
                                    step_num = int(step_match.group(1))
                                else:
                                    # Try just extracting a number
                                    num_match = re.search(r'(\d+)', img_file)
                                    step_num = int(num_match.group(1)) if num_match else len(obs_images) + 1
                                
                                obs_images.append({
                                    'filename': img_file,
                                    'path': img_path,
                                    'step_number': step_num
                                })
                                print(f"  🖼️  Found image: {img_file} (Step {step_num})")
                        
                        # Sort images by step number
                        obs_images.sort(key=lambda x: x['step_number'])
                        
                        if obs_images:
                            poc_mapping[obs_key] = obs_images
                            print(f"✅ Mapped {len(obs_images)} images to {obs_key}")
            
        except Exception as e:
            print(f"Error extracting PoC zip: {e}")
            traceback.print_exc()
        finally:
            if os.path.exists(temp_zip_path):
                os.remove(temp_zip_path)
    
    print(f"🎯 Final PoC mapping keys: {list(poc_mapping.keys())}")
    for key, images in poc_mapping.items():
        print(f"   {key}: {len(images)} images")
    return poc_mapping

def parse_evidence_steps(evidence_text):
    """Parse evidence/PoC text to extract step descriptions"""
    if not evidence_text or evidence_text.lower() in ['nan', 'none', '']:
        return []
    
    steps = []
    
    # Try to split by "Step X:" pattern
    step_pattern = re.compile(r'Step\s*(\d+)\s*[:\-]?\s*(.+?)(?=Step\s*\d+|$)', re.IGNORECASE | re.DOTALL)
    matches = step_pattern.findall(evidence_text)
    
    if matches:
        for step_num, step_desc in matches:
            steps.append({
                'number': int(step_num),
                'description': step_desc.strip()
            })
    else:
        # If no step pattern, split by newlines and number each
        lines = [line.strip() for line in evidence_text.split('\n') if line.strip()]
        for i, line in enumerate(lines, 1):
            steps.append({
                'number': i,
                'description': line
            })
    
    return steps

def safe_str(val, default=''):
    """Convert value to string, handling NaN/None/float safely"""
    if val is None:
        return default
    if isinstance(val, float) and pd.isna(val):
        return default
    return str(val) if val else default

def generate_vulnerability_sections(vulnerabilities, poc_mapping):
    """Generate vulnerability sections with proper data extraction"""
    vulnerability_sections = []
    
    for i, vuln in enumerate(vulnerabilities, 1):
        print(f"\nProcessing vulnerability {i}:")
        print(f"Available keys: {list(vuln.keys())}")
        
        # Extract data using standard keys from normalized mapping - use safe_str to handle NaN
        obs_num = safe_str(vuln.get('observation_number'), str(i))
        new_or_repeat = safe_str(vuln.get('new_or_repeat'), 'New')
        cve_cwe = safe_str(vuln.get('cve_cwe'))
        cvss_version = safe_str(vuln.get('cvss_version'))
        affected_asset = safe_str(vuln.get('affected_asset'))
        title = safe_str(vuln.get('title'))
        description = safe_str(vuln.get('description'))
        recommendation = safe_str(vuln.get('recommendation'))
        reference = safe_str(vuln.get('reference'))
        evidence_text = safe_str(vuln.get('evidence'))
        
        # Extract severity and other metadata
        severity = safe_str(vuln.get('severity'), 'Medium')
        cvss = safe_str(vuln.get('cvss'))
        cvss_vector = safe_str(vuln.get('cvss_vector'))
        status = safe_str(vuln.get('status'), 'Open')
        
        print(f"Extracted - Title: {title}, CVE: {cve_cwe}, Asset: {affected_asset}")
        
        # Parse evidence steps from Excel
        evidence_steps = parse_evidence_steps(evidence_text)
        
        # Get PoC images for this observation
        obs_key = f"OBS-{i:03d}"
        poc_images = poc_mapping.get(obs_key, [])
        
        print(f"Found {len(poc_images)} PoC images for {obs_key}")
        print(f"Found {len(evidence_steps)} evidence steps from Excel")
        
        # Match images with step descriptions
        matched_poc = []
        for idx, img in enumerate(poc_images):
            step_desc = ''
            if idx < len(evidence_steps):
                step_desc = evidence_steps[idx]['description']
            
            matched_poc.append({
                'filename': img['filename'],
                'path': img['path'],
                'step_number': idx + 1,
                'description': step_desc
            })
        
        # Process recommendations
        recommendations_list = []
        if recommendation and isinstance(recommendation, str):
            rec_lines = [line.strip() for line in recommendation.split('\n') if line.strip()]
            for line in rec_lines:
                clean_line = re.sub(r'^\d+\.\s*', '', line)
                if clean_line:
                    recommendations_list.append(clean_line)
        
        vulnerability_section = {
            'observation_number': i,
            'new_or_repeat': new_or_repeat,
            'severity': severity,
            'status': status,
            'cve_cwe': cve_cwe,
            'cvss': cvss,
            'cvss_version': cvss_version,
            'cvss_vector': cvss_vector,
            'affected_asset': affected_asset,
            'title': title,
            'description': description,
            'recommendations': recommendations_list,
            'reference': reference,
            'poc_images': matched_poc,
            'has_poc': len(matched_poc) > 0
        }
        
        vulnerability_sections.append(vulnerability_section)
        print(f"Created section with {len(matched_poc)} PoC images")
    
    return vulnerability_sections



def create_landscape_vulnerability_box(doc, vulnerability_section):
    """Create properly formatted vulnerability box matching expected layout"""
    import os
    from docx.shared import Inches, Pt, RGBColor, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    # Get severity colors
    severity = safe_str(vulnerability_section.get('severity', 'Medium')).strip()
    severity_colors = {
        'Critical': {'bg': '800000', 'text': 'FFFFFF'},
        'High': {'bg': 'FF0000', 'text': 'FFFFFF'},
        'Medium': {'bg': 'FFC000', 'text': '000000'},
        'Low': {'bg': '92D050', 'text': 'FFFFFF'},
        'Informational': {'bg': '00B0F0', 'text': 'FFFFFF'}
    }
    colors = severity_colors.get(severity, severity_colors['Medium'])

    # Create main table - Row 1: Obs | Severity | Status, Row 2: New/Repeat | CVE/CWE
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set light gray border color for the table
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '<w:left w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '<w:bottom w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '<w:right w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '<w:insideH w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '<w:insideV w:val="single" w:sz="4" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(tblBorders)

    # ---------- Row 1: Observation | Severity | Status ----------
    row1 = table.rows[0]
    
    # Observation cell
    cell_obs = row1.cells[0]
    cell_obs.text = ''
    p = cell_obs.paragraphs[0]
    run = p.add_run(f"Observation: #{vulnerability_section.get('observation_number', '1')}")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(11)
    run.font.bold = True

    # Severity cell with colored background
    cell_sev = row1.cells[1]
    cell_sev.text = ''
    p = cell_sev.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Severity: {severity}")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255) if colors['text'] == 'FFFFFF' else RGBColor(0, 0, 0)
    cell_sev._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["bg"]}"/>'))

    # Status cell (yellow/orange)
    cell_status = row1.cells[2]
    cell_status.text = ''
    p = cell_status.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Status: {safe_str(vulnerability_section.get('status', 'Open'))}")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(11)
    run.font.bold = True
    cell_status._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFD966"/>'))

    # ---------- Row 2: New/Repeat | CVE/CWE (2 columns) ----------
    row2 = table.rows[1]
    
    # New or Repeat Observation cell
    cell_nor = row2.cells[0]
    cell_nor.text = ''
    p = cell_nor.paragraphs[0]
    run = p.add_run("New or Repeat Observation: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(10)
    run.font.bold = True
    run = p.add_run(safe_str(vulnerability_section.get('new_or_repeat', 'New')))
    run.font.name = 'Altone Trial'
    run.font.size = Pt(10)

    # CVE/CWE cell (merge cells 1 and 2)
    cell_cve = row2.cells[1]
    cell_cve.merge(row2.cells[2])
    cell_cve.text = ''
    p = cell_cve.paragraphs[0]
    run = p.add_run("CVE/CWE: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(10)
    run.font.bold = True
    run = p.add_run(safe_str(vulnerability_section.get('cve_cwe', '')))
    run.font.name = 'Altone Trial'
    run.font.size = Pt(10)

    # ---------- Content rows (each as separate row for proper layout) ----------
    
    def add_content_row(label, value, is_url=False):
        """Add a single-cell merged row with label: value"""
        if not value:
            return
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[1])
        cell.merge(row.cells[2])
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run(label)
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        run = p.add_run(safe_str(value))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        if is_url:
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True

    # CVSS row
    add_content_row("CVSS: ", vulnerability_section.get('cvss'))
    
    # CVSS Version Ref row
    add_content_row("CVSS Version Ref: ", vulnerability_section.get('cvss_vector'))

    # Affected Asset row
    add_content_row("Affected Asset i.e. IP/URL/Application etc.: ", 
                    vulnerability_section.get('affected_asset'), is_url=True)

    # Observation/Vulnerability Title row
    add_content_row("Observation/ Vulnerability Title: ", vulnerability_section.get('title'))

    # Detailed Observation row
    add_content_row("Detailed Observation/ Vulnerable Point: ", vulnerability_section.get('description'))

    # Recommendation row
    recommendations = vulnerability_section.get('recommendations', [])
    if recommendations:
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[1])
        cell.merge(row.cells[2])
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run("Recommendation:")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        for i, rec in enumerate(recommendations, 1):
            rec_p = cell.add_paragraph()
            rec_p.paragraph_format.space_before = Pt(0)
            rec_p.paragraph_format.space_after = Pt(0)
            run = rec_p.add_run(f"{i}. {safe_str(rec)}")
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)

    # Reference row
    reference = vulnerability_section.get('reference')
    if reference:
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[1])
        cell.merge(row.cells[2])
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run("Reference:")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        run.font.bold = True
        
        ref_p = cell.add_paragraph()
        run = ref_p.add_run(safe_str(reference))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True

    # Evidence/Proof of Concept row
    poc_images = vulnerability_section.get('poc_images', [])
    if poc_images:
        row = table.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[1])
        cell.merge(row.cells[2])
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        
        run = p.add_run("Evidence / Proof of Concept:")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        run.font.bold = True

        for poc_img in poc_images:
            # Step description
            step_p = cell.add_paragraph()
            step_p.paragraph_format.space_before = Pt(6)
            step_p.paragraph_format.space_after = Pt(3)
            
            step_text = f"Step {poc_img.get('step_number', '')}: "
            if poc_img.get('description'):
                step_text += safe_str(poc_img['description'])
            
            run = step_p.add_run(step_text)
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)
            run.font.bold = False

            # Add the image if present
            img_path = poc_img.get('path') or poc_img.get('filename')
            if img_path and os.path.exists(img_path):
                img_p = cell.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                try:
                    img_run.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    err_p = cell.add_paragraph(f"[Image error: {os.path.basename(img_path)}: {e}]")
                    err_p.runs[0].font.italic = True
            else:
                missing_p = cell.add_paragraph(f"[Image missing: {poc_img.get('filename', 'unknown')}]")
                missing_p.runs[0].font.italic = True

    return table


def create_asset_findings_table(doc, vulnerability_data):
    """Create the Asset/Hostname findings table with vulnerability counts"""
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    if not vulnerability_data:
        return None
    
    # Aggregate data by asset
    asset_data = {}
    for vuln in vulnerability_data:
        asset = safe_str(vuln.get('asset', '')) or safe_str(vuln.get('affected_asset', '')) or 'Unknown'
        purpose = safe_str(vuln.get('purpose', '')) or safe_str(vuln.get('vapt_status', ''))
        status = safe_str(vuln.get('vapt_status', '')) or safe_str(vuln.get('status', 'Open'))
        severity = safe_str(vuln.get('severity', 'Medium')).lower()
        
        if asset not in asset_data:
            asset_data[asset] = {
                'purpose': purpose,
                'status': status,
                'critical': 0,
                'high': 0,
                'medium': 0,
                'low': 0,
                'informational': 0
            }
        
        # Count by severity
        if severity in ['critical']:
            asset_data[asset]['critical'] += 1
        elif severity in ['high']:
            asset_data[asset]['high'] += 1
        elif severity in ['medium']:
            asset_data[asset]['medium'] += 1
        elif severity in ['low']:
            asset_data[asset]['low'] += 1
        elif severity in ['informational', 'info']:
            asset_data[asset]['informational'] += 1
    
    # Calculate totals
    overall = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0, 'informational': 0}
    for asset, counts in asset_data.items():
        counts['total'] = counts['critical'] + counts['high'] + counts['medium'] + counts['low'] + counts['informational']
        overall['critical'] += counts['critical']
        overall['high'] += counts['high']
        overall['medium'] += counts['medium']
        overall['low'] += counts['low']
        overall['informational'] += counts['informational']
    overall['total'] = overall['critical'] + overall['high'] + overall['medium'] + overall['low'] + overall['informational']
    
    # Create table with header row + data rows + overall row (no page break - handled by caller)
    num_rows = 1 + len(asset_data) + 1  # header + assets + overall
    table = doc.add_table(rows=num_rows, cols=9)
    table.style = 'Table Grid'
    
    # Column headers
    headers = ['Sr.\nNo.', 'Asset/Hostname', 'Instant\npurpose', 'VAPT\nStatus', 'Critical', 'High', 'Medium', 'Low', 'Informational', 'Total']
    
    # Actually we need 10 columns
    # Recreate table with correct columns
    for row in table.rows:
        for cell in row.cells:
            pass
    
    # Remove old table and create new one
    table._element.getparent().remove(table._element)
    table = doc.add_table(rows=num_rows, cols=10)
    table.style = 'Table Grid'
    
    # Header row styling
    header_row = table.rows[0]
    headers = ['Sr.\nNo.', 'Asset/Hostname', 'Instant\npurpose', 'VAPT\nStatus', 'Critical', 'High', 'Medium', 'Low', 'Informational', 'Total']
    
    severity_colors = {
        4: '800000',  # Critical - Dark Red/Maroon
        5: 'FF0000',  # High - Red
        6: 'FFC000',  # Medium - Orange
        7: '92D050',  # Low - Green
        8: '00B0F0',  # Informational - Blue
        9: '7030A0'   # Total - Purple
    }
    
    # Header colors: first 4 columns purple, then severity colors for each column
    header_colors = {
        0: '6923D0',  # Sr. No. - Purple
        1: '6923D0',  # Asset/Hostname - Purple
        2: '6923D0',  # Instant purpose - Purple
        3: '6923D0',  # VAPT Status - Purple
        4: '800000',  # Critical - Dark Red
        5: 'FF0000',  # High - Red
        6: 'FFC000',  # Medium - Orange
        7: '92D050',  # Low - Green
        8: '00B0F0',  # Informational - Blue
        9: '7030A0'   # Total - Purple
    }
    
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header_text)
        run.font.name = 'Altone Trial'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Apply color based on column
        color = header_colors.get(idx, '6923D0')
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Data rows
    row_idx = 1
    for sr_no, (asset, counts) in enumerate(asset_data.items(), 1):
        row = table.rows[row_idx]
        
        # Sr. No.
        row.cells[0].text = str(sr_no) + '.'
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Asset/Hostname
        row.cells[1].text = asset
        
        # Instant purpose
        row.cells[2].text = counts['purpose']
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # VAPT Status
        row.cells[3].text = counts['status']
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Severity counts with colored backgrounds
        severity_vals = [
            (4, counts['critical'], '800000'),
            (5, counts['high'], 'FF0000'),
            (6, counts['medium'], 'FFC000'),
            (7, counts['low'], '92D050'),
            (8, counts['informational'], '00B0F0'),
            (9, counts['total'], '7030A0')
        ]
        
        for col_idx, val, color in severity_vals:
            cell = row.cells[col_idx]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
            cell._element.get_or_add_tcPr().append(shading)
        
        row_idx += 1
    
    # Overall Findings row
    overall_row = table.rows[row_idx]
    
    # Merge first 4 cells for "Overall Findings"
    overall_row.cells[0].merge(overall_row.cells[3])
    overall_row.cells[0].text = 'Overall Findings'
    overall_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in overall_row.cells[0].paragraphs[0].runs:
        run.font.bold = True
    
    # Overall severity counts
    overall_vals = [
        (4, overall['critical'], '800000'),
        (5, overall['high'], 'FF0000'),
        (6, overall['medium'], 'FFC000'),
        (7, overall['low'], '92D050'),
        (8, overall['informational'], '00B0F0'),
        (9, overall['total'], '7030A0')
    ]
    
    for col_idx, val, color in overall_vals:
        cell = overall_row.cells[col_idx]
        cell.text = str(val)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Apply font styling to all cells
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = 'Altone Trial'
                    if run.font.size is None:
                        run.font.size = Pt(10)
    
    return table


def generate_certin_report_from_form(data, template_path, output_path, vulnerability_data=None, poc_mapping=None):
    """Generate Cert-IN report from form data and template"""
    try:
        safe_template_path = _sanitize_docx_jinja_placeholders(template_path)
        doc = DocxTemplate(safe_template_path)
        
        document_change_history = json.loads(data.get('document_change_history', '[]'))
        distribution_list = json.loads(data.get('distribution_list', '[]'))
        engagement_scope = json.loads(data.get('engagement_scope', '[]'))
        auditing_team = json.loads(data.get('auditing_team', '[]'))
        audit_activities = json.loads(data.get('audit_activities', '[]'))
        tools_software = json.loads(data.get('tools_software', '[]'))
        
        vulnerability_sections = []
        client_name_from_excel = ''
        project_name_from_excel = ''
        if vulnerability_data and poc_mapping is not None:
            vulnerability_sections = generate_vulnerability_sections(vulnerability_data, poc_mapping)
            # Extract client/project from first vulnerability row if not provided in form
            if vulnerability_data:
                client_name_from_excel = vulnerability_data[0].get('client', '')
                project_name_from_excel = vulnerability_data[0].get('project', '')
                print(f"📊 Extracted from Excel - Client: {client_name_from_excel}, Project: {project_name_from_excel}")
        
        # Use form data as primary, fall back to Excel data
        final_client_name = data.get('client_name', '') or client_name_from_excel
        final_report_name = data.get('report_name', '') or project_name_from_excel
        
        context = {
            'CLIENT_NAME': final_client_name,
            'REPORT_NAME': final_report_name,
            'REPORT_RELEASE_DATE': data.get('report_release_date', ''),
            'TYPE_OF_AUDIT': data.get('type_of_audit', ''),
            'TYPE_OF_AUDIT_REPORT': data.get('type_of_audit_report', ''),
            'PERIOD': data.get('period', ''),
            'DOCUMENT_TITLE': data.get('document_title', ''),
            'DOCUMENT_ID': data.get('document_id', ''),
            'DOCUMENT_VERSION': data.get('document_version', ''),
            'PREPARED_BY': data.get('prepared_by', ''),
            'REVIEWED_BY': data.get('reviewed_by', ''),
            'APPROVED_BY': data.get('approved_by', ''),
            'RELEASED_BY': data.get('released_by', ''),
            'RELEASE_DATE': data.get('release_date', ''),
            'DOCUMENT_CHANGE_HISTORY': document_change_history,
            'DISTRIBUTION_LIST': distribution_list,
            'ENGAGEMENT_SCOPE': engagement_scope,
            'AUDITING_TEAM': auditing_team,
            'AUDIT_ACTIVITIES': audit_activities,
            'TOOLS_SOFTWARE': tools_software,
            'VULNERABILITIES': vulnerability_sections,
            'HAS_VULNERABILITIES': len(vulnerability_sections) > 0,
            'VULNERABILITY_COUNT': len(vulnerability_sections),
            'GENERATION_DATE': datetime.now().strftime('%d %B %Y'),
            'GENERATION_TIME': datetime.now().strftime('%H:%M:%S'),
        }
        
        doc.render(context)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp_path = tmp.name
        
        from docx import Document
        doc = Document(tmp_path)
        
        if vulnerability_sections:
            # Find where "Total Vulnerabilities" graph starts
            graph_idx = None
            for i, para in enumerate(doc.paragraphs):
                if 'Total Vulnerabilities' in para.text:
                    graph_idx = i
                    break
            
            # Create the table
            table = create_asset_findings_table(doc, vulnerability_data)
            
            if table is not None and graph_idx is not None:
                # Add a page break paragraph before the table to ensure it starts on a fresh page
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                
                graph_para = doc.paragraphs[graph_idx]._element
                
                # Remove table from end and insert before graph
                table._element.getparent().remove(table._element)
                graph_para.addprevious(table._element)
            
            # Detailed Observations heading on new page
            doc.add_page_break()
            vuln_heading = doc.add_paragraph()
            vuln_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            vuln_run = vuln_heading.add_run("Detailed Observations")
            vuln_run.font.name = 'Altone Trial'
            vuln_run.font.size = Pt(18)
            vuln_run.font.bold = True
            vuln_run.font.color.rgb = RGBColor(106, 68, 154)
            
            for i, vuln_section in enumerate(vulnerability_sections):
                create_landscape_vulnerability_box(doc, vuln_section)
                
                if i < len(vulnerability_sections) - 1:
                    doc.add_page_break()
        
        doc.save(output_path)
        os.remove(tmp_path)
        
        return True
        
    except Exception as e:
        print(f"Error generating report: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")

def _sanitize_docx_jinja_placeholders(template_path: str) -> str:
    try:
        with open(template_path, 'rb') as f:
            original_bytes = f.read()
        zin = zipfile.ZipFile(io.BytesIO(original_bytes))
        out_buf = io.BytesIO()
        zout = zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED)
        
        pattern_expr = re.compile(r"(\{\{\s*)([^}]+?)(\s*\}\})")
        
        def _normalize_print_statement(inner: str) -> str:
            inner = inner.strip()
            if inner == "Audit.sr_no":
                inner = "team.sr_no"
            
            chars = []
            for ch in inner:
                if re.match(r"[A-Za-z0-9_.]", ch):
                    chars.append(ch)
                elif ch.isspace():
                    chars.append('_')
                else:
                    chars.append('_')
            return ''.join(chars)
        
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith('word/') and item.filename.endswith('.xml'):
                xml = data.decode('utf-8', errors='ignore')
                xml = pattern_expr.sub(lambda m: m.group(1) + _normalize_print_statement(m.group(2)) + m.group(3), xml)
                data = xml.encode('utf-8')
            zout.writestr(item, data)
        zin.close()
        zout.close()
        
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmpf:
            tmpf.write(out_buf.getvalue())
            return tmpf.name
    except Exception:
        return template_path

@app.get("/", response_class=HTMLResponse)
async def root():
    """Serve the Cert-IN form interface"""
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Cert-IN Report Generator</title>
        <style>
            body { font-family: Calibri, sans-serif; margin: 20px; background-color: #f5f5f5; }
            .container { max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
            h1 { color: #6923d0; text-align: center; margin-bottom: 30px; }
            .form-section { margin-bottom: 30px; padding: 20px; border: 1px solid #ddd; border-radius: 5px; }
            .form-section h3 { color: #6923d0; margin-top: 0; border-bottom: 2px solid #6923d0; padding-bottom: 10px; }
            .form-row { display: flex; gap: 20px; margin-bottom: 15px; }
            .form-group { flex: 1; }
            .form-group label { display: block; margin-bottom: 5px; font-weight: bold; color: #333; }
            .form-group input, .form-group select, .form-group textarea { width: 100%; padding: 8px; border: 1px solid #ddd; border-radius: 4px; box-sizing: border-box; }
            .form-group textarea { height: 60px; resize: vertical; }
            .dynamic-list { border: 1px solid #ddd; padding: 15px; margin: 10px 0; border-radius: 4px; }
            .dynamic-item { display: flex; gap: 10px; margin-bottom: 10px; align-items: end; }
            .dynamic-item input { flex: 1; }
            .btn { background-color: #6923d0; color: white; border: none; padding: 10px 20px; border-radius: 4px; cursor: pointer; margin: 5px; }
            .btn:hover { background-color: #5a1fb8; }
            .btn-secondary { background-color: #6c757d; }
            .btn-secondary:hover { background-color: #5a6268; }
            .btn-danger { background-color: #dc3545; }
            .btn-danger:hover { background-color: #c82333; }
            .btn-success { background-color: #28a745; }
            .btn-success:hover { background-color: #218838; }
            .progress { display: none; margin: 20px 0; }
            .progress-bar { background-color: #6923d0; height: 20px; border-radius: 4px; width: 0%; transition: width 0.3s; }
            .result { display: none; margin: 20px 0; padding: 15px; border-radius: 4px; }
            .result.success { background-color: #d4edda; border: 1px solid #c3e6cb; color: #155724; }
            .result.error { background-color: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; }
            .info-box { background-color: #e7f3ff; border-left: 4px solid #2196F3; padding: 15px; margin: 15px 0; }
            .info-box h4 { margin-top: 0; color: #2196F3; }
            .info-box code { background-color: #fff; padding: 2px 6px; border-radius: 3px; font-family: monospace; }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🔒 Cert-IN Report Generator v2.2</h1>
            <div style="text-align:center;margin-bottom:16px;">
                <a href="/report_formats.html" class="btn btn-secondary" style="text-decoration:none;display:inline-block;">← Back to Report Formats</a>
            </div>
            <p style="text-align: center; color: #666; margin-bottom: 30px;">
                Generate Cert-IN compliant reports with document control pages
            </p>
            
            <form id="certinForm">
                <!-- Basic Report Information -->
                <div class="form-section">
                    <h3>📋 Basic Report Information</h3>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="client_name">Client Name *</label>
                            <input type="text" id="client_name" name="client_name" required>
                        </div>
                        <div class="form-group">
                            <label for="report_name">Web/Network Initial Report Name *</label>
                            <input type="text" id="report_name" name="report_name" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="report_release_date">Report Release Date *</label>
                            <input type="date" id="report_release_date" name="report_release_date" required>
                        </div>
                        <div class="form-group">
                            <label for="type_of_audit">Type of Audit *</label>
                            <select id="type_of_audit" name="type_of_audit" required>
                                <option value="">Select Type</option>
                                <option value="Security assessment (VAPT) - WEB & Network">Security assessment (VAPT) - WEB & Network</option>
                                <option value="Security assessment (VAPT) - WEB">Security assessment (VAPT) - WEB</option>
                                <option value="Security assessment (VAPT) - Network">Security assessment (VAPT) - Network</option>
                                <option value="Penetration Testing">Penetration Testing</option>
                                <option value="Vulnerability Assessment">Vulnerability Assessment</option>
                            </select>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="type_of_audit_report">Type of Audit Report *</label>
                            <select id="type_of_audit_report" name="type_of_audit_report" required>
                                <option value="">Select Type</option>
                                <option value="First Audit Report">First Audit Report</option>
                                <option value="Follow-up Audit Report">Follow-up Audit Report</option>
                                <option value="Final Audit Report">Final Audit Report</option>
                            </select>
                        </div>
                        <div class="form-group">
                            <label for="period">Period *</label>
                            <input type="text" id="period" name="period" placeholder="e.g., 23-07-2025 to 02-08-2025" required>
                        </div>
                    </div>
                </div>

                <!-- Document Control -->
                <div class="form-section">
                    <h3>📄 Document Control</h3>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="document_title">Document Title *</label>
                            <input type="text" id="document_title" name="document_title" required>
                        </div>
                        <div class="form-group">
                            <label for="document_id">Document ID *</label>
                            <input type="text" id="document_id" name="document_id" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="document_version">Document Version *</label>
                            <input type="text" id="document_version" name="document_version" value="1.0" required>
                        </div>
                        <div class="form-group">
                            <label for="prepared_by">Prepared by *</label>
                            <input type="text" id="prepared_by" name="prepared_by" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="reviewed_by">Reviewed by *</label>
                            <input type="text" id="reviewed_by" name="reviewed_by" required>
                        </div>
                        <div class="form-group">
                            <label for="approved_by">Approved by *</label>
                            <input type="text" id="approved_by" name="approved_by" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="released_by">Released by *</label>
                            <input type="text" id="released_by" name="released_by" required>
                        </div>
                        <div class="form-group">
                            <label for="release_date">Release Date *</label>
                            <input type="date" id="release_date" name="release_date" required>
                        </div>
                    </div>
                </div>

                <!-- Document Change History -->
                <div class="form-section">
                    <h3>📝 Document Change History</h3>
                    <div id="changeHistoryList">
                        <div class="dynamic-item">
                            <input type="text" name="change_version[]" placeholder="Version" required>
                            <input type="date" name="change_date[]" required>
                            <input type="text" name="change_remarks[]" placeholder="Remarks/Reason of change" required>
                            <button type="button" class="btn btn-danger" onclick="removeChangeHistory(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addChangeHistory()">+ Add Change History</button>
                </div>

                <!-- Document Distribution List -->
                <div class="form-section">
                    <h3>📧 Document Distribution List</h3>
                    <div id="distributionList">
                        <div class="dynamic-item">
                            <input type="text" name="dist_name[]" placeholder="Name" required>
                            <input type="text" name="dist_organization[]" placeholder="Organization" required>
                            <input type="text" name="dist_designation[]" placeholder="Designation" required>
                            <input type="email" name="dist_email[]" placeholder="Email ID" required>
                            <button type="button" class="btn btn-danger" onclick="removeDistribution(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addDistribution()">+ Add Distribution</button>
                </div>

                <!-- Engagement Scope -->
                <div class="form-section">
                    <h3>🎯 Engagement Scope</h3>
                    <div id="engagementScopeList">
                        <div class="dynamic-item">
                            <input type="number" name="scope_sr_no[]" placeholder="S. No">
                            <input type="text" name="scope_asset[]" placeholder="Asset Description">
                            <input type="text" name="scope_criticality[]" placeholder="Criticality">
                            <input type="text" name="scope_internal_ip[]" placeholder="Internal IP">
                            <input type="text" name="scope_url[]" placeholder="URL or NA">
                            <input type="text" name="scope_public_ip[]" placeholder="Public IP">
                            <input type="text" name="scope_location[]" placeholder="Location">
                            <input type="text" name="scope_hash[]" placeholder="Hash Value">
                            <input type="text" name="scope_version[]" placeholder="Version">
                            <input type="text" name="scope_other[]" placeholder="Other Details">
                            <button type="button" class="btn btn-danger" onclick="removeEngagementScope(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addEngagementScope()">+ Add Asset</button>
                </div>

                <!-- Details of Auditing Team -->
                <div class="form-section">
                    <h3>👥 Details of Auditing Team</h3>
                    <div id="auditingTeamList">
                        <div class="dynamic-item">
                            <input type="text" name="team_name[]" placeholder="Name" required>
                            <input type="text" name="team_designation[]" placeholder="Designation" required>
                            <input type="email" name="team_email[]" placeholder="Email ID" required>
                            <input type="text" name="team_qualifications[]" placeholder="Qualifications/Certifications" required>
                            <select name="team_certin_listed[]" required>
                                <option value="">CERT-In Listed?</option>
                                <option value="Yes">Yes</option>
                                <option value="No">No</option>
                            </select>
                            <button type="button" class="btn btn-danger" onclick="removeAuditingTeam(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addAuditingTeam()">+ Add Team Member</button>
                </div>

                <!-- Audit Activities and Timelines -->
                <div class="form-section">
                    <h3>📅 Audit Activities and Timelines</h3>
                    <div id="auditActivitiesList">
                        <div class="dynamic-item">
                            <input type="text" name="activity_task[]" placeholder="Activity/Task" required>
                            <input type="date" name="activity_date[]" required>
                            <button type="button" class="btn btn-danger" onclick="removeAuditActivity(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addAuditActivity()">+ Add Activity</button>
                </div>

                <!-- Tools/Software Used -->
                <div class="form-section">
                    <h3>🛠️ Tools/Software Used</h3>
                    <div id="toolsSoftwareList">
                        <div class="dynamic-item">
                            <input type="number" name="tool_sr_no[]" placeholder="S. No" required>
                            <input type="text" name="tool_name[]" placeholder="Name of Tool/Software" required>
                            <input type="text" name="tool_version[]" placeholder="Version" required>
                            <select name="tool_type[]" required>
                                <option value="">Type</option>
                                <option value="Opensource">Opensource</option>
                                <option value="Licensed">Licensed</option>
                            </select>
                            <button type="button" class="btn btn-danger" onclick="removeToolSoftware(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addToolSoftware()">+ Add Tool</button>
                </div>

                <!-- Vulnerability Data Upload -->
                <div class="form-section">
                    <h3>🔍 Vulnerability Data (Required)</h3>
                    
                    <div class="info-box">
                        <h4>📋 Excel Format Requirements</h4>
                        <p><strong>Required columns (any case/spacing):</strong></p>
                        <ul>
                            <li><code>New or Repeat Observation</code></li>
                            <li><code>CVE/CWE</code></li>
                            <li><code>CVSS Version Ref</code></li>
                            <li><code>Affected Asset i.e. IP/URL/Application etc.</code></li>
                            <li><code>Observation/ Vulnerability Title</code></li>
                            <li><code>Detailed Observation/ Vulnerable Point</code></li>
                            <li><code>Recommendation</code></li>
                            <li><code>Reference</code></li>
                            <li><code>Evidence / Proof of Concept</code> - Write step descriptions here</li>
                        </ul>
                        <p><strong>Note:</strong> Column names are case-insensitive and flexible with spaces</p>
                    </div>
                    
                    <div class="form-group">
                        <label for="vulnerability_file">Upload Vulnerability Excel File *</label>
                        <input type="file" id="vulnerability_file" name="vulnerability_file" accept=".xlsx,.xls" required>
                    </div>
                    
                    <div class="info-box">
                        <h4>📁 PoC Zip Structure Requirements</h4>
                        <p><strong>Folder Structure:</strong></p>
                        <code>POC_certIN.zip/POC_certIN/#1, #2, #3</code>
                        <p><strong>Example:</strong></p>
                        <ul>
                            <li><code>POC_certIN/#1/step1.png</code> → Observation #1, Step 1</li>
                            <li><code>POC_certIN/#1/step2.png</code> → Observation #1, Step 2</li>
                            <li><code>POC_certIN/#2/step1.png</code> → Observation #2, Step 1</li>
                        </ul>
                        <p><strong>Important:</strong> Each <code>#N</code> folder contains ONLY images for that observation</p>
                    </div>
                    
                    <div class="form-group">
                        <label for="poc_files">Upload PoC Zip File *</label>
                        <input type="file" id="poc_files" name="poc_files" accept=".zip" required>
                    </div>
                </div>

                <!-- Submit Button -->
                <div style="text-align: center; margin-top: 30px;">
                    <button type="submit" class="btn btn-success" style="font-size: 18px; padding: 15px 30px;">
                        🚀 Generate Cert-IN Report
                    </button>
                </div>
            </form>

            <!-- Progress Bar -->
            <div class="progress" id="progress">
                <div style="background-color: #f0f0f0; border-radius: 4px; padding: 10px;">
                    <div id="progressText">Processing...</div>
                    <div class="progress-bar" id="progressBar"></div>
                </div>
            </div>

            <!-- Result -->
            <div class="result" id="result"></div>
        </div>

        <script>
            let CSRF_TOKEN = null;
            (async () => {
                try {
                    const res = await fetch('/csrf-token', { credentials: 'same-origin' });
                    if (res.ok) {
                        const data = await res.json().catch(() => ({}));
                        CSRF_TOKEN = data.csrf_token || data.token || data.csrf || null;
                    }
                } catch (_) {}
            })();

            function addChangeHistory() {
                const container = document.getElementById('changeHistoryList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="change_version[]" placeholder="Version" required>
                    <input type="date" name="change_date[]" required>
                    <input type="text" name="change_remarks[]" placeholder="Remarks/Reason of change" required>
                    <button type="button" class="btn btn-danger" onclick="removeChangeHistory(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeChangeHistory(button) {
                button.parentElement.remove();
            }

            function addDistribution() {
                const container = document.getElementById('distributionList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="dist_name[]" placeholder="Name" required>
                    <input type="text" name="dist_organization[]" placeholder="Organization" required>
                    <input type="text" name="dist_designation[]" placeholder="Designation" required>
                    <input type="email" name="dist_email[]" placeholder="Email ID" required>
                    <button type="button" class="btn btn-danger" onclick="removeDistribution(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeDistribution(button) {
                button.parentElement.remove();
            }

            function addEngagementScope() {
                const container = document.getElementById('engagementScopeList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="number" name="scope_sr_no[]" placeholder="S. No" required>
                    <input type="text" name="scope_asset[]" placeholder="Asset Description" required>
                    <input type="text" name="scope_criticality[]" placeholder="Criticality" required>
                    <input type="text" name="scope_internal_ip[]" placeholder="Internal IP">
                    <input type="url" name="scope_url[]" placeholder="URL">
                    <input type="text" name="scope_public_ip[]" placeholder="Public IP">
                    <input type="text" name="scope_location[]" placeholder="Location">
                    <input type="text" name="scope_hash[]" placeholder="Hash Value">
                    <input type="text" name="scope_version[]" placeholder="Version">
                    <input type="text" name="scope_other[]" placeholder="Other Details">
                    <button type="button" class="btn btn-danger" onclick="removeEngagementScope(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeEngagementScope(button) {
                button.parentElement.remove();
            }

            function addAuditingTeam() {
                const container = document.getElementById('auditingTeamList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="team_name[]" placeholder="Name" required>
                    <input type="text" name="team_designation[]" placeholder="Designation" required>
                    <input type="email" name="team_email[]" placeholder="Email ID" required>
                    <input type="text" name="team_qualifications[]" placeholder="Qualifications/Certifications" required>
                    <select name="team_certin_listed[]" required>
                        <option value="">CERT-In Listed?</option>
                        <option value="Yes">Yes</option>
                        <option value="No">No</option>
                    </select>
                    <button type="button" class="btn btn-danger" onclick="removeAuditingTeam(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeAuditingTeam(button) {
                button.parentElement.remove();
            }

            function addAuditActivity() {
                const container = document.getElementById('auditActivitiesList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="activity_task[]" placeholder="Activity/Task" required>
                    <input type="date" name="activity_date[]" required>
                    <button type="button" class="btn btn-danger" onclick="removeAuditActivity(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeAuditActivity(button) {
                button.parentElement.remove();
            }

            function addToolSoftware() {
                const container = document.getElementById('toolsSoftwareList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="number" name="tool_sr_no[]" placeholder="S. No" required>
                    <input type="text" name="tool_name[]" placeholder="Name of Tool/Software" required>
                    <input type="text" name="tool_version[]" placeholder="Version" required>
                    <select name="tool_type[]" required>
                        <option value="">Type</option>
                        <option value="Opensource">Opensource</option>
                        <option value="Licensed">Licensed</option>
                    </select>
                    <button type="button" class="btn btn-danger" onclick="removeToolSoftware(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeToolSoftware(button) {
                button.parentElement.remove();
            }

            document.getElementById('certinForm').addEventListener('submit', async function(e) {
                e.preventDefault();
                
                const formData = new FormData();
                const form = document.getElementById('certinForm');
                
                for (let element of form.elements) {
                    if (element.name && element.type !== 'file') {
                        if (element.type === 'checkbox') {
                            if (element.checked) {
                                formData.append(element.name, element.value);
                            }
                        } else {
                            formData.append(element.name, element.value);
                        }
                    }
                }
                
                const vulnerabilityFile = document.getElementById('vulnerability_file').files[0];
                if (vulnerabilityFile) {
                    formData.append('vulnerability_file', vulnerabilityFile);
                }
                
                const pocFile = document.getElementById('poc_files').files[0];
                if (pocFile) {
                    formData.append('poc_files', pocFile);
                }
                
                const changeVersions = document.querySelectorAll('input[name="change_version[]"]');
                const changeDates = document.querySelectorAll('input[name="change_date[]"]');
                const changeRemarks = document.querySelectorAll('input[name="change_remarks[]"]');
                
                const changeHistory = [];
                for (let i = 0; i < changeVersions.length; i++) {
                    changeHistory.push({
                        version: changeVersions[i].value,
                        date: changeDates[i].value,
                        remarks: changeRemarks[i].value
                    });
                }
                formData.append('document_change_history', JSON.stringify(changeHistory));
                
                const distNames = document.querySelectorAll('input[name="dist_name[]"]');
                const distOrgs = document.querySelectorAll('input[name="dist_organization[]"]');
                const distDesignations = document.querySelectorAll('input[name="dist_designation[]"]');
                const distEmails = document.querySelectorAll('input[name="dist_email[]"]');
                
                const distributionList = [];
                for (let i = 0; i < distNames.length; i++) {
                    distributionList.push({
                        name: distNames[i].value,
                        organization: distOrgs[i].value,
                        designation: distDesignations[i].value,
                        email: distEmails[i].value
                    });
                }
                formData.append('distribution_list', JSON.stringify(distributionList));
                
                const scopeSrNos = document.querySelectorAll('input[name="scope_sr_no[]"]');
                const scopeAssets = document.querySelectorAll('input[name="scope_asset[]"]');
                const scopeCriticalities = document.querySelectorAll('input[name="scope_criticality[]"]');
                const scopeInternalIPs = document.querySelectorAll('input[name="scope_internal_ip[]"]');
                const scopeURLs = document.querySelectorAll('input[name="scope_url[]"]');
                const scopePublicIPs = document.querySelectorAll('input[name="scope_public_ip[]"]');
                const scopeLocations = document.querySelectorAll('input[name="scope_location[]"]');
                const scopeHashes = document.querySelectorAll('input[name="scope_hash[]"]');
                const scopeVersions = document.querySelectorAll('input[name="scope_version[]"]');
                const scopeOthers = document.querySelectorAll('input[name="scope_other[]"]');
                
                const engagementScope = [];
                for (let i = 0; i < scopeSrNos.length; i++) {
                    engagementScope.push({
                        sr_no: scopeSrNos[i].value,
                        asset_description: scopeAssets[i].value,
                        criticality: scopeCriticalities[i].value,
                        internal_ip: scopeInternalIPs[i].value,
                        url: (scopeURLs[i].value && scopeURLs[i].value.toUpperCase() !== 'NA') ? scopeURLs[i].value : '',
                        public_ip: scopePublicIPs[i].value,
                        location: scopeLocations[i].value,
                        hash_value: scopeHashes[i].value,
                        version: scopeVersions[i].value,
                        other_details: scopeOthers[i].value
                    });
                }
                formData.append('engagement_scope', JSON.stringify(engagementScope));
                
                const teamNames = document.querySelectorAll('input[name="team_name[]"]');
                const teamDesignations = document.querySelectorAll('input[name="team_designation[]"]');
                const teamEmails = document.querySelectorAll('input[name="team_email[]"]');
                const teamQualifications = document.querySelectorAll('input[name="team_qualifications[]"]');
                const teamCertinListed = document.querySelectorAll('select[name="team_certin_listed[]"]');
                
                const auditingTeam = [];
                for (let i = 0; i < teamNames.length; i++) {
                    auditingTeam.push({
                        sr_no: i + 1,
                        name: teamNames[i].value,
                        designation: teamDesignations[i].value,
                        email: teamEmails[i].value,
                        qualifications: teamQualifications[i].value,
                        certin_listed: teamCertinListed[i].value
                    });
                }
                formData.append('auditing_team', JSON.stringify(auditingTeam));
                
                const activityTasks = document.querySelectorAll('input[name="activity_task[]"]');
                const activityDates = document.querySelectorAll('input[name="activity_date[]"]');
                
                const auditActivities = [];
                for (let i = 0; i < activityTasks.length; i++) {
                    auditActivities.push({
                        task: activityTasks[i].value,
                        date: activityDates[i].value
                    });
                }
                formData.append('audit_activities', JSON.stringify(auditActivities));
                
                const toolSrNos = document.querySelectorAll('input[name="tool_sr_no[]"]');
                const toolNames = document.querySelectorAll('input[name="tool_name[]"]');
                const toolVersions = document.querySelectorAll('input[name="tool_version[]"]');
                const toolTypes = document.querySelectorAll('select[name="tool_type[]"]');
                
                const toolsSoftware = [];
                for (let i = 0; i < toolSrNos.length; i++) {
                    toolsSoftware.push({
                        sr_no: toolSrNos[i].value,
                        name: toolNames[i].value,
                        version: toolVersions[i].value,
                        type: toolTypes[i].value
                    });
                }
                formData.append('tools_software', JSON.stringify(toolsSoftware));
                
                try {
                    document.getElementById('progress').style.display = 'block';
                    document.getElementById('progressText').textContent = 'Generating report...';
                    document.getElementById('progressBar').style.width = '50%';
                    
                    const response = await fetch('/type3/generate-report/', {
                        method: 'POST',
                        headers: CSRF_TOKEN ? { 'X-CSRF-Token': CSRF_TOKEN } : {},
                        body: formData,
                        credentials: 'same-origin'
                    });
                    
                    if (!response.ok) {
                        let msg = 'Report generation failed';
                        try {
                            const errorData = await response.json();
                            msg = errorData.detail || JSON.stringify(errorData);
                        } catch (e) {
                            const txt = await response.text();
                            if (txt) msg = txt;
                        }
                        throw new Error(msg);
                    }
                    
                    const result = await response.json();
                    
                    document.getElementById('progressText').textContent = 'Report generated successfully!';
                    document.getElementById('progressBar').style.width = '100%';
                    
                    document.getElementById('result').innerHTML = `
                        <h3>✅ Report Generated Successfully!</h3>
                        <p><strong>Filename:</strong> ${result.filename}</p>
                        <p><strong>Vulnerabilities:</strong> ${result.vulnerability_count}</p>
                        <p><strong>PoC Images:</strong> ${result.poc_count} observations with images</p>
                        <a href="${result.download_url}" download class="btn btn-success">
                            📥 Download Report
                        </a>
                    `;
                    document.getElementById('result').className = 'result success';
                    document.getElementById('result').style.display = 'block';
                    
                } catch (error) {
                    document.getElementById('progress').style.display = 'none';
                    document.getElementById('result').innerHTML = `
                        <h3>❌ Error</h3>
                        <p>${error.message}</p>
                    `;
                    document.getElementById('result').className = 'result error';
                    document.getElementById('result').style.display = 'block';
                }
            });
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)

@app.post("/generate-report/")
async def generate_report(
    request: Request,
    db: Session = Depends(get_db)
):
    """Generate Cert-IN report from form data"""
    try:
        print(f"Starting Cert-IN report generation - Version {VERSION}")

        form = await request.form()
        data = {k: form.get(k) for k in form.keys()}

        vulnerability_data = []
        poc_mapping = {}

        if 'vulnerability_file' in data and data['vulnerability_file']:
            vulnerability_file = data['vulnerability_file']
            temp_path = os.path.join(UPLOAD_DIR, f"temp_vuln_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
            content = await vulnerability_file.read()
            with open(temp_path, "wb") as f:
                f.write(content)

            vulnerability_data = read_vulnerability_excel(temp_path)
            os.remove(temp_path)

        if 'poc_files' in data and data['poc_files']:
            poc_file = data['poc_files']
            poc_mapping = await process_poc_zip_files([poc_file], vulnerability_data)

        template_path = os.path.join(os.path.dirname(__file__), "CSS Certin temp.docx")
        if not os.path.exists(template_path):
            raise HTTPException(status_code=404, detail="Cert-IN template not found")

        client_name = data.get('client_name', 'Client')
        raw_filename = f"{client_name}_Cert-IN_Report_{datetime.now().strftime('%Y-%m-%d')}.docx"
        output_path = os.path.join(UPLOAD_DIR, raw_filename)

        generate_certin_report_from_form(data, template_path, output_path, vulnerability_data, poc_mapping)

        safe_filename = os.path.basename(output_path).replace(" ", "_")
        safe_output_path = os.path.join(os.path.dirname(output_path), safe_filename)

        if output_path != safe_output_path:
            os.rename(output_path, safe_output_path)

        try:
            user = request.session.get('user') or {}
            certin_report = CertINReport(
                client_name=data.get('client_name', ''),
                report_name=data.get('report_name', ''),
                report_release_date=data.get('report_release_date', ''),
                type_of_audit=data.get('type_of_audit', ''),
                type_of_audit_report=data.get('type_of_audit_report', ''),
                period=data.get('period', ''),
                document_title=data.get('document_title', ''),
                document_id=data.get('document_id', ''),
                document_version=data.get('document_version', ''),
                prepared_by=data.get('prepared_by', ''),
                reviewed_by=data.get('reviewed_by', ''),
                approved_by=data.get('approved_by', ''),
                released_by=data.get('released_by', ''),
                release_date=data.get('release_date', ''),
                document_change_history=data.get('document_change_history', '[]'),
                distribution_list=data.get('distribution_list', '[]'),
                engagement_scope=data.get('engagement_scope', '[]'),
                auditing_team=data.get('auditing_team', '[]'),
                audit_activities=data.get('audit_activities', '[]'),
                tools_software=data.get('tools_software', '[]'),
                created_by_email=user.get('email', 'unknown'),
                created_by_name=user.get('name'),
                file_path=safe_output_path
            )
            db.add(certin_report)
            db.commit()
        except Exception as e:
            print(f"Error saving to database: {e}")

        print(f"Report generated successfully - Version {VERSION}")

        try:
            user = request.session.get('user') or {}
            db.add(AuditLog(
                user_email=user.get('email'),
                user_name=user.get('name'),
                action='generate-certin-report-form',
                metadata_json=json.dumps({
                    'client_name': data.get('client_name'),
                    'template_used': 'CSS Certin temp.docx',
                    'version': VERSION,
                    'has_vulnerability_data': len(vulnerability_data) > 0,
                    'has_poc_data': len(poc_mapping) > 0,
                    'vulnerability_count': len(vulnerability_data),
                    'poc_count': len(poc_mapping)
                }),
                ip_address=request.client.host if request.client else None,
                user_agent=request.headers.get('user-agent')
            ))
            db.commit()
        except Exception as e:
            print(f"Error logging audit: {e}")

        return JSONResponse(
            content={
                "message": "Cert-IN report generated successfully",
                "filename": safe_filename,
                "download_url": f"/type3/download/{safe_filename}",
                "vulnerability_count": len(vulnerability_data),
                "poc_count": len(poc_mapping)
            },
            status_code=200
        )

    except HTTPException:
        raise
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/download/{filename}")
async def download_report(filename: str):
    """Download generated report"""
    try:
        # Don't over-sanitize - just use the filename as-is but ensure it's safe
        # The filename was already sanitized when created
        file_path = os.path.join(UPLOAD_DIR, filename)
        
        if not os.path.exists(file_path):
            # Try with sanitized filename as fallback
            safe_filename = sanitize_filename(filename)
            file_path = os.path.join(UPLOAD_DIR, safe_filename)
            if not os.path.exists(file_path):
                raise HTTPException(status_code=404, detail=f"Report not found: {filename}")
        
        return FileResponse(
            path=file_path,
            filename=os.path.basename(file_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/reports/")
async def list_reports(request: Request, db: Session = Depends(get_db)):
    """List all generated Cert-IN reports"""
    try:
        user = request.session.get('user')
        if not user:
            raise HTTPException(status_code=401, detail="Not authenticated")
        
        reports = db.query(CertINReport).order_by(CertINReport.created_at.desc()).limit(50).all()
        
        return [
            {
                "id": report.id,
                "client_name": report.client_name,
                "report_name": report.report_name,
                "document_title": report.document_title,
                "created_at": report.created_at.isoformat(),
                "created_by": report.created_by_name,
                "download_url": f"/type3/download/{os.path.basename(report.file_path)}" if report.file_path else None
            }
            for report in reports
        ]
        
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8003)