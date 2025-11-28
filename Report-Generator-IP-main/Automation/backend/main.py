from fastapi import FastAPI, UploadFile, File, HTTPException, Form, Request, Depends
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
import os
import sys
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from datetime import datetime
import traceback
import io
import shutil
import tempfile
import matplotlib.pyplot as plt
import numpy as np
from collections import Counter
import urllib.parse
from zipfile import ZipFile
import zipfile
import glob
import hashlib
import re
import base64
import json
from docx.oxml import OxmlElement
from matplotlib.patches import Patch

# Try to import PIL for image processing, but don't fail if it's not available
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: PIL/Pillow library not available. Logo resizing will be disabled.")

# Add parent directory to path to import excel_parser
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../..'))
from excel_parser import parse_excel_data, format_for_template

# Version tracking for code changes
# Format: MAJOR.MINOR.PATCH
# Increment MAJOR for significant structural changes
# Increment MINOR for feature additions or improvements
# Increment PATCH for bug fixes
VERSION = "1.0.7"  # Updated PATCH version for splitting parsing logic

app = FastAPI()
from sqlalchemy.orm import Session
from db import get_db, AuditLog, DashboardDataset, ProjectHistory, ProjectEvaluationEvent


# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to store uploaded files
UPLOAD_DIR = "uploads"
# Define the fixed path for the document control file
DOC_CONTROL_FILE = os.path.join(UPLOAD_DIR, "Comprehensive_Vulnerability_Template (2).xlsx")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(os.path.join(UPLOAD_DIR, "screenshots"), exist_ok=True)

# Define the order of severities
severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4, "Unknown": 5}

# Serve static files from the parent directory (Automation2.0)
# Remove conflicting broad static mount; rely on top-level app mounts
# app.mount("/static", StaticFiles(directory="../"), name="static")

@app.get("/", response_class=HTMLResponse)
async def root():
    with open(os.path.join(os.path.dirname(__file__), "../report_formats.html"), encoding="utf-8") as f:
        return f.read()

# --- Dashboard APIs (list datasets & logs) ---
from fastapi import APIRouter
router = APIRouter()

# Create a separate router for MitKat routes
mitkat_router = APIRouter()

def is_dashboard_allowed(request: Request) -> bool:
    user = request.session.get('user')
    if not user:
        return False
    if os.getenv("TEST_MODE", "false").lower() == "true":
        return True
    return user.get('email', '').lower() in {"sarvesh.salgaonkar@cybersmithsecure.com", "smith.gonsalves@cybersmithsecure.com", "developer@cybersmithsecure.com"}

@router.post("/dashboard/upload")
async def upload_dashboard_dataset(
    request: Request,
    db: Session = Depends(get_db),
    file: UploadFile = File(...),
    title: str = Form(...),
    project_name: str = Form("")
):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    # Persist file under uploads/dashboard_datasets
    datasets_dir = os.path.join(UPLOAD_DIR, "dashboard_datasets")
    os.makedirs(datasets_dir, exist_ok=True)
    # Sanitize filename to prevent XSS/path traversal
    original_name = file.filename or "upload.bin"
    safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", original_name)
    file_path = os.path.join(datasets_dir, f"{datetime.now().strftime('%Y%m%d%H%M%S')}_{safe_name}")
    with open(file_path, "wb") as f:
        f.write(await file.read())
    user = request.session.get('user') or {}
    ds = DashboardDataset(
        title=title,
        project_name=project_name,
        file_path=file_path,
        uploaded_by_email=user.get('email', 'unknown'),
        uploaded_by_name=user.get('name')
    )
    db.add(ds)
    db.add(AuditLog(
        user_email=user.get('email'),
        user_name=user.get('name'),
        action='dashboard-upload',
        metadata_json=json.dumps({'title': title, 'project_name': project_name, 'filename': file.filename}),
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get('user-agent')
    ))
    db.commit()
    db.refresh(ds)
    return {
        "id": ds.id,
        "title": ds.title,
        "project_name": ds.project_name,
        "uploaded_by_email": ds.uploaded_by_email,
        "uploaded_by_name": ds.uploaded_by_name,
        "uploaded_at": ds.uploaded_at.isoformat()
    }

@router.get("/dashboard-datasets")
async def list_dashboard_datasets(request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    rows = db.query(DashboardDataset).order_by(DashboardDataset.uploaded_at.desc()).limit(200).all()
    return [
        {
            "id": r.id,
            "title": r.title,
            "project_name": r.project_name,
            "uploaded_by_email": r.uploaded_by_email,
            "uploaded_by_name": r.uploaded_by_name,
            "uploaded_at": r.uploaded_at.isoformat()
        }
        for r in rows
    ]

@router.get("/dashboard/latest-dataset")
async def latest_dashboard_dataset(request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    row = db.query(DashboardDataset).order_by(DashboardDataset.uploaded_at.desc()).first()
    if not row:
        return {"status": "empty"}
    
    # Read the Excel file and return the data
    try:
        import pandas as pd
        df = pd.read_excel(row.file_path)
        df.columns = df.columns.str.strip()
        
        # Normalize headers
        header_map = {
            'empid': 'EmpID', 'emp id': 'EmpID', 'employee id': 'EmpID',
            'reportedby': 'ReportedBy', 'reported by': 'ReportedBy', 'name': 'ReportedBy',
            'project': 'Project',
            'severity': 'Severity',
            'date': 'Date', 'reportedon': 'Date',
            'vulnerabilityname': 'VulnerabilityName', 'vulnerability name': 'VulnerabilityName'
        }
        
        # Convert to list of dictionaries
        rows = []
        for _, row_data in df.iterrows():
            row_dict = {}
            for col in df.columns:
                normalized_col = header_map.get(col.strip().lower(), col)
                row_dict[normalized_col] = row_data[col]
            rows.append(row_dict)
        
        return {
            "id": row.id,
            "title": row.title,
            "project_name": row.project_name,
            "uploaded_by_email": row.uploaded_by_email,
            "uploaded_by_name": row.uploaded_by_name,
            "uploaded_at": row.uploaded_at.isoformat(),
            "rows": rows
        }
    except Exception as e:
        print(f"Error reading dashboard data: {e}")
        return {"status": "error", "message": str(e)}

@router.get("/dashboard-datasets/{dataset_id}/file")
async def download_dashboard_dataset_file(dataset_id: int, request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    ds = db.query(DashboardDataset).filter(DashboardDataset.id == dataset_id).first()
    if not ds or not os.path.exists(ds.file_path):
        raise HTTPException(status_code=404, detail="Dataset not found")
    # Ensure the file serves with a safe filename
    download_name = os.path.basename(ds.file_path)
    download_name = re.sub(r"[^A-Za-z0-9._-]", "_", download_name)
    return FileResponse(path=ds.file_path, filename=download_name, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

@router.get("/audit-logs")
async def list_audit_logs(request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    rows = db.query(AuditLog).order_by(AuditLog.created_at.desc()).limit(300).all()
    return [
        {
            "id": r.id,
            "user_email": r.user_email,
            "user_name": r.user_name,
            "action": r.action,
            "metadata_json": r.metadata_json,
            "ip_address": r.ip_address,
            "user_agent": r.user_agent,
            "created_at": r.created_at.isoformat()
        }
        for r in rows
    ]

@router.post("/project-history/update")
async def update_project_history(
    request: Request,
    db: Session = Depends(get_db),
    project_name: str = Form(...),
    total_vulnerabilities: int = Form(...),
    unique_vulnerabilities: int = Form(...),
    critical_count: int = Form(0),
    high_count: int = Form(0),
    medium_count: int = Form(0),
    low_count: int = Form(0),
    informational_count: int = Form(0),
    vulnerability_details: str = Form("")
):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    
    user = request.session.get('user') or {}
    
    # Check if project already exists
    existing_project = db.query(ProjectHistory).filter(
        ProjectHistory.project_name == project_name
    ).first()
    
    if existing_project:
        # Update existing project
        existing_project.total_vulnerabilities += total_vulnerabilities
        existing_project.unique_vulnerabilities = unique_vulnerabilities  # Update with latest count
        existing_project.critical_count += critical_count
        existing_project.high_count += high_count
        existing_project.medium_count += medium_count
        existing_project.low_count += low_count
        existing_project.informational_count += informational_count
        existing_project.last_evaluated = datetime.utcnow()
        existing_project.evaluation_count += 1
        existing_project.vulnerability_details = vulnerability_details
        
        # Record event row as well
        event = ProjectEvaluationEvent(
            project_name=project_name,
            total_vulnerabilities=total_vulnerabilities,
            unique_vulnerabilities=unique_vulnerabilities,
            critical_count=critical_count,
            high_count=high_count,
            medium_count=medium_count,
            low_count=low_count,
            informational_count=informational_count,
            vulnerability_details=vulnerability_details,
            uploaded_by_email=user.get('email'),
            uploaded_by_name=user.get('name')
        )
        db.add(event)
        db.commit()
        db.refresh(existing_project)
        
        return {
            "id": existing_project.id,
            "project_name": existing_project.project_name,
            "total_vulnerabilities": existing_project.total_vulnerabilities,
            "unique_vulnerabilities": existing_project.unique_vulnerabilities,
            "critical_count": existing_project.critical_count,
            "high_count": existing_project.high_count,
            "medium_count": existing_project.medium_count,
            "low_count": existing_project.low_count,
            "informational_count": existing_project.informational_count,
            "first_evaluated": existing_project.first_evaluated.isoformat(),
            "last_evaluated": existing_project.last_evaluated.isoformat(),
            "evaluation_count": existing_project.evaluation_count,
            "uploaded_by_email": existing_project.uploaded_by_email,
            "uploaded_by_name": existing_project.uploaded_by_name,
            "vulnerability_details": existing_project.vulnerability_details
        }
    else:
        # Create new project
        new_project = ProjectHistory(
            project_name=project_name,
            total_vulnerabilities=total_vulnerabilities,
            unique_vulnerabilities=unique_vulnerabilities,
            critical_count=critical_count,
            high_count=high_count,
            medium_count=medium_count,
            low_count=low_count,
            informational_count=informational_count,
            uploaded_by_email=user.get('email', 'unknown'),
            uploaded_by_name=user.get('name'),
            vulnerability_details=vulnerability_details
        )
        
        db.add(new_project)
        # Record first event row
        event = ProjectEvaluationEvent(
            project_name=project_name,
            total_vulnerabilities=total_vulnerabilities,
            unique_vulnerabilities=unique_vulnerabilities,
            critical_count=critical_count,
            high_count=high_count,
            medium_count=medium_count,
            low_count=low_count,
            informational_count=informational_count,
            vulnerability_details=vulnerability_details,
            uploaded_by_email=user.get('email'),
            uploaded_by_name=user.get('name')
        )
        db.add(event)
        db.commit()
        db.refresh(new_project)
        
        return {
            "id": new_project.id,
            "project_name": new_project.project_name,
            "total_vulnerabilities": new_project.total_vulnerabilities,
            "unique_vulnerabilities": new_project.unique_vulnerabilities,
            "critical_count": new_project.critical_count,
            "high_count": new_project.high_count,
            "medium_count": new_project.medium_count,
            "low_count": new_project.low_count,
            "informational_count": new_project.informational_count,
            "first_evaluated": new_project.first_evaluated.isoformat(),
            "last_evaluated": new_project.last_evaluated.isoformat(),
            "evaluation_count": new_project.evaluation_count,
            "uploaded_by_email": new_project.uploaded_by_email,
            "uploaded_by_name": new_project.uploaded_by_name,
            "vulnerability_details": new_project.vulnerability_details
        }

@router.get("/project-history")
async def get_project_history(request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    
    projects = db.query(ProjectHistory).order_by(ProjectHistory.last_evaluated.desc()).all()
    
    # Calculate overall totals
    total_projects = len(projects)
    overall_total_vulnerabilities = sum(p.total_vulnerabilities for p in projects)
    overall_unique_vulnerabilities = sum(p.unique_vulnerabilities for p in projects)
    overall_critical = sum(p.critical_count for p in projects)
    overall_high = sum(p.high_count for p in projects)
    overall_medium = sum(p.medium_count for p in projects)
    overall_low = sum(p.low_count for p in projects)
    overall_informational = sum(p.informational_count for p in projects)
    
    # Build monthly summaries based on events
    events = db.query(ProjectEvaluationEvent).order_by(ProjectEvaluationEvent.created_at.desc()).all()
    monthly = {}
    for ev in events:
        month_key = ev.created_at.strftime('%Y-%m')
        if month_key not in monthly:
            monthly[month_key] = {
                'month': month_key,
                'total_vulnerabilities': 0,
                'projects': set()
            }
        monthly[month_key]['total_vulnerabilities'] += (ev.total_vulnerabilities or 0)
        monthly[month_key]['projects'].add(ev.project_name)

    monthly_summary = [
        {
            'month': k,
            'total_vulnerabilities': v['total_vulnerabilities'],
            'projects_evaluated': len(v['projects'])
        }
        for k, v in sorted(monthly.items(), key=lambda x: x[0], reverse=True)
    ]

    return {
        "summary": {
            "total_projects": total_projects,
            "total_vulnerabilities": overall_total_vulnerabilities,
            "unique_vulnerabilities": overall_unique_vulnerabilities,
            "critical_count": overall_critical,
            "high_count": overall_high,
            "medium_count": overall_medium,
            "low_count": overall_low,
            "informational_count": overall_informational
        },
        "monthly": monthly_summary,
        "projects": [
            {
                "id": p.id,
                "project_name": p.project_name,
                "total_vulnerabilities": p.total_vulnerabilities,
                "unique_vulnerabilities": p.unique_vulnerabilities,
                "critical_count": p.critical_count,
                "high_count": p.high_count,
                "medium_count": p.medium_count,
                "low_count": p.low_count,
                "informational_count": p.informational_count,
                "first_evaluated": p.first_evaluated.isoformat(),
                "last_evaluated": p.last_evaluated.isoformat(),
                "evaluation_count": p.evaluation_count,
                "uploaded_by_email": p.uploaded_by_email,
                "uploaded_by_name": p.uploaded_by_name,
                "vulnerability_details": p.vulnerability_details
            }
            for p in projects
        ]
    }

@router.get("/project-history/monthly/{month_key}")
async def get_project_history_monthly(month_key: str, request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    # Expect month_key in YYYY-MM
    try:
        start = datetime.strptime(month_key + "-01", "%Y-%m-%d")
        # compute first day of next month
        if start.month == 12:
            end = datetime(start.year + 1, 1, 1)
        else:
            end = datetime(start.year, start.month + 1, 1)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid month format. Use YYYY-MM")

    # Pull events in the month range
    events = db.query(ProjectEvaluationEvent).filter(
        ProjectEvaluationEvent.created_at >= start,
        ProjectEvaluationEvent.created_at < end
    ).all()

    severity_counts = {
        'Critical': 0,
        'High': 0,
        'Medium': 0,
        'Low': 0,
        'Informational': 0
    }
    project_map = {}

    for ev in events:
        # Update overall severity
        severity_counts['Critical'] += (ev.critical_count or 0)
        severity_counts['High'] += (ev.high_count or 0)
        severity_counts['Medium'] += (ev.medium_count or 0)
        severity_counts['Low'] += (ev.low_count or 0)
        severity_counts['Informational'] += (ev.informational_count or 0)

        # Update per-project
        p = project_map.setdefault(ev.project_name, {
            'project_name': ev.project_name,
            'total_vulnerabilities': 0,
            'critical_count': 0,
            'high_count': 0,
            'medium_count': 0,
            'low_count': 0,
            'informational_count': 0
        })
        p['total_vulnerabilities'] += (ev.total_vulnerabilities or 0)
        p['critical_count'] += (ev.critical_count or 0)
        p['high_count'] += (ev.high_count or 0)
        p['medium_count'] += (ev.medium_count or 0)
        p['low_count'] += (ev.low_count or 0)
        p['informational_count'] += (ev.informational_count or 0)

    projects = sorted(project_map.values(), key=lambda x: x['total_vulnerabilities'], reverse=True)

    return {
        'month': month_key,
        'projects_evaluated': len(projects),
        'total_vulnerabilities': sum(p['total_vulnerabilities'] for p in projects),
        'severity_counts': severity_counts,
        'projects': projects
    }

@router.get("/project-history/events")
async def get_project_history_events(request: Request, db: Session = Depends(get_db)):
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    events = db.query(ProjectEvaluationEvent).order_by(ProjectEvaluationEvent.created_at.desc()).limit(1000).all()
    out = {}
    for ev in events:
        out.setdefault(ev.project_name, []).append({
            'created_at': ev.created_at.isoformat(),
            'total_vulnerabilities': ev.total_vulnerabilities,
            'unique_vulnerabilities': ev.unique_vulnerabilities,
            'critical_count': ev.critical_count,
            'high_count': ev.high_count,
            'medium_count': ev.medium_count,
            'low_count': ev.low_count,
            'informational_count': ev.informational_count,
            'uploaded_by_email': ev.uploaded_by_email,
            'uploaded_by_name': ev.uploaded_by_name
        })
    return out

app.include_router(router)
app.include_router(mitkat_router)

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as f:
            f.write(await file.read())
        return JSONResponse(content={"filename": file.filename}, status_code=200)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-screenshots/")
async def upload_screenshots(files: list[UploadFile] = File(...)):
    """
    Upload multiple screenshot files.
    Returns a dictionary of the saved filenames.
    """
    try:
        saved_files = []
        for file in files:
            # Only allow image files
            if not file.content_type.startswith('image/'):
                continue
                
            # Sanitize filename and lock to screenshots directory
            original_name = file.filename or "image.png"
            safe_name = re.sub(r"[^A-Za-z0-9._-]", "_", original_name)
            screenshots_dir = os.path.join(UPLOAD_DIR, "screenshots")
            os.makedirs(screenshots_dir, exist_ok=True)
            file_location = os.path.join(screenshots_dir, safe_name)
            with open(file_location, "wb") as f:
                f.write(await file.read())
            saved_files.append(safe_name)
        
        return JSONResponse(
            content={"filenames": saved_files},
            status_code=200
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# Function to determine severity from CVSS score
def get_severity_from_cvss(cvss):
    print(f"Determining severity for CVSS: {cvss}, type: {type(cvss)}")  # Debug output
    
    if cvss is None:
        print("CVSS is None, defaulting to Informational")
        return "Informational"
    
    try:
        cvss_float = float(cvss)
        if cvss_float >= 9.0 and cvss_float <= 10.0:
            print(f"CVSS {cvss_float} classified as Critical")
            return "Critical"
        elif cvss_float >= 7.0 and cvss_float < 9.0:
            print(f"CVSS {cvss_float} classified as High")
            return "High"
        elif cvss_float >= 4.0 and cvss_float < 7.0:
            print(f"CVSS {cvss_float} classified as Medium")
            return "Medium"
        elif cvss_float > 0.0 and cvss_float < 4.0:
            print(f"CVSS {cvss_float} classified as Low")
            return "Low"
        else:
            print(f"CVSS {cvss_float} classified as Informational")
            return "Informational"
    except (ValueError, TypeError) as e:
        print(f"Error converting CVSS to float: {e}, defaulting to Informational")
        return "Informational"

# Function to get severity colors based on severity level
def get_severity_colors(severity):
    colors = {
        "Critical": {"row1_bg": "#990000", "row2_bg": "#FF3333", "font": "#990000"},
        "High": {"row1_bg": "#FF0000", "row2_bg": "#FF6666", "font": "#FF0000"},
        "Medium": {"row1_bg": "#FFCC00", "row2_bg": "#FFCC66", "font": "#FFCC00"},
        "Low": {"row1_bg": "#00b050", "row2_bg": "#99CC33", "font": "#00b050"},
        "Informational": {"row1_bg": "#0070c0", "row2_bg": "#66CCFF", "font": "#0070c0"}
    }
    return colors.get(severity, colors["Informational"])

# Function to generate more detailed vulnerability content based on vulnerability name
def get_vulnerability_description(vuln_name):
    """
    Generate a detailed description paragraph for the vulnerability based on OWASP guidance.
    """
    descriptions = {
        "SQL Injection": """SQL Injection is a code injection technique that exploits vulnerabilities in the interface between web applications and database servers. The vulnerability occurs when user input is incorrectly filtered and directly included in SQL statements, allowing attackers to manipulate the structure of the SQL query. This can lead to unauthorized access to sensitive data, bypassing authentication mechanisms, and in severe cases, compromising the underlying server or performing denial of service attacks. The primary cause is insufficient validation and sanitization of user-supplied input used in database queries.""",
        
        "Cross-Site Scripting": """Cross-Site Scripting (XSS) is a security vulnerability that allows attackers to inject malicious client-side scripts into web pages viewed by other users. These attacks occur when an application includes untrusted data in a new web page without proper validation or encoding. XSS vulnerabilities can be found in various forms: Reflected XSS where malicious script is reflected off the web server in an error message or search result; Stored XSS where the malicious script is permanently stored on target servers; and DOM-based XSS where vulnerability exists in client-side code. The impact of XSS can range from session theft to complete compromise of user accounts.""",
        
        "File Upload": """Unrestricted File Upload vulnerabilities allow attackers to upload files of dangerous types or content to the server. This vulnerability occurs when an application fails to properly validate the type, content, size, or naming convention of uploaded files. Attackers can exploit this to upload web shells, malware, or other malicious files that allow them to execute arbitrary code on the server. This may lead to complete system compromise, data theft, or defacement of websites. The vulnerability is particularly dangerous when uploaded files are stored in web-accessible locations, allowing direct execution through browser requests.""",
        
        "Insecure Direct Object References": """Insecure Direct Object References (IDOR) occur when an application provides direct access to objects based on user-supplied input without appropriate authorization checks. This vulnerability allows attackers to bypass authorization and directly access resources in the system, potentially exposing sensitive data that should be protected. IDOR vulnerabilities commonly manifest as manipulated parameters within requests that reveal sensitive objects like files, database records, or user accounts. This can lead to unauthorized access to sensitive information, modification of data belonging to other users, or privilege escalation allowing attackers to gain administrative access to the application.""",
        
        "Sensitive Data Exposure": """Sensitive Data Exposure occurs when an application fails to adequately protect sensitive information such as financial data, healthcare information, personally identifiable information (PII), or authentication credentials. This vulnerability can manifest through various channels including insufficient encryption of data at rest or in transit, improper handling of sensitive information in logs or error messages, and inadequate access controls. Attackers can exploit these weaknesses to access sensitive data through packet sniffing, man-in-the-middle attacks, or by directly accessing poorly protected storage. The impact can include identity theft, financial fraud, privacy violations, and regulatory compliance failures."""
    }
    
    # Try to find a match in our predefined descriptions
    for key in descriptions:
        if key.lower() in vuln_name.lower():
            return descriptions[key]
    
    # Default description for any vulnerability not specifically defined
    return f"""This vulnerability ({vuln_name}) allows attackers to potentially compromise the application security through improper input validation and processing mechanisms. The vulnerability exists due to insufficient security controls in the application code, which could allow malicious actors to manipulate system behavior in unintended ways. Such vulnerabilities often arise from coding practices that do not account for all possible edge cases or user inputs, creating potential attack vectors. When discovered, these issues should be addressed promptly to prevent potential exploitation."""

def get_vulnerability_impact(vuln_name):
    """
    Generate a detailed impact paragraph for the vulnerability based on OWASP guidance.
    """
    impacts = {
        "SQL Injection": """The impact of SQL Injection vulnerabilities can be severe and far-reaching for an organization. Successful exploitation can lead to unauthorized access to sensitive database information, including personal user data, financial records, and intellectual property. Attackers may bypass authentication mechanisms, gaining unauthorized administrative access to the application. In more severe cases, attackers might execute commands on the database server, potentially leading to complete server compromise. This can result in significant financial losses, regulatory penalties for data breaches, damage to company reputation, and loss of customer trust. Business operations may be disrupted if data is stolen, manipulated, or deleted entirely.""",
        
        "Cross-Site Scripting": """The exploitation of Cross-Site Scripting vulnerabilities can have significant consequences for both users and the organization. Attackers can hijack user sessions, steal authentication cookies, and impersonate legitimate users to access sensitive information or perform unauthorized actions. They may also inject malicious content that defaces websites, leads to redirection to malicious sites, or installs malware on user devices. For organizations, this can lead to reputational damage, loss of customer trust, potential data breaches, and regulatory compliance issues. XSS attacks targeting administrative users can result in complete application compromise, allowing attackers to access backend systems, modify server-side logic, or obtain sensitive business data.""",
        
        "File Upload": """Unrestricted file upload vulnerabilities can have devastating impacts on organization security. Successfully exploited vulnerabilities allow attackers to gain remote code execution capabilities on the server, potentially leading to complete system compromise. This can result in unauthorized access to databases, internal networks, and sensitive enterprise systems. Attackers may establish persistence through backdoors, maintain long-term access to compromised systems, and use the server as a platform to attack other internal systems. Organizations face significant risks including data theft, intellectual property loss, financial loss, service disruption due to regulatory penalties. The reputational damage from such breaches can lead to loss of customer confidence and business opportunities.""",
        
        "Insecure Direct Object References": """The business impact of Insecure Direct Object References can be substantial. Exploitation allows attackers to access unauthorized data across user boundaries, potentially exposing sensitive personal information, financial data, or proprietary business information. This can lead to privacy violations, identity theft, financial fraud, or industrial espionage. Organizations may face significant regulatory penalties for data protection failures, particularly under frameworks like GDPR, HIPAA, or PCI-DSS. The reputational damage from such data breaches can severely impact customer trust and result in customer churn. Additionally, competitors may gain access to confidential business information, leading to competitive disadvantage and potential revenue loss.""",
        
        "Sensitive Data Exposure": """Exposure of sensitive data can have catastrophic consequences for both individuals and organizations. When personal information is compromised, individuals may become victims of identity theft, financial fraud, or reputational damage. For organizations, the impacts include regulatory fines and penalties under data protection laws such as GDPR, HIPAA, or PCI-DSS, which can amount to millions of dollars. Legal costs may escalate through lawsuits from affected parties seeking damages. Business reputation suffers significantly, leading to loss of customer trust, decreased market share, and reduced stock value for public companies. Operational disruptions can occur while addressing the breach, and long-term business relationships may be damaged if partner data is compromised."""
    }
    
    # Try to find a match in our predefined impacts
    for key in impacts:
        if key.lower() in vuln_name.lower():
            return impacts[key]
    
    # Default impact for any vulnerability not specifically defined
    return f"""Exploitation of this vulnerability ({vuln_name}) could lead to unauthorized access, data breaches, reputation damage, and potential regulatory compliance violations. Customers' trust may be eroded, resulting in business loss. The organization may face financial penalties from regulatory bodies for failing to protect sensitive information adequately. Remediation costs can be substantial, including not only the technical fixes but also potential legal fees, customer notification expenses, and increased security monitoring. In severe cases, business operations might be disrupted, leading to revenue loss and long-term damage to the organization's brand and market position."""

def get_vulnerability_recommendation(vuln_name):
    """
    Generate detailed remediation recommendations based on OWASP guidance.
    """
    recommendations = {
        "SQL Injection": """To remediate SQL Injection vulnerabilities, implement parameterized queries (prepared statements) for all database operations, which ensures that user input is treated as data rather than executable code. Apply the principle of least privilege to database accounts used by applications, limiting their capabilities to only necessary operations. Implement input validation on both client and server sides, with a whitelist approach that accepts only known-good input. Utilize stored procedures and ORM (Object-Relational Mapping) frameworks that handle SQL escaping automatically. Implement a WAF (Web Application Firewall) as an additional layer of protection to filter malicious SQL injection attempts. Regularly scan application code for SQL injection vulnerabilities and conduct penetration testing to verify the effectiveness of implemented controls.""",
        
        "Cross-Site Scripting": """To effectively mitigate Cross-Site Scripting vulnerabilities, implement context-sensitive output encoding for all user-supplied data before rendering it in HTML, JavaScript, CSS, or URL contexts. Use Content Security Policy (CSP) headers to restrict sources of executable scripts and prevent execution of inline JavaScript. Apply input validation with a whitelist approach to filter potentially dangerous inputs before processing. Utilize modern web frameworks that automatically handle XSS protection through built-in template escaping mechanisms. For cookie security, implement the HttpOnly flag to prevent JavaScript access and use the Secure flag for transmission over HTTPS only. Keep all frameworks and libraries updated to ensure known XSS vulnerabilities are patched. Regularly conduct security testing including both automated scanning and manual code review to identify potential XSS vulnerabilities.""",
        
        "File Upload": """To securely handle file uploads, implement strict file type validation using content inspection (checking file headers) rather than relying only on file extensions. Enforce file size limits to prevent denial of service attacks through oversized file uploads. Store uploaded files outside the web root or application directory to prevent direct execution, and if files must be publicly accessible, store them in a different domain or subdomain. Generate random file names to prevent directory traversal attempts and consistent naming patterns that could be guessed. Scan uploaded files with antivirus software or integrate with security services that can detect malicious content. Implement proper access controls to ensure uploaded files are only accessible by authorized users. Consider using a CDN or dedicated file storage service that specializes in secure file handling for high-risk environments.""",
        
        "Insecure Direct Object References": """To remediate Insecure Direct Object References, implement proper access control checks for every request that accesses data objects, verifying that the authenticated user has appropriate permissions for the requested resource. Replace direct references to objects with indirect references or tokens that are mapped server-side to actual resource identifiers. Use server-side session objects to store and verify user authorization to access specific resources. Create a centralized authorization mechanism that provides consistent enforcement across the application. Implement the principle of least privilege, ensuring users only have access to resources necessary for their legitimate functions. Log access attempts to sensitive resources and configure alerts for potential access control violations. Regularly review and test access control mechanisms through security assessments and penetration testing.""",
        
        "Sensitive Data Exposure": """To protect sensitive data effectively, implement strong encryption for data both at rest and in transit using industry-standard algorithms and appropriate key management. Use TLS (at least version 1.2 or higher) for all connections that transmit sensitive information, and enforce this through HSTS headers. Store passwords using strong adaptive hashing algorithms with salts (such as bcrypt, Argon2, or PBKDF2). Disable caching for responses containing sensitive data using appropriate cache-control headers. Minimize the collection and retention of sensitive data – only collect what is absolutely necessary for business functions and purge unnecessary data regularly. Implement proper key management practices including secure key generation, storage, distribution, and rotation. Classify data based on sensitivity and apply appropriate controls based on classification. Regularly audit systems for sensitive data handling compliance and conduct penetration testing focused on data protection controls."""
    }
    
    # Try to find a match in our predefined recommendations
    for key in recommendations:
        if key.lower() in vuln_name.lower():
            return recommendations[key]
    
    # Default recommendation for any vulnerability not specifically defined
    return f"""To remediate this {vuln_name} vulnerability, implement proper input validation with a strong allowlist approach that verifies all user inputs against expected formats and values. Follow the principle of least privilege throughout the application, ensuring components only have access to resources necessary for their legitimate function. Develop and maintain secure coding practices including regular code reviews and security-focused testing. Keep all software components, including frameworks and libraries, updated with security patches. Implement defense in depth with multiple layers of security controls, so that if one layer fails, others will still provide protection. Conduct regular security assessments and penetration testing to identify vulnerabilities before they can be exploited. Create a security-focused development lifecycle that incorporates security at every stage from design to deployment and maintenance."""

def normalize_filename(fname):
    """Normalize filenames for robust matching: strip, lower, keep extension."""
    return fname.strip().lower() if fname else ''

# Function to parse the uploaded Excel file for vulnerability data
def parse_vulnerabilities_excel(file_path):
    """
    Parse the uploaded Excel file using standardized parser and return a dict grouped by IP.
    Each key is an IP, and the value is a list of vulnerability dicts for that IP.
    """
    print(f"Parsing vulnerability data from: {file_path}")
    try:
        # Try to use standardized parser first
        try:
            parsed_data = parse_excel_data(file_path)
            print(f"✅ Using standardized Excel parser")
            print(f"   Metadata: {parsed_data['metadata']}")
            print(f"   Vulnerabilities: {len(parsed_data['vulnerabilities'])}")
            
            vulnerabilities_by_ip = {}
            for index, vuln in enumerate(parsed_data['vulnerabilities']):
                # Get IP from affected_asset or ip_url_app
                ip = vuln.get('affected_asset', '') or vuln.get('ip_url_app', '')
                if not ip:
                    print(f"Row {index} missing IP, skipping.")
                    continue
                
                # Convert steps to steps_with_screenshots format
                steps_with_screenshots = []
                if vuln.get('steps'):
                    for step in vuln['steps']:
                        steps_with_screenshots.append({
                            'text': step['content'],
                            'screenshot': ''  # Screenshots handled separately via PoC mapping
                        })
                
                # Convert to main.py format
                vulnerability = {
                    'ip': ip,
                    'name': vuln.get('observation', ''),
                    'description': vuln.get('detailed_observation', '') or vuln.get('observation_summary', ''),
                    'impact': '',  # Not in standardized format, will be generated
                    'severity': vuln.get('severity', 'Medium'),
                    'cvss': vuln.get('cvss', ''),
                    'vulnerable_url': vuln.get('ip_url_app', ''),
                    'vulnerable_parameter': '',  # Not in standardized format
                    'remediation': vuln.get('recommendation', ''),
                    'steps_with_screenshots': steps_with_screenshots,
                    'sr_no': str(vuln.get('sr_no', f"VUL-{index+1:03d}")).replace('VULN-', 'VUL-').strip(),
                    'no_vuln': False
                }
                
                if ip not in vulnerabilities_by_ip:
                    vulnerabilities_by_ip[ip] = []
                vulnerabilities_by_ip[ip].append(vulnerability)
                print(f"Parsed vulnerability Sr No: {vulnerability['sr_no']} for IP: {ip} at index {index}")
            
            print(f"Total IPs extracted (standardized parser): {len(vulnerabilities_by_ip)}")
            return vulnerabilities_by_ip
            
        except Exception as std_parse_error:
            print(f"⚠️  Standardized parser failed, falling back to legacy parser: {std_parse_error}")
            traceback.print_exc()
        
        # Fallback to original parsing logic
        df = pd.read_excel(file_path)
        print(f"Vulnerability Excel file read successfully. Columns: {df.columns.tolist()}")
        df.columns = df.columns.str.strip()
        vulnerabilities_by_ip = {}
        for index, row in df.iterrows():
            try:
                ip = str(row.get('IP', '')).strip()
                if not ip:
                    print(f"Row {index} missing IP, skipping.")
                    continue
                sr_no_val = str(row.get('Sr No', '')).strip()
                if sr_no_val.lower() == 'no vulnerability' or sr_no_val.lower() == 'no vulnerabilities':
                    # Special handling for IPs with no vulnerabilities
                    vulnerability = {
                        'ip': ip,
                        'name': 'No Vulnerability Found',
                        'description': 'No vulnerabilities were identified for this IP during the assessment period.',
                        'impact': '',
                        'severity': 'Informational',
                        'cvss': '',
                        'vulnerable_url': '',
                        'vulnerable_parameter': '',
                        'remediation': '',
                        'steps_with_screenshots': [],
                        'sr_no': sr_no_val,
                        'no_vuln': True
                    }
                    if ip not in vulnerabilities_by_ip:
                        vulnerabilities_by_ip[ip] = []
                    vulnerabilities_by_ip[ip].append(vulnerability)
                    print(f"Parsed NO VULNERABILITY for IP: {ip} at index {index}")
                    continue
                cvss_score = row.get('CVSS Score', row.get('CVSS', None))
                vulnerable_url = row.get('Vulnerable URL', row.get('Vulnerable Parameter', None))
                parameter_value = row.get('Vulnerable Parameter', None)
                if isinstance(parameter_value, float) and pd.isna(parameter_value):
                    vulnerable_parameter = None
                else:
                    vulnerable_parameter = str(parameter_value)
                steps_with_screenshots = []
                steps_text = str(row.get('Steps', '')).splitlines()
                screenshot_cols = [col for col in df.columns if col.strip().lower().startswith('screenshot')]
                for i, step_text in enumerate(steps_text):
                    screenshot_filename = ''
                    if i < len(screenshot_cols):
                        screenshot_filename = str(row.get(screenshot_cols[i], '')).strip().lower()
                    steps_with_screenshots.append({
                        'text': step_text.strip(),
                        'screenshot': screenshot_filename
                    })
                severity = get_severity_from_cvss(cvss_score)
                vuln_name = row.get('Vulnerability Name', '')
                description = row.get('Description', '') or get_vulnerability_description(vuln_name)
                remediation = row.get('Remediation', '') or get_vulnerability_recommendation(vuln_name)
                impact = row.get('Impact', '') or get_vulnerability_impact(vuln_name)
                vulnerability = {
                    'ip': ip,
                    'name': vuln_name,
                    'description': description,
                    'impact': impact,
                    'severity': severity,
                    'cvss': cvss_score,
                    'vulnerable_url': vulnerable_url,
                    'vulnerable_parameter': vulnerable_parameter,
                    'remediation': remediation,
                    'steps_with_screenshots': steps_with_screenshots,
                    'sr_no': str(row.get('Sr No', f"VUL-{index+1:03d}")).replace('VULN-', 'VUL-').strip(),
                }
                if ip not in vulnerabilities_by_ip:
                    vulnerabilities_by_ip[ip] = []
                vulnerabilities_by_ip[ip].append(vulnerability)
                print(f"Parsed vulnerability Sr No: {vulnerability['sr_no']} for IP: {ip} at index {index}")
            except Exception as e:
                print(f"Error processing row {index}: {e}")
                continue
        print(f"Total IPs extracted: {len(vulnerabilities_by_ip)}")
        return vulnerabilities_by_ip
    except Exception as e:
        print(f"Error parsing vulnerability Excel: {e}")
        return {}

# Function to parse the fixed document control file
def parse_doc_control_excel(file_path):
    """
    Parse the fixed document control Excel file and return document control details
    including assessment scope, dates, note, and amendment log.
    """
    print(f"Parsing document control data from: {file_path}")
    doc_control_data = {
        'assessment_scope': '',
        'assessment_start_date': '',
        'assessment_end_date': '',
        'assessment_note': '',  # Add assessment_note
        'amendment_log': [],
        'reviewed_by': '',
        'authorized_by': ''
    }
    try:
        df = pd.read_excel(file_path)
        print(f"Document Control Excel file read successfully. Columns: {df.columns.tolist()}")

        # Clean column names by stripping whitespace
        df.columns = df.columns.str.strip()
        
        # Find the header row containing "Version", "Date", "Reviewed By", "Brief description of the change"
        header_row_index = None
        for index, row in df.iterrows():
            # Check if key headers are present in this row
            if all(col in row.index for col in ["Version", "Date", "Reviewed By", "Brief description of the change"]):
                header_row_index = index
                print(f"Found document control header row at index: {header_row_index}")
                break
        
        if header_row_index is None:
            print("Document control header row not found. Cannot extract document control details or amendment log.")
            return doc_control_data # Return empty data

        # Assuming assessment details are in the row *after* the header row (first data row)
        first_data_row_index = header_row_index + 1
        if first_data_row_index < len(df):
            first_data_row = df.iloc[first_data_row_index]
            
            # Handle Assessment Scope
            if 'Assessment Scope' in first_data_row:
                doc_control_data['assessment_scope'] = str(first_data_row['Assessment Scope']).strip()
                print(f"Found Assessment Scope: {doc_control_data['assessment_scope']}")
            # Handle Assessment Note
            if 'Assessment Note' in first_data_row:
                doc_control_data['assessment_note'] = str(first_data_row['Assessment Note']).strip()
                print(f"Found Assessment Note: {doc_control_data['assessment_note']}")
            # Handle Assessment Start Date
            if 'Assessment Start Date' in first_data_row:
                start_date = first_data_row['Assessment Start Date']
                print(f"Raw Start Date (Doc Control): {start_date}, Type: {type(start_date)}")
                if isinstance(start_date, pd.Timestamp):
                    doc_control_data['assessment_start_date'] = start_date.strftime('%dth %B %Y')
                else:
                    # Try to parse the date string if it's in the format "10th March 2025"
                    try:
                         # Remove 'th', 'st', 'nd', 'rd' from the day
                        date_str = str(start_date).strip()
                        date_str = date_str.replace('th', '').replace('st', '').replace('nd', '').replace('rd', '')
                        # Parse the date
                        parsed_date = pd.to_datetime(date_str, format='%d %B %Y')
                        doc_control_data['assessment_start_date'] = parsed_date.strftime('%dth %B %Y')
                    except:
                        doc_control_data['assessment_start_date'] = str(start_date).strip()
                print(f"Formatted Start Date (Doc Control): {doc_control_data['assessment_start_date']}")

            # Handle Assessment End Date
            if 'Assessment End Date' in first_data_row:
                end_date = first_data_row['Assessment End Date']
                print(f"Raw End Date (Doc Control): {end_date}, Type: {type(end_date)}")
                if isinstance(end_date, pd.Timestamp):
                    doc_control_data['assessment_end_date'] = end_date.strftime('%dth %B %Y')
                else:
                     # Try to parse the date string if it's in the format "30th March 2025"
                    try:
                         # Remove 'th', 'st', 'nd', 'rd' from the day
                        date_str = str(end_date).strip()
                        date_str = date_str.replace('th', '').replace('st', '').replace('nd', '').replace('rd', '')
                        # Parse the date
                        parsed_date = pd.to_datetime(date_str, format='%d %B %Y')
                        doc_control_data['assessment_end_date'] = parsed_date.strftime('%dth %B %Y')
                    except:
                        doc_control_data['assessment_end_date'] = str(end_date).strip()
                print(f"Formatted End Date (Doc Control): {doc_control_data['assessment_end_date']}")
            
            # Extract "Reviewed By" and "Authorized By" from the first data row
            if first_data_row_index < len(df):
                first_data_row = df.iloc[first_data_row_index]
                doc_control_data['reviewed_by'] = str(first_data_row.get('Reviewed By', '')).strip()
                doc_control_data['authorized_by'] = str(first_data_row.get('Authorized By', '')).strip()
                print(f"Found Reviewed By: {doc_control_data['reviewed_by']}")
                print(f"Found Authorized By: {doc_control_data['authorized_by']}")

            print(f"Final assessment details from doc control:")
            print(f"Scope: {doc_control_data['assessment_scope']}")
            print(f"Start Date: {doc_control_data['assessment_start_date']}")
            print(f"End Date: {doc_control_data['assessment_end_date']}")
            print(f"Note: {doc_control_data['assessment_note']}")

        # Extract amendment log entries from rows below the header row
        amendment_log_start_index = header_row_index + 1
        for index in range(amendment_log_start_index, len(df)):
            row = df.iloc[index]
            # Stop if the "Version" cell is empty/NaN
            if pd.isna(row.get('Version', None)) or str(row.get('Version', '')).strip() == '':
                break # Stop reading amendment log
            
            # Handle Date formatting to remove time
            log_date = row.get('Date', '')
            formatted_date = str(log_date).strip()
            if isinstance(log_date, pd.Timestamp):
                formatted_date = log_date.strftime('%Y-%m-%d') # Format as YYYY-MM-DD
            
            log_entry = {
                'Version': str(row.get('Version', '')).strip(),
                'Date': formatted_date,
                'Reviewed By': str(row.get('Reviewed By', '')).strip(),
                'Brief description of the change': str(row.get('Brief description of the change', '')).strip()
            }
            doc_control_data['amendment_log'].append(log_entry)
            print(f"Added amendment log entry: {log_entry}")

        print(f"Total amendment log entries extracted: {len(doc_control_data['amendment_log'])}")
        
        return doc_control_data

    except FileNotFoundError:
        print(f"Error: Document control file not found at {file_path}")
        # Continue without document control data if file is missing
        return doc_control_data
    except Exception as e:
        print(f"Error parsing document control Excel file: {e}")
        # Continue without document control data if parsing fails
        return doc_control_data

def create_summary_table(doc, vulnerabilities, ip=None):
    """
    Adds a summary table for the given IP's vulnerabilities after the IP heading.
    All cell contents are center-aligned and bold. The first column header is 'Sr. No.'
    """
    doc.add_paragraph()
    headers = ['Sr. No.', 'Vulnerability Name', 'Vulnerability Risk Type']
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(1.5)
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        p = cell.paragraphs[0]
        p.text = header_text
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
    for idx, vuln in enumerate(vulnerabilities, 1):
        row_cells = table.add_row().cells
        # Sr. No.
        cell0 = row_cells[0]
        cell0.text = str(idx)
        for para in cell0.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
        # Vulnerability Name
        cell1 = row_cells[1]
        cell1.text = vuln.get('name', 'N/A')
        for para in cell1.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
        # Vulnerability Risk Type (Severity)
        severity = vuln.get('severity', 'Unknown')
        cell2 = row_cells[2]
        cell2.text = severity
        colors = get_severity_colors(severity)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
        cell2._element.get_or_add_tcPr().append(shading_elm)
        for para in cell2.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Always white font for severity
    doc.add_paragraph()

def create_vulnerability_table(doc, vulnerability, display_sr_no=None, image_map=None, ip=None):
    print(f"Creating table for vulnerability: {vulnerability['name']}, Severity: {vulnerability['severity']}, IP: {ip or vulnerability.get('ip')}")
    row_labels = [
        "",  # Row 1: Vulnerability Name
        "IP",  # Row 2: IP (changed from 'IP Address' to 'IP')
        "",  # Row 3: Vulnerable URL
        "Vulnerable Parameter",  # Row 4
        "CVSS Score",  # Row 5
        "Severity",  # Row 6
        "Vulnerability Description",  # Row 7
        "Vulnerability Impact",  # Row 8
        "Recommendation",  # Row 9
        "Proof of Concept / Steps to Reproduce"  # Row 10
    ]
    table = doc.add_table(rows=len(row_labels), cols=1)
    table.style = 'Table Grid'
    severity = vulnerability["severity"]
    colors = get_severity_colors(severity)
    row1_bg_hex = colors["row1_bg"].lstrip('#')
    row2_bg_hex = colors["row2_bg"].lstrip('#')
    font_hex = colors["font"].lstrip('#')
    row1_r, row1_g, row1_b = tuple(int(row1_bg_hex[i:i+2], 16) for i in (0, 2, 4))
    row2_r, row2_g, row2_b = tuple(int(row2_bg_hex[i:i+2], 16) for i in (0, 2, 4))
    font_r, font_g, font_b = tuple(int(font_hex[i:i+2], 16) for i in (0, 2, 4))
    # Row 1: Vulnerability Name
    row = table.rows[0]
    cell = row.cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    sr_display = display_sr_no if display_sr_no else vulnerability.get('sr_no', '')
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(f"{sr_display}: {vulnerability['name']}")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Row 2: IP (color-coded)
    row = table.rows[1]
    cell = row.cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(f"IP: {ip or vulnerability.get('ip', 'N/A')}")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    # Row 3: Vulnerable URL
    row = table.rows[2]
    cell = row.cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row2_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Vulnerable URL: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run = paragraph.add_run(vulnerability['vulnerable_url'] or "N/A")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(255, 255, 255)
    vuln_param = vulnerability.get('vulnerable_parameter', '')
    if isinstance(vuln_param, float) and pd.isna(vuln_param):
        vuln_param = "N/A"
    row_contents = [
        str(vuln_param),
        str(vulnerability['cvss']) if vulnerability['cvss'] is not None else "N/A",
        str(vulnerability['severity']) if vulnerability['severity'] else "N/A",
        str(vulnerability['description']) if vulnerability['description'] else "N/A",
        str(vulnerability.get('impact', "N/A")),
        str(vulnerability.get('remediation', "N/A")),
        ""
    ]
    for i in range(3, len(row_labels)):
        row_idx = i
        row = table.rows[row_idx]
        cell = row.cells[0]
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f"{row_labels[i]}: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(font_r, font_g, font_b)
        if i < len(row_labels) - 1:
            run = paragraph.add_run(row_contents[i-3])
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = False
        else:
            steps = vulnerability.get('steps_with_screenshots', [])
            if not steps:
                run = paragraph.add_run("No steps provided.")
                run.font.name = 'Altone Trial'
                run.font.size = Pt(11)
            else:
                for step_idx, step in enumerate(steps):
                    step_para = cell.add_paragraph()
                    step_para.paragraph_format.left_indent = Pt(10)
                    step_text = step['text']
                    run = step_para.add_run(f"{step_text}")
                    run.font.name = 'Altone Trial'
                    run.font.size = Pt(11)
                    run.font.bold = False
                    screenshot_path = None
                    screenshot_name = step.get('screenshot', '').strip().lower()
                    if image_map and screenshot_name:
                        ip_key = ip or vulnerability.get('ip')
                        sr_no = display_sr_no or vulnerability.get('sr_no')
                        import re
                        match = re.search(r'step(\d+)', screenshot_name)
                        if match:
                            step_num = int(match.group(1))
                        else:
                            step_num = step_idx + 1
                        print(f"[DEBUG] Looking for image: IP={ip_key}, VULN={sr_no}, STEP={step_num}")
                        if image_map and ip_key in image_map and sr_no in image_map[ip_key]:
                            screenshot_path = image_map[ip_key][sr_no].get(step_num)
                            print(f"[DEBUG] Found path: {screenshot_path}")
                    if screenshot_path and os.path.exists(screenshot_path):
                        screenshot_para = cell.add_paragraph()
                        screenshot_para.paragraph_format.left_indent = Pt(20)
                        screenshot_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        img_run = screenshot_para.add_run()
                        img_run.add_picture(screenshot_path, width=Inches(4.5))  # Consistent width for all images
                    else:
                        missing_para = cell.add_paragraph()
                        missing_para.paragraph_format.left_indent = Pt(20)
                        missing_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = missing_para.add_run(f"[Screenshot missing: {step.get('screenshot')}]")
                        run.font.name = 'Altone Trial'
                        run.font.size = Pt(10)
                        run.font.italic = True
                    if step_idx < len(steps) - 1:
                        spacer = cell.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(6)
    return table

def create_vulnerability_table_with_poc(doc, vulnerability, display_sr_no=None, image_map=None):
    # Create the main vulnerability table as before
    table = create_vulnerability_table(doc, vulnerability, display_sr_no, image_map)
    # After the table, add PoC steps and images
    para = doc.add_paragraph(f"Proof of Concept / Steps to Reproduce for {vulnerability['name']}:")
    para.runs[0].bold = True
    para.runs[0].font.size = Pt(14)
    steps = vulnerability.get('steps_with_screenshots', [])
    for step_idx, step in enumerate(steps, 1):
        step_para = doc.add_paragraph(f"Step {step_idx}: {step['text']}")
        step_para.paragraph_format.left_indent = Pt(10)
        screenshot_path = find_image_for_step(step['screenshot'], image_map, display_sr_no, step_idx)
        if screenshot_path and os.path.exists(screenshot_path):
            print(f"[DEBUG] Adding image from ZIP: {screenshot_path}")
            img_run = doc.add_paragraph().add_run()
            img_run.add_picture(screenshot_path, width=Inches(4))
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]")
    return table

def calculate_optimal_logo_size(image_path, max_width=2.3, max_height=2.0):
    """
    Calculate the optimal logo size to fit within the template while maintaining aspect ratio.
    
    Args:
        image_path (str): Path to the logo image
        max_width (float): Maximum width in inches (increase this value if logos appear too small)
        max_height (float): Maximum height in inches (increase this value for taller logos)
        
    Returns:
        float: Optimal width in inches
        
    Note:
        - The default parameters (2.5 x 1.2 inches) are suitable for most company logos
        - For wider logos, increase max_width
        - For taller logos, increase max_height
        - The minimum width is set to 1.5 inches to ensure logos are clearly visible
    """
    # Check if PIL is available
    if not PIL_AVAILABLE:
        print("PIL not available, using default logo width")
        return max_width * 0.9  # Use a slightly smaller width as a safe default
    
    try:
        # Open the image to get its dimensions
        with Image.open(image_path) as img:
            img_width, img_height = img.size
            
        # Convert pixels to inches (assuming 96 DPI which is standard for many screens)
        # This is a rough estimate, as exact DPI may vary based on the image
        width_in_inches = img_width / 96
        height_in_inches = img_height / 96
            
        print(f"Original logo dimensions: {img_width}x{img_height}px ({width_in_inches:.2f}x{height_in_inches:.2f} inches)")
        
        # If the image is already within our size limits, use its actual size,
        # but ensure it's not too small (at least 1.5 inches wide for good visibility)
        if width_in_inches <= max_width and height_in_inches <= max_height:
            if width_in_inches < 1.5:
                print(f"Logo is too small, scaling up to minimum width: 1.5 inches")
                return 1.5
            print(f"Logo is already within size limits, using original size: {width_in_inches:.2f} inches wide")
            return width_in_inches
            
        # Calculate aspect ratio
        aspect_ratio = img_width / img_height
        
        # Calculate potential width and height based on max dimensions
        width_from_height = max_height * aspect_ratio
        height_from_width = max_width / aspect_ratio
        
        # Determine which dimension is the limiting factor
        if width_from_height <= max_width:
            # Height is the limiting factor
            optimal_width = width_from_height
        else:
            # Width is the limiting factor
            optimal_width = max_width
            
        # Ensure minimum size of 1.5 inches width
        if optimal_width < 1.5:
            optimal_width = 1.5
            
        print(f"Resizing logo to: {optimal_width:.2f} inches wide")
        print(f"Logo display size will be approximately {optimal_width:.2f} x {optimal_width/aspect_ratio:.2f} inches")
        return optimal_width
    except Exception as e:
        print(f"Error calculating optimal logo size: {e}")
        return max_width  # Default to max_width if there's an error

# Updated function signature to accept amendment_log and assessment details
def generate_word_report(vulnerabilities_by_ip, doc_control_data, template_path, output_path, client_name="Client", image_map=None, logo_path=None):
    try:
        tpl = DocxTemplate(template_path)
        current_year = datetime.now().year
        current_quarter = (datetime.now().month - 1) // 3 + 1
        client_code = client_name.upper()[:2]
        doc_id = f"C-{client_code}-IP-Q{current_quarter}-{current_year}"
        assessment_scope = doc_control_data.get('assessment_scope', '')
        assessment_start_date = doc_control_data.get('assessment_start_date', '')
        assessment_end_date = doc_control_data.get('assessment_end_date', '')
        assessment_note = doc_control_data.get('assessment_note', '')
        amendment_log = doc_control_data.get('amendment_log', [])
        date_range = f"{assessment_start_date} to {assessment_end_date}"
        company_logo = None
        default_logo_path = os.path.join(os.path.dirname(__file__), "default_logo.png")
        try:
            if logo_path and os.path.exists(logo_path):
                optimal_width = calculate_optimal_logo_size(logo_path)
                company_logo = InlineImage(tpl, logo_path, width=Inches(optimal_width))
            elif os.path.exists(default_logo_path):
                optimal_width = calculate_optimal_logo_size(default_logo_path)
                company_logo = InlineImage(tpl, default_logo_path, width=Inches(optimal_width))
        except Exception as e:
            print(f"Error processing logo: {e}. Continuing without logo.")
            company_logo = None
        context = {
            'client_name': client_name,
            'assessment_scope': assessment_scope,
            'date_range': date_range,
            'document_id': doc_id,
            'amendment_log': amendment_log,
            'reviewed_by': doc_control_data.get('reviewed_by', ''),
            'authorized_by': doc_control_data.get('authorized_by', ''),
            'report_date': datetime.now().strftime('%d %B %Y'),
            'assessment_note': assessment_note,
            'company_logo': company_logo,
            'Risk_Assessment_Analysis': '_RISK_ASSESSMENT_ANALYSIS_PLACEHOLDER_',
            'Vulnerabilities_Found': '_VULNERABILITIES_FOUND_PLACEHOLDER_',
        }
        tpl.render(context)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tpl.save(tmp.name)
            tmp_path = tmp.name
        doc = Document(tmp_path)
        # Insert overall vulnerabilities charts at placeholder
        insert_overall_vulnerabilities_charts(doc, vulnerabilities_by_ip)
        # Insert risk assessment analysis table at placeholder
        insert_risk_assessment_analysis_table(doc, vulnerabilities_by_ip)
        for ip, vulns in vulnerabilities_by_ip.items():
            doc.add_page_break()
            heading = doc.add_paragraph(f"IP {ip}")
            heading.runs[0].bold = True
            heading.runs[0].font.size = Pt(18)
            heading.runs[0].font.name = 'Altone Trial'
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Assign display_sr_no for this IP
            for idx, vuln in enumerate(vulns, 1):
                vuln['display_sr_no'] = f"VUL-{idx:03d}"
            # If this IP is a 'no vulnerability' IP, skip summary table and add special table
            if len(vulns) == 1 and vulns[0].get('no_vuln'):
                create_no_vulnerability_table(doc, ip, vulns[0].get('steps_with_screenshots', []))
                continue
            # Insert summary table for this IP
            create_summary_table(doc, vulns, ip=ip)
            # Insert per-IP bar chart
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as ip_chart_file:
                ip_chart_path = create_ip_severity_bar_chart(ip, vulns, ip_chart_file.name)
                img_paragraph = doc.add_paragraph()
                run = img_paragraph.add_run()
                run.add_picture(ip_chart_path, width=Inches(4.5))
                img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Insert vulnerability tables for this IP
            for vuln in vulns:
                doc.add_page_break()  # Always start each vulnerability table on a new page
                create_vulnerability_table(doc, vuln, vuln['display_sr_no'], image_map, ip=ip)
        doc.add_page_break()
        doc.add_paragraph()
        doc.add_paragraph()
        summary_heading = doc.add_paragraph()
        run = summary_heading.add_run("SUMMARY OF FINDINGS &\nCONCLUSION:")
        doc.add_paragraph()
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = 'Altone Trial'
        run.font.color.rgb = RGBColor(106, 68, 154)
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        try:
            summary_heading.style = doc.styles['Heading 1']
        except Exception:
            pass
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_after = Pt(24)
        run = p1.add_run(
            "Finally, it must be remembered that security is an ongoing process, and that this report will provide an idea of the current vulnerabilities we were able to detect. There is no guarantee that new vulnerabilities will not be found and exploited in the future."
        )
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.space_after = Pt(24)
        run = p2.add_run("The assessment was only possible because ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("joint support & coordination")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" from the information security team of organization for ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("sharing & coordinating")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" during the assessment period. It is advised to refer the ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("Technical Report")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" for understanding in-depth of vulnerabilities that were discovered by technical team of ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("CyberSmithSECURE Pvt. Ltd.")
        run.bold = True
        run.underline = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p3.paragraph_format.space_after = Pt(24)
        run = p3.add_run("The Security Researchers of the ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p3.add_run("CyberSmithSECURE")
        run.bold = True
        run.underline = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p3.add_run(" performed Vulnerability Testing. We jointly recommend that all suggested measures in this document be performed to ensure the overall security of the target device. The following targeted sectors were identified by the security researchers for the scope of this testing.")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p4.paragraph_format.space_after = Pt(24)
        run = p4.add_run("We thank internal Information Security team for their support & cooperation during the time of assessment.")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(106, 68, 154)
        doc.save(output_path)
        return output_path
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating Word report: {str(e)}")

def get_script_dir():
    return os.path.dirname(os.path.realpath(__file__))

# Function to check if file exists in uploads directory
def check_screenshots_directory(vulnerability, screenshots_dir="uploads/screenshots"):
    """
    Check if screenshot directory exists for a vulnerability, create it if not
    Returns the path to the screenshots directory
    """
    try:
        # Create directory structure based on severity and vulnerability name
        severity = vulnerability.get('severity', 'Unknown')
        vuln_name = vulnerability.get('name', 'Unknown').replace(' ', '_').replace('/', '_')
        
        # Create path: uploads/screenshots/[Severity]/[Vulnerability_Name]
        screenshot_path = os.path.join(screenshots_dir, severity, vuln_name)
        os.makedirs(screenshot_path, exist_ok=True)
        
        return screenshot_path
    except Exception as e:
        print(f"Error creating screenshots directory: {e}")
        return None

def find_image_case_insensitive(directory, filename):
    """
    Search for an image file in the directory, case-insensitive, and try common extensions.
    """
    base, ext = os.path.splitext(filename)
    possible_exts = ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']
    candidates = [base + e for e in possible_exts]
    files = os.listdir(directory)
    for candidate in candidates:
        for f in files:
            if f.lower() == candidate.lower():
                return os.path.join(directory, f)
    return None

def index_images_recursively(root_dir):
    image_map = {}
    for dirpath, _, filenames in os.walk(root_dir):
        for fname in filenames:
            if fname.lower().endswith((".png", ".jpg", ".jpeg")):
                norm = normalize_filename(fname)
                full_path = os.path.join(dirpath, fname)
                with open(full_path, 'rb') as f:
                    data = f.read()
                    h = hashlib.md5(data).hexdigest()
                    print(f"[IMAGE MAP] {fname} => {norm} | Size: {len(data)} | Hash: {h}")
                    f.seek(0)
                image_map[norm] = full_path
    print(f"[DEBUG] Final image map: {list(image_map.keys())}")
    return image_map

def insert_steps_with_images(doc, vulnerability, image_map):
    doc.add_paragraph(f"Proof of Concept / Steps to Reproduce for {vulnerability['name']}").runs[0].bold = True
    for idx, step in enumerate(vulnerability.get("steps_with_screenshots", []), 1):
        doc.add_paragraph(f"Step {idx}: {step['text']}")
        norm = normalize_filename(step['screenshot'])
        img_path = image_map.get(norm)
        print(f"[PoC] Step {idx} for '{vulnerability['name']}' → Looking for '{norm}' → {img_path}")
        if img_path and os.path.exists(img_path):
            with open(img_path, 'rb') as f:
                img_data = f.read()
                print(f"[PoC] FOUND: {img_path} | Hash: {hashlib.md5(img_data).hexdigest()} | Size: {len(img_data)}")
                f.seek(0)
                doc.add_paragraph().add_run().add_picture(f, width=Inches(4))
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]")

@app.post("/generate-report/")
async def generate_report(
    file: UploadFile = File(...),  # Vulnerability file
    assessment_file: UploadFile = File(...),  # Assessment details file
    poc_images: UploadFile = File(None),  # POC screenshots ZIP file (optional)
    company_logo: UploadFile = File(None),  # Company logo file (optional)
    client_name: str = Form("Client"),
    request: Request = None,
    db: Session = Depends(get_db)
):
    try:
        print(f"Starting report generation - Version {VERSION}")
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)
        screenshots_dir = os.path.join(UPLOAD_DIR, "screenshots")
        if not os.path.exists(screenshots_dir):
            os.makedirs(screenshots_dir)
        logos_dir = os.path.join(UPLOAD_DIR, "logos")
        if not os.path.exists(logos_dir):
            os.makedirs(logos_dir)
        vulnerability_file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(vulnerability_file_location, "wb") as f:
            f.write(await file.read())
        print(f"Saved uploaded vulnerability file to: {vulnerability_file_location}")
        assessment_file_location = os.path.join(UPLOAD_DIR, assessment_file.filename)
        with open(assessment_file_location, "wb") as f:
            f.write(await assessment_file.read())
        print(f"Saved uploaded assessment file to: {assessment_file_location}")
        logo_path = None
        if company_logo is not None:
            logo_filename = f"{client_name.replace(' ', '_')}_logo{os.path.splitext(company_logo.filename)[-1]}"
            logo_path = os.path.join(logos_dir, logo_filename)
            with open(logo_path, "wb") as f:
                f.write(await company_logo.read())
            print(f"Saved company logo to: {logo_path}")
        poc_images_dir = None
        if poc_images is not None:
            tmpdirname = tempfile.mkdtemp()
            zip_path = os.path.join(tmpdirname, poc_images.filename)
            with open(zip_path, 'wb') as f:
                f.write(await poc_images.read())
            with ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(tmpdirname)
            print('--- Extracted ZIP Directory Tree ---')
            for dirpath, dirnames, filenames in os.walk(tmpdirname):
                print(f"DIR: {dirpath}")
                print(f"  DIRNAMES: {dirnames}")
                print(f"  FILENAMES: {filenames}")
            print('--- End of Directory Tree ---')
            poc_images_dir = tmpdirname
        vulnerabilities_by_ip = parse_vulnerabilities_excel(vulnerability_file_location)
        doc_control_data = parse_doc_control_excel(assessment_file_location)
        # Save optional dashboard file for later dashboard usage
        # No dashboard ingestion here; dashboard is a separate feature
        
        # Make template path absolute
        template_path = os.path.join(get_script_dir(), "CSS-NEW_Technical_Sample_Report_Automated_v2.1-IP.docx")

        report_filename = f"{client_name} VAPT Report {datetime.now().strftime('%Y-%d-%m')}.docx"
        output_path = os.path.join(UPLOAD_DIR, report_filename)
        image_map = None
        if poc_images_dir:
            image_map = index_images_from_poc_zip(poc_images_dir)
            if image_map:
                print('Available images in ZIP:')
                for ip, vulns in image_map.items():
                    print(f'IP: {ip} -> {list(vulns.keys())}')
        generate_word_report(vulnerabilities_by_ip, doc_control_data, template_path, output_path, client_name, image_map, logo_path)
        print(f"Report generation completed successfully - Version {VERSION}")
        # Audit log
        try:
            user = request.session.get('user') if request else None
            db.add(AuditLog(
                user_email=(user or {}).get('email'),
                user_name=(user or {}).get('name'),
                action='generate-report-type1',
                metadata_json=json.dumps({'vuln_file': file.filename, 'assessment_file': assessment_file.filename}),
                ip_address=request.client.host if request and request.client else None,
                user_agent=request.headers.get('user-agent') if request else None
            ))
            db.commit()
        except Exception:
            pass
        return FileResponse(
            path=output_path,
            filename=report_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

def count_vulnerabilities_by_severity(vulnerabilities):
    """
    Counts the number of vulnerabilities for each severity level.
    """
    severity_counts = Counter()
    for vuln in vulnerabilities:
        severity_counts[vuln.get('severity', 'Unknown')] += 1
        
    return severity_counts

def create_severity_bar_chart(severity_counts, output_path):
    """
    Generates a bar chart showing the count of vulnerabilities by severity.
    Saves the chart as an image file.
    """
    severities = list(severity_counts.keys())
    counts = list(severity_counts.values())
    
    # Define colors matching the order (approximate hex values)
    colors = [
        '#990000' if s == 'Critical' else
        '#FF0000' if s == 'High' else
        '#FFCC00' if s == 'Medium' else
        '#00b050' if s == 'Low' else
        '#0070c0' # Informational/Unknown
        for s in severities
    ]
    
    # Sort data by severity order for consistency display
    sorted_severities, sorted_counts, sorted_colors = zip(*sorted(zip(severities, counts, colors), key=lambda x: severity_order[x[0]]))
    
    x_pos = np.arange(len(sorted_severities))
    
    plt.figure(figsize=(8, 6))
    bars = plt.bar(x_pos, sorted_counts, color=sorted_colors)
    plt.xticks(x_pos, sorted_severities, rotation=10, ha='right', fontsize=14)
    plt.ylabel('Number of Vulnerabilities', fontsize=14)
    plt.title('Total Vulnerabilities Found', fontsize=18)
    
    # Add counts inside the bars with larger font size
    for i, bar in enumerate(bars):
        height = bar.get_height()
        plt.text(bar.get_x() + bar.get_width()/2, height/2, str(sorted_counts[i]),
                 ha='center', va='center', fontsize=16, weight='bold', color='white')
    
    plt.tight_layout() # Adjust layout to prevent labels overlapping
    plt.savefig(output_path)
    plt.close()
    
    return output_path

def create_severity_donut_chart(severity_counts, output_path):
    """
    Generates a donut chart showing the distribution of vulnerabilities by severity.
    Saves the chart as an image file with improved styling to match a typical cybersecurity report.
    """
    severities = list(severity_counts.keys())
    counts = list(severity_counts.values())
    total_findings = sum(counts)
    
    # Use colors from get_severity_colors for consistency
    color_map = {severity: get_severity_colors(severity)['row1_bg'] for severity in severity_order.keys()}
    # Default color for 'Unknown' if present
    color_map['Unknown'] = '#808080'  # Grey for unknown
    
    # Sort data by severity order
    sorted_severities, sorted_counts = zip(*sorted(severity_counts.items(), key=lambda item: severity_order.get(item[0], 5)))
    sorted_colors = [color_map.get(s) for s in sorted_severities]
    
    fig, ax = plt.subplots(figsize=(8, 6))  # Adjusted figure size for better proportions
    
    wedges, texts = ax.pie(
        sorted_counts,
        colors=sorted_colors, 
        wedgeprops=dict(width=0.4, edgecolor='white'),  # Thinner donut with white edges for separation
        startangle=90,  # Start at the top
        labels=None,  # We'll add custom labels
        textprops={'fontsize': 14, 'weight': 'bold', 'color': 'white'}  # Larger font for wedge labels
    )

    # Add count labels manually on each wedge with larger font size
    import math
    for i, w in enumerate(wedges):
        ang = (w.theta2 + w.theta1)/2.
        x = math.cos(math.radians(ang)) * 0.8
        y = math.sin(math.radians(ang)) * 0.8
        ax.text(x, y, str(sorted_counts[i]), ha='center', va='center', fontsize=16, weight='bold', color='white')

    # Draw a circle in the center to make it a donut
    centre_circle = plt.Circle((0, 0), 0.60, fc='white')  # Adjusted center circle size
    fig.gca().add_artist(centre_circle)
    
    # Add total findings text in the center
    ax.text(
        0, 0,
        f'Total\nFindings\n{total_findings}',
        ha='center',
        va='center',
        fontsize=14,
        weight='bold',
        color='black'
    )
    
    # Add legend with counts, positioned below the chart to avoid overlap, with larger font size
    legend_labels = [f'{s} ({c})' for s, c in zip(sorted_severities, sorted_counts)]
    ax.legend( # type: ignore
        wedges,
        legend_labels,
        title="Severity Levels",
        loc="center",
        bbox_to_anchor=(0.5, -0.1),  # Place legend below the chart
        ncol=len(sorted_severities),  # Display in one row
        fontsize=14,
        title_fontsize=16,
        frameon=False  # Remove the legend frame for a cleaner look
    )
    
    # Set the title with consistent styling
    plt.title(
        'Overall Vulnerabilities Identified',
        fontsize=18,
        weight='bold',
        pad=20,
        color='#660099'  # Match the heading color used in the Word document
    )
    
    ax.axis('equal')  # Equal aspect ratio to ensure that pie is drawn as a circle
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches='tight')  # Ensure the chart legend is not cut off
    plt.close()
    
    return output_path 

def index_images_from_poc_zip(root_dir):
    import re
    image_map = {}
    # If root_dir contains a single folder, descend into it
    entries = [e for e in os.listdir(root_dir) if os.path.isdir(os.path.join(root_dir, e))]
    if len(entries) == 1:
        root_dir = os.path.join(root_dir, entries[0])
    # Each top-level folder is an IP
    for ip_folder in os.listdir(root_dir):
        ip_path = os.path.join(root_dir, ip_folder)
        if not os.path.isdir(ip_path):
            continue
        ip_key = ip_folder.strip()
        image_map[ip_key] = {}
        # Inside each IP folder, walk through severity folders (optional) and vuln folders
        for severity in os.listdir(ip_path):
            severity_path = os.path.join(ip_path, severity)
            if not os.path.isdir(severity_path):
                continue
            for vuln_id in os.listdir(severity_path):
                vuln_path = os.path.join(severity_path, vuln_id)
                if os.path.isdir(vuln_path) and vuln_id.lower().startswith("vul-"):
                    if vuln_id not in image_map[ip_key]:
                        image_map[ip_key][vuln_id] = {}
                    for fname in os.listdir(vuln_path):
                        if fname.lower().endswith((".png", ".jpg", ".jpeg")):
                            match = re.search(r'step(\d+)', fname.lower())
                            if match:
                                step_num = int(match.group(1))
                                image_map[ip_key][vuln_id][step_num] = os.path.join(vuln_path, fname)
    print('--- Final Image Map ---')
    for ip, vulns in image_map.items():
        print(f'IP: {ip}')
        for vuln_id, steps in vulns.items():
            print(f'  {vuln_id}: {list(steps.keys())}')
    print('--- End Image Map ---')
    return image_map

def build_steps_with_images(steps_with_screenshots, image_map, tpl, sr_no):
    steps = []
    sr_no_key = sr_no.strip().lower() if sr_no else ''
    for idx, step in enumerate(steps_with_screenshots, 1):
        screenshot_name = step.get('screenshot', '')
        match = re.search(r'step(\\d+)', screenshot_name.strip().lower())
        if match:
            step_num = int(match.group(1))
        else:
            step_num = idx
        key = f"{sr_no_key}_step{step_num}"
        img_path = image_map.get(key) if image_map else None
        print(f"[POC IMAGE DEBUG] Step {idx}, Key: {key}, Path: {img_path}")
        if img_path and os.path.exists(img_path):
            print(f"[POC IMAGE FOUND] Using image: {img_path}")
            image = InlineImage(tpl, img_path, width=Inches(5))
        else:
            print(f"[POC IMAGE MISSING] {img_path}")
            image = f"[Screenshot missing: {screenshot_name}]"
        steps.append({'text': step['text'], 'image': image})
    return steps

def find_image_for_step(screenshot_name, image_map, sr_no, step_idx):
    screenshot_name_key = screenshot_name.strip().lower()
    sr_no_key = sr_no.strip().lower() if sr_no else ''
    match = re.search(r'step(\d+)', screenshot_name_key)
    if match:
        step_num = int(match.group(1))
    else:
        step_num = step_idx
    key = f"{sr_no_key}_step{step_num}"
    print(f"Looking for key: {key} in image_map")
    if key in image_map:
        print(f"Found image for {key}: {image_map[key]}")
    else:
        print(f"Image for {key} NOT FOUND")
    return image_map.get(key)

# After parsing vulnerabilities, ensure each has a display_sr_no
def assign_display_sr_no(vulnerabilities):
    for idx, vuln in enumerate(vulnerabilities, 1):
        if not vuln.get('display_sr_no'):
            sr_no = str(vuln.get('sr_no', f"VUL-{idx:03d}")).replace('VULN-', 'VUL-').strip()
            vuln['display_sr_no'] = sr_no
    return vulnerabilities 

def insert_poc_steps_section(doc, vulnerability, image_map):
    # Always add a new section for PoC steps
    doc.add_paragraph("Proof of Concept / Steps to Reproduce:").runs[0].bold = True
    steps = vulnerability.get('steps_with_screenshots', [])
    sr_no = vulnerability.get('display_sr_no') or vulnerability.get('sr_no')
    for idx, step in enumerate(steps, 1):
        step_para = doc.add_paragraph(f"Step {idx}: {step['text']}")
        step_para.paragraph_format.left_indent = Pt(10)
        screenshot_path = find_image_for_step(step['screenshot'], image_map, sr_no, idx)
        if screenshot_path and os.path.exists(screenshot_path):
            print(f"[DEBUG] Adding image from ZIP: {screenshot_path}")
            img_run = doc.add_paragraph().add_run()
            img_run.add_picture(screenshot_path, width=Inches(5))
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]") 

def insert_risk_assessment_analysis_table(doc, vulnerabilities_by_ip):
    """
    Inserts a table at the _RISK_ASSESSMENT_ANALYSIS_PLACEHOLDER_ marker summarizing vulnerabilities per IP and severity.
    """
    placeholder = '_RISK_ASSESSMENT_ANALYSIS_PLACEHOLDER_'
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            placeholder_paragraph = p
            break
    if not placeholder_paragraph:
        print(f"Warning: Placeholder {placeholder} not found in the template.")
        return

    # Table columns
    columns = [
        'Sr. No.', 'Hostname', 'Instant Purpose', 'VAPT Status',
        'Critical', 'High', 'Medium', 'Low', 'Informational', 'Total'
    ]
    severity_keys = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    # Prepare data rows
    data_rows = []
    all_severity_totals = {k: 0 for k in severity_keys}
    grand_total = 0
    for idx, (ip, vulns) in enumerate(vulnerabilities_by_ip.items(), 1):
        row = [str(idx), ip, 'Public IP', 'Completed']
        sev_counts = {k: 0 for k in severity_keys}
        if len(vulns) == 1 and vulns[0].get('no_vuln'):
            # All counts remain 0
            pass
        else:
            for v in vulns:
                sev = v.get('severity', 'Informational')
                if sev in sev_counts:
                    sev_counts[sev] += 1
        total = sum(sev_counts.values())
        for k in severity_keys:
            row.append(str(sev_counts[k]))
            all_severity_totals[k] += sev_counts[k]
        row.append(str(total))
        grand_total += total
        data_rows.append(row)
    # Add the total row
    total_row = [''] * 4 + [str(all_severity_totals[k]) for k in severity_keys] + [str(grand_total)]
    # Insert the table
    table = doc.add_table(rows=1 + len(data_rows) + 1, cols=len(columns))
    table.style = 'Table Grid'
    # Set column widths (adjusted for better single-line fit)
    widths = [Inches(0.8), Inches(2.2), Inches(1.7), Inches(1.7), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.2), Inches(1.0)]
    for i, w in enumerate(widths):
        table.columns[i].width = w
    # Header row
    header_row = table.rows[0]
    tr = header_row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(1.2 * 1440)))  # 1.2 inches in twips
    trHeight.set(qn('w:hRule'), 'exact')
    trPr.append(trHeight)
    for i, col in enumerate(columns):
        cell = header_row.cells[i]
        p = cell.paragraphs[0]
        p.text = col
        run = p.runs[0]
        run.font.bold = True
        # Center align header text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = 1  # Center vertical alignment
        # Color code headers for first four columns
        if col in ['Sr. No.', 'Hostname', 'Instant Purpose', 'VAPT Status']:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="8ea9db"/>')
            cell._element.get_or_add_tcPr().append(shading_elm)
        # Color code severity columns
        elif col in severity_keys:
            colors = get_severity_colors(col)
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
            cell._element.get_or_add_tcPr().append(shading_elm)
            run.font.color.rgb = RGBColor(255, 255, 255)
            # Set text direction to vertical (top-to-bottom, left-to-right)
            tcPr = cell._element.get_or_add_tcPr()
            text_direction = parse_xml(r'<w:textDirection w:val="btLr" %s/>' % nsdecls('w'))
            tcPr.append(text_direction)
        elif col == 'Total':
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
            cell._element.get_or_add_tcPr().append(shading_elm)
            run.font.color.rgb = RGBColor(255, 255, 255)
        else:
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
            cell._element.get_or_add_tcPr().append(shading_elm)
    # Data rows
    for row_idx, row_data in enumerate(data_rows, 1):
        cells = table.rows[row_idx].cells
        for i, val in enumerate(row_data):
            cells[i].text = val
            # Center align data
            for para in cells[i].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[i].vertical_alignment = 1  # Center vertical alignment
            # Color code data for first four columns
            if columns[i] in ['Sr. No.', 'Hostname', 'Instant Purpose', 'VAPT Status']:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="bdd7ee"/>')
                cells[i]._element.get_or_add_tcPr().append(shading_elm)
            # Color code severity columns
            elif columns[i] in severity_keys:
                colors = get_severity_colors(columns[i])
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
                cells[i]._element.get_or_add_tcPr().append(shading_elm)
                # White font for dark backgrounds
                p = cells[i].paragraphs[0]
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            elif columns[i] == 'Total':
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
                cells[i]._element.get_or_add_tcPr().append(shading_elm)
                p = cells[i].paragraphs[0]
                if p.runs:
                    p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            else:
                shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
                cells[i]._element.get_or_add_tcPr().append(shading_elm)
    # Total row
    total_cells = table.rows[len(data_rows)+1].cells
    # Merge the first four cells for 'Overall Findings'
    merged_cell = total_cells[0].merge(total_cells[1]).merge(total_cells[2]).merge(total_cells[3])
    merged_cell.text = 'Overall Findings'
    p = merged_cell.paragraphs[0]
    p.runs[0].font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = 1
    # Apply background color #bdd7ee to the merged cell
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="bdd7ee"/>')
    merged_cell._element.get_or_add_tcPr().append(shading_elm)
    # Fill the rest of the columns with totals
    for i, k in enumerate(severity_keys, start=4):
        total_cells[i].text = str(all_severity_totals[k])
        colors = get_severity_colors(k)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
        total_cells[i]._element.get_or_add_tcPr().append(shading_elm)
        p = total_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_cells[i].vertical_alignment = 1
        if p.runs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    # Grand total
    total_cells[9].text = str(grand_total)
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
    total_cells[9]._element.get_or_add_tcPr().append(shading_elm)
    p = total_cells[9].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total_cells[9].vertical_alignment = 1
    if p.runs:
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
    # Insert table at placeholder
    parent = placeholder_paragraph._p.getparent()
    parent.insert(parent.index(placeholder_paragraph._p), table._element)
    parent.remove(placeholder_paragraph._p)

def create_combined_vulnerability_charts(severity_counts, output_path):
    """
    Generates a single image with the bar chart and donut chart side by side,
    and a centered legend at the bottom showing all severity levels.
    """
    import math
    import matplotlib.pyplot as plt
    from matplotlib.patches import Patch
    # Ensure all severity levels are present
    all_severities = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    all_colors = [get_severity_colors(s)['row1_bg'] for s in all_severities]
    counts = [severity_counts.get(s, 0) for s in all_severities]
    total_findings = sum(counts)

    fig = plt.figure(figsize=(12, 6))
    gs = fig.add_gridspec(2, 2, height_ratios=[5, 1])
    # Bar chart (left)
    ax_bar = fig.add_subplot(gs[0, 0])
    x_pos = range(len(all_severities))
    bars = ax_bar.bar(x_pos, counts, color=all_colors)
    ax_bar.set_xticks(x_pos)
    ax_bar.set_xticklabels(all_severities, rotation=10, ha='right', fontsize=14)
    ax_bar.set_ylabel('Number of Vulnerabilities', fontsize=14)
    ax_bar.set_title('Total Vulnerabilities Found', fontsize=18)
    # Add counts inside bars
    for i, bar in enumerate(bars):
        height = bar.get_height()
        ax_bar.text(bar.get_x() + bar.get_width()/2, height/2, str(counts[i]),
                    ha='center', va='center', fontsize=16, weight='bold', color='white')
    # Donut chart (right)
    ax_donut = fig.add_subplot(gs[0, 1])
    wedges, _ = ax_donut.pie(
        counts,
        colors=all_colors,
        wedgeprops=dict(width=0.4, edgecolor='white'),
        startangle=90,
        labels=None,
        textprops={'fontsize': 14, 'weight': 'bold', 'color': 'white'}
    )
    for i, w in enumerate(wedges):
        ang = (w.theta2 + w.theta1)/2.
        x = math.cos(math.radians(ang)) * 0.8
        y = math.sin(math.radians(ang)) * 0.8
        ax_donut.text(x, y, str(counts[i]), ha='center', va='center', fontsize=16, weight='bold', color='white')
    centre_circle = plt.Circle((0, 0), 0.60, fc='white')
    ax_donut.add_artist(centre_circle)
    ax_donut.text(0, 0, f'Total\nFindings\n{total_findings}', ha='center', va='center', fontsize=14, weight='bold', color='black')
    ax_donut.set_title('Overall Vulnerabilities Identified', fontsize=18, weight='bold', pad=20, color='#660099')
    ax_donut.axis('equal')
    # Hide axes
    ax_bar.spines['top'].set_visible(False)
    ax_bar.spines['right'].set_visible(False)
    ax_bar.spines['left'].set_visible(False)
    ax_bar.spines['bottom'].set_visible(False)
    ax_bar.tick_params(left=False, bottom=False)
    ax_bar.yaxis.set_ticks_position('none')
    ax_bar.xaxis.set_ticks_position('none')
    # Legend (bottom, centered across both charts)
    legend_ax = fig.add_subplot(gs[1, :])
    legend_ax.axis('off')
    legend_handles = [Patch(facecolor=all_colors[i], label=all_severities[i]) for i in range(len(all_severities))]
    legend = legend_ax.legend(
        handles=legend_handles,
        loc='center',
        ncol=len(all_severities),
        fontsize=16,
        frameon=False,
        bbox_to_anchor=(0.5, 0.5)
    )
    # Remove axes for legend
    legend_ax.set_xticks([])
    legend_ax.set_yticks([])
    plt.tight_layout(rect=[0, 0.08, 1, 1])
    plt.savefig(output_path, bbox_inches='tight')
    plt.close()
    return output_path

def insert_overall_vulnerabilities_charts(doc, vulnerabilities_by_ip):
    """
    Inserts a single combined chart image at the _VULNERABILITIES_FOUND_PLACEHOLDER_ marker.
    """
    import tempfile
    placeholder = '_VULNERABILITIES_FOUND_PLACEHOLDER_'
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if placeholder in p.text:
            placeholder_paragraph = p
            break
    if not placeholder_paragraph:
        print(f"Warning: Placeholder {placeholder} not found in the template.")
        return
    # Gather all vulnerabilities into a flat list, EXCLUDING 'no_vuln' entries
    all_vulns = []
    for vulns in vulnerabilities_by_ip.values():
        for v in vulns:
            if not v.get('no_vuln'):
                all_vulns.append(v)
    from collections import Counter
    severity_counts = Counter()
    for vuln in all_vulns:
        severity_counts[vuln.get('severity', 'Unknown')] += 1
    # Generate the combined chart as a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as chart_file:
        chart_path = create_combined_vulnerability_charts(severity_counts, chart_file.name)
        # Insert heading (styled similar to the image)
        heading = placeholder_paragraph.insert_paragraph_before("Overall Vulnerability Identified")
        heading.style = placeholder_paragraph.style
        for run in heading.runs:
            run.font.size = Pt(32)
            run.font.bold = True
            run.font.color.rgb = RGBColor(106, 68, 154)  # #6A449A
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Insert the combined chart image
        img_paragraph = placeholder_paragraph.insert_paragraph_before()
        run = img_paragraph.add_run()
        run.add_picture(chart_path, width=Inches(6.5))
        img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Remove the placeholder
        parent = placeholder_paragraph._p.getparent()
        parent.remove(placeholder_paragraph._p)

# Add a new API endpoint to save the default logo
@app.post("/upload-default-logo/")
async def upload_default_logo(logo_data: dict):
    """
    Upload a default logo image to be used when the user doesn't provide one.
    Expects a JSON payload with a base64 encoded image.
    """
    try:
        if 'base64_image' not in logo_data:
            raise HTTPException(status_code=400, detail="No image data provided")
        
        # Extract the base64 string and strip any header data
        base64_data = logo_data['base64_image']
        if ',' in base64_data:
            base64_data = base64_data.split(',', 1)[1]
        
        # Decode the base64 string to binary
        image_data = base64.b64decode(base64_data)
        
        # Save to the backend directory
        logo_path = os.path.join(os.path.dirname(__file__), "default_logo.png")
        with open(logo_path, "wb") as f:
            f.write(image_data)
        
        return JSONResponse(content={"status": "success", "message": "Default logo saved"})
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e)) 

def is_no_vulnerability_row(row):
    sr_no = str(row.get('Sr No', '')).strip().lower()
    return sr_no == 'no vulnerability' or sr_no == 'no vulnerabilities'

def create_no_vulnerability_table(doc, ip, steps_with_screenshots=None):
    """
    Create a special vulnerability table for IPs with no vulnerabilities.
    """
    table = doc.add_table(rows=4, cols=1)
    table.style = 'Table Grid'
    # Row 1: No Vulnerability Found (Purple)
    cell = table.rows[0].cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    para = cell.paragraphs[0]
    run = para.add_run("No Vulnerability Found")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Row 2: IP
    cell = table.rows[1].cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    para = cell.paragraphs[0]
    run = para.add_run(f"IP: {ip}")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Row 3: Vulnerability Description
    cell = table.rows[2].cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    para = cell.paragraphs[0]
    run = para.add_run("Vulnerability Description: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(106, 68, 154)
    run2 = para.add_run("No vulnerabilities were identified for this IP during the assessment.")
    run2.font.name = 'Altone Trial'
    run2.font.size = Pt(11)
    run2.font.bold = False
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Row 4: Proof of Concept
    cell = table.rows[3].cells[0]
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    para = cell.paragraphs[0]
    run = para.add_run("Proof of Concept: ")
    run.font.name = 'Altone Trial'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(106, 68, 154)
    if steps_with_screenshots and len(steps_with_screenshots) > 0:
        for step in steps_with_screenshots:
            step_para = cell.add_paragraph(step['text'])
            step_para.paragraph_format.left_indent = Pt(10)
            step_para.runs[0].font.size = Pt(11)
    else:
        para.add_run("No Steps Provided").font.size = Pt(11)
    return table

def create_ip_severity_bar_chart(ip, vulnerabilities, output_path):
    """
    Generates a bar chart image for a single IP, showing severity-wise vulnerability counts with legend and color-coding.
    """
    import matplotlib.pyplot as plt
    from matplotlib.patches import Patch
    from collections import Counter
    all_severities = ['Critical', 'High', 'Medium', 'Low', 'Informational']
    all_colors = [get_severity_colors(s)['row1_bg'] for s in all_severities]
    # Count severities for this IP, excluding 'no_vuln' entries
    counts = Counter()
    for v in vulnerabilities:
        if not v.get('no_vuln'):
            counts[v.get('severity', 'Informational')] += 1
    y = [counts.get(s, 0) for s in all_severities]
    fig, ax = plt.subplots(figsize=(6, 4))
    bars = ax.bar(all_severities, y, color=all_colors)
    # Remove Y-axis label
    ax.set_ylabel("")
    # Set title as requested
    ax.set_title('Total Vulnerabilities Found', fontsize=14)
    # Add counts inside bars
    for i, bar in enumerate(bars):
        height = bar.get_height()
        ax.text(bar.get_x() + bar.get_width()/2, height/2, str(y[i]),
                ha='center', va='center', fontsize=12, weight='bold', color='white')
    # Legend at the bottom, centered
    legend_handles = [Patch(facecolor=all_colors[i], label=all_severities[i]) for i in range(len(all_severities))]
    ax.legend(handles=legend_handles, loc='lower center', bbox_to_anchor=(0.5, -0.22), ncol=len(all_severities), fontsize=11, frameon=False)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(False)
    ax.tick_params(left=False, bottom=False)
    ax.yaxis.set_ticks_position('none')
    ax.xaxis.set_ticks_position('none')
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches='tight')
    plt.close()

# Audit Logs Endpoint
@router.get("/audit-logs")
async def get_audit_logs(request: Request, db: Session = Depends(get_db)):
    """Get audit logs for the dashboard"""
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    
    try:
        # Get audit logs from database
        logs = db.query(AuditLog).order_by(AuditLog.created_at.desc()).limit(100).all()
        
        return [
            {
                "id": log.id,
                "user_email": log.user_email,
                "user_name": log.user_name,
                "action": log.action,
                "metadata_json": log.metadata_json,
                "created_at": log.created_at.isoformat()
            }
            for log in logs
        ]
    except Exception as e:
        print(f"Error fetching audit logs: {e}")
        return []

# Usage Analytics Endpoint
@router.get("/usage-analytics")
async def get_usage_analytics(request: Request, db: Session = Depends(get_db)):
    """Get usage analytics for the dashboard"""
    if not is_dashboard_allowed(request):
        raise HTTPException(status_code=403, detail="Forbidden")
    
    try:
        from sqlalchemy import text
        
        # Get usage statistics
        result = db.execute(text("""
            SELECT 
                user_email,
                action_type,
                COUNT(*) as action_count,
                MAX(created_at) as last_activity
            FROM usage_analytics
            WHERE created_at >= NOW() - INTERVAL '30 days'
            GROUP BY user_email, action_type
            ORDER BY action_count DESC
        """))
        
        analytics = []
        for row in result.fetchall():
            analytics.append({
                "user_email": row.user_email,
                "action_type": row.action_type,
                "action_count": row.action_count,
                "last_activity": row.last_activity.isoformat() if row.last_activity else None
            })
        
        return analytics
    except Exception as e:
        print(f"Error fetching usage analytics: {e}")
        return []

# ============================================================================
# MITKAT REPORT GENERATOR (New Report Type)
# ============================================================================

def parse_mitkat_vulnerabilities_excel(file_path):
    """
    Parse Excel file for MitKat report with support for Revalidation Status and Screenshot columns.
    Expected columns:
    - Sr No, Observation/Vulnerability title, Affected Asset, CVE/CWE, Severity, 
    - Detailed observation, Recommendation, Reference, Steps, Screenshot columns,
    - Revalidation Status, Screenshot (new columns)
    """
    print(f"Parsing MitKat vulnerability data from: {file_path}")
    vulnerabilities = []
    
    try:
        # Try standardized parser first
        try:
            parsed_data = parse_excel_data(file_path)
            print(f"✅ Using standardized Excel parser for MitKat")
            
            for idx, vuln in enumerate(parsed_data['vulnerabilities'], 1):
                # Extract revalidation status and screenshot from raw data
                df = parsed_data.get('raw_df')
                revalidation_status = ''
                screenshot_path = ''
                
                if df is not None and idx <= len(df):
                    row = df.iloc[idx - 1]
                    # Find revalidation status column (case-insensitive)
                    for col in df.columns:
                        if 'revalidation' in col.lower() and 'status' in col.lower():
                            revalidation_status = str(row.get(col, '')).strip()
                            break
                    # Find screenshot column (case-insensitive, but not the step screenshots)
                    for col in df.columns:
                        if col.lower() == 'screenshot' and 'step' not in col.lower():
                            screenshot_path = str(row.get(col, '')).strip()
                            break
                
                # Convert steps to steps_with_screenshots format
                steps_with_screenshots = []
                if vuln.get('steps'):
                    for step in vuln['steps']:
                        steps_with_screenshots.append({
                            'text': step['content'],
                            'screenshot': step.get('screenshot', '')
                        })
                
                # Extract purpose from Excel if available
                purpose = ''
                if df is not None and idx <= len(df):
                    row = df.iloc[idx - 1]
                    # Find purpose column (case-insensitive)
                    for col in df.columns:
                        if 'purpose' in col.lower():
                            purpose = str(row.get(col, '')).strip()
                            break
                
                vulnerability = {
                    'sr_no': str(vuln.get('sr_no', idx)),
                    'observation': vuln.get('observation', ''),
                    'affected_asset': vuln.get('affected_asset', '') or vuln.get('ip_url_app', ''),
                    'cve_cwe': vuln.get('cve_cwe', ''),
                    'severity': vuln.get('severity', 'Medium'),
                    'cvss': vuln.get('cvss', ''),
                    'detailed_observation': vuln.get('detailed_observation', '') or vuln.get('observation_summary', ''),
                    'recommendation': vuln.get('recommendation', ''),
                    'reference': vuln.get('reference', ''),
                    'steps_with_screenshots': steps_with_screenshots,
                    'revalidation_status': revalidation_status,
                    'screenshot': screenshot_path,
                    'new_or_repeat': vuln.get('new_or_re', 'New Observation'),
                    'purpose': purpose or 'Web Application',  # Default to 'Web Application' if not found
                    'status': 'Completed'  # Default status
                }
                vulnerabilities.append(vulnerability)
            
            print(f"✅ Parsed {len(vulnerabilities)} vulnerabilities for MitKat report")
            return vulnerabilities
            
        except Exception as std_error:
            print(f"⚠️ Standardized parser failed: {std_error}, falling back to legacy parser")
            traceback.print_exc()
        
        # Fallback to legacy parsing
        df = pd.read_excel(file_path)
        df.columns = df.columns.str.strip()
        
        for idx, row in df.iterrows():
            try:
                # Find columns (case-insensitive)
                def find_col(possible_names):
                    for name in possible_names:
                        for col in df.columns:
                            if name.lower() in col.lower():
                                return col
                    return None
                
                sr_no = str(row.get(find_col(['Sr No', 'Sr.No', 'Serial No']) or 'Sr No', idx + 1)).strip()
                observation = str(row.get(find_col(['Observation', 'Vulnerability Title', 'Title']) or 'Observation', '')).strip()
                affected_asset = str(row.get(find_col(['Affected Asset', 'IP', 'URL', 'Application']) or 'Affected Asset', '')).strip()
                cve_cwe = str(row.get(find_col(['CVE/CWE', 'CVE', 'CWE']) or 'CVE/CWE', '')).strip()
                severity = str(row.get(find_col(['Severity', 'Risk']) or 'Severity', 'Medium')).strip()
                cvss = str(row.get(find_col(['CVSS', 'CVSS Score']) or 'CVSS', '')).strip()
                detailed_obs = str(row.get(find_col(['Detailed Observation', 'Description', 'Details']) or 'Detailed Observation', '')).strip()
                recommendation = str(row.get(find_col(['Recommendation', 'Remediation']) or 'Recommendation', '')).strip()
                reference = str(row.get(find_col(['Reference', 'References']) or 'Reference', '')).strip()
                revalidation_status = str(row.get(find_col(['Revalidation Status']) or 'Revalidation Status', '')).strip()
                screenshot = str(row.get(find_col(['Screenshot']) or 'Screenshot', '')).strip()
                purpose = str(row.get(find_col(['Purpose', 'purpose']) or 'Purpose', 'Web Application')).strip()
                if not purpose or purpose.lower() in ['nan', 'none', '']:
                    purpose = 'Web Application'
                
                # Parse steps
                steps_with_screenshots = []
                steps_text = str(row.get(find_col(['Steps', 'Evidence', 'PoC']) or 'Steps', '')).strip()
                if steps_text and steps_text.lower() not in ['nan', 'none', '']:
                    step_lines = [line.strip() for line in steps_text.split('\n') if line.strip()]
                    for i, step_line in enumerate(step_lines, 1):
                        steps_with_screenshots.append({
                            'text': step_line,
                            'screenshot': ''
                        })
                
                vulnerability = {
                    'sr_no': sr_no,
                    'observation': observation,
                    'affected_asset': affected_asset,
                    'cve_cwe': cve_cwe,
                    'severity': severity,
                    'cvss': cvss,
                    'detailed_observation': detailed_obs,
                    'recommendation': recommendation,
                    'reference': reference,
                    'steps_with_screenshots': steps_with_screenshots,
                    'revalidation_status': revalidation_status,
                    'screenshot': screenshot,
                    'new_or_repeat': 'New Observation',
                    'purpose': purpose,
                    'status': 'Completed'
                }
                vulnerabilities.append(vulnerability)
                
            except Exception as e:
                print(f"Error processing row {idx}: {e}")
                continue
        
        print(f"✅ Parsed {len(vulnerabilities)} vulnerabilities (legacy parser)")
        return vulnerabilities
        
    except Exception as e:
        print(f"Error parsing MitKat Excel: {e}")
        traceback.print_exc()
        return []

def index_mitkat_poc_images(poc_zip_dir):
    """
    Index POC images from ZIP file.
    Expected structure: POC.zip/VUL-001/step1.png, step2.png, etc.
    OR: POC.zip/#1/step1.png, step2.png, etc.
    """
    image_map = {}
    
    if not poc_zip_dir or not os.path.exists(poc_zip_dir):
        return image_map
    
    print(f"Indexing MitKat POC images from: {poc_zip_dir}")
    
    # Walk through directory structure
    for root, dirs, files in os.walk(poc_zip_dir):
        # Check if this is a vulnerability folder (VUL-XXX or #X format)
        folder_name = os.path.basename(root)
        
        # Match VUL-001, VUL-002, etc. or #1, #2, etc.
        vuln_match = re.match(r'(?:VUL-|#)(\d+)', folder_name, re.IGNORECASE)
        if vuln_match:
            vuln_num = int(vuln_match.group(1))
            vuln_key = f"VUL-{vuln_num:03d}"
            
            if vuln_key not in image_map:
                image_map[vuln_key] = {}
            
            # Find all images in this folder
            for fname in sorted(files):
                if fname.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                    # Extract step number from filename
                    step_match = re.search(r'step[_\-\s]?(\d+)', fname.lower())
                    if step_match:
                        step_num = int(step_match.group(1))
                    else:
                        # Try to extract any number
                        num_match = re.search(r'(\d+)', fname)
                        step_num = int(num_match.group(1)) if num_match else len(image_map[vuln_key]) + 1
                    
                    img_path = os.path.join(root, fname)
                    image_map[vuln_key][step_num] = img_path
                    print(f"  Mapped {vuln_key}/Step {step_num}: {fname}")
    
    print(f"✅ Indexed POC images for {len(image_map)} vulnerabilities")
    return image_map

def create_mitkat_observations_table(doc, vulnerabilities, poc_image_map):
    """
    Create the Observations table for MitKat report (pages 9, 10, 11, etc.)
    Each vulnerability gets its own detailed table.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    for idx, vuln in enumerate(vulnerabilities, 1):
        # Create main vulnerability table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.0)
        table.columns[1].width = Inches(5.5)
        
        # Header row
        header_row = table.rows[0]
        
        # Left cell: Sr No, Observation, Affected Asset
        left_cell = header_row.cells[0]
        left_cell.text = ''
        p = left_cell.paragraphs[0]
        
        # Sr No
        run = p.add_run(f"Sr. No.: {vuln.get('sr_no', idx)}")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        p.add_run().add_break()
        
        # Observation/Vulnerability Title
        run = p.add_run(f"Observation/Vulnerability Title: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(vuln.get('observation', 'N/A'))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        p.add_run().add_break()
        
        # Affected Asset
        run = p.add_run(f"Affected Asset: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(vuln.get('affected_asset', 'N/A'))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        
        # Right cell: CVE/CWE, Severity, Status
        right_cell = header_row.cells[1]
        right_cell.text = ''
        p = right_cell.paragraphs[0]
        
        # CVE/CWE
        run = p.add_run(f"CVE/CWE: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(vuln.get('cve_cwe', 'N/A'))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        p.add_run().add_break()
        
        # Severity (with color coding)
        severity = vuln.get('severity', 'Medium')
        colors = get_severity_colors(severity)
        run = p.add_run(f"Severity: {severity}")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
        right_cell._element.get_or_add_tcPr().append(shading_elm)
        p.add_run().add_break()
        
        # New or Repeat Observation
        run = p.add_run(f"New or Repeat Observation: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(vuln.get('new_or_repeat', 'New Observation'))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        
        # Add detailed observation row
        detail_row = table.add_row()
        detail_cell = detail_row.cells[0]
        detail_cell.merge(detail_row.cells[1])
        detail_cell.text = ''
        p = detail_cell.paragraphs[0]
        run = p.add_run("Detailed Observation/Vulnerable Point: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(vuln.get('detailed_observation', 'N/A'))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        
        # Add recommendation row
        rec_row = table.add_row()
        rec_cell = rec_row.cells[0]
        rec_cell.merge(rec_row.cells[1])
        rec_cell.text = ''
        p = rec_cell.paragraphs[0]
        run = p.add_run("Recommendation: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        run = p.add_run(vuln.get('recommendation', 'N/A'))
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = False
        
        # Add reference row if present
        if vuln.get('reference'):
            ref_row = table.add_row()
            ref_cell = ref_row.cells[0]
            ref_cell.merge(ref_row.cells[1])
            ref_cell.text = ''
            p = ref_cell.paragraphs[0]
            run = p.add_run("Reference: ")
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = True
            run = p.add_run(vuln.get('reference', ''))
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = False
            run.font.color.rgb = RGBColor(0, 0, 255)
            run.font.underline = True
        
        # Add Proof of Concept/Steps row
        poc_row = table.add_row()
        poc_cell = poc_row.cells[0]
        poc_cell.merge(poc_row.cells[1])
        poc_cell.text = ''
        p = poc_cell.paragraphs[0]
        run = p.add_run("References to evidence/Proof of Concept: ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.bold = True
        
        # Add steps with screenshots
        steps = vuln.get('steps_with_screenshots', [])
        # Extract vuln_key - cannot use backslashes in f-string expressions
        sr_no_str = str(vuln.get('sr_no', idx))
        match = re.search(r'(\d+)', sr_no_str)
        vuln_key = f"VUL-{int(match.group(1)):03d}" if match else None
        
        for step_idx, step in enumerate(steps, 1):
            step_p = poc_cell.add_paragraph()
            step_p.paragraph_format.left_indent = Pt(10)
            run = step_p.add_run(f"Step {step_idx}: {step.get('text', '')}")
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = False
            
            # Add screenshot if available
            screenshot_path = None
            if poc_image_map and vuln_key and vuln_key in poc_image_map:
                screenshot_path = poc_image_map[vuln_key].get(step_idx)
            
            if screenshot_path and os.path.exists(screenshot_path):
                img_p = poc_cell.add_paragraph()
                img_p.paragraph_format.left_indent = Pt(20)
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_p.add_run()
                img_run.add_picture(screenshot_path, width=Inches(5))
            elif step.get('screenshot'):
                missing_p = poc_cell.add_paragraph()
                missing_p.paragraph_format.left_indent = Pt(20)
                run = missing_p.add_run(f"[Screenshot missing: {step.get('screenshot')}]")
                run.font.name = 'Altone Trial'
                run.font.size = Pt(10)
                run.font.italic = True
        
        # Add Revalidation Status row if present
        if vuln.get('revalidation_status'):
            reval_row = table.add_row()
            reval_cell = reval_row.cells[0]
            reval_cell.merge(reval_row.cells[1])
            reval_cell.text = ''
            p = reval_cell.paragraphs[0]
            run = p.add_run(f"Revalidation Status: {vuln.get('revalidation_status')}")
            run.font.name = 'Altone Trial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        # Add spacing after table
        doc.add_paragraph()
        
        # Add page break for next vulnerability (except last one)
        if idx < len(vulnerabilities):
            doc.add_page_break()

def generate_mitkat_report(vulnerabilities, front_page_data, doc_control_data, template_path, output_path, poc_image_map=None):
    """
    Generate MitKat report from template and data.
    """
    try:
        from docxtpl import DocxTemplate, InlineImage
        
        tpl = DocxTemplate(template_path)
        
        # Prepare context for template
        context = {
            # Front page data
            'CLIENT_NAME': front_page_data.get('client_name', 'Client'),
            'REPORT_TITLE': front_page_data.get('report_title', 'VAPT Final Report'),
            'REPORT_TO': front_page_data.get('report_to', ''),
            
            # Document control data
            'REPORT_RELEASE_DATE': doc_control_data.get('report_release_date', ''),
            'TYPE_OF_AUDIT': doc_control_data.get('type_of_audit', ''),
            'TYPE_OF_AUDIT_REPORT': doc_control_data.get('type_of_audit_report', ''),
            'PERIOD': doc_control_data.get('period', ''),
            'DOCUMENT_TITLE': doc_control_data.get('document_title', ''),
            'DOCUMENT_ID': doc_control_data.get('document_id', ''),
            'DOCUMENT_VERSION': doc_control_data.get('document_version', '1.0'),
            'PREPARED_BY': doc_control_data.get('prepared_by', ''),
            'REVIEWED_BY': doc_control_data.get('reviewed_by', ''),
            'APPROVED_BY': doc_control_data.get('approved_by', ''),
            'RELEASED_BY': doc_control_data.get('released_by', ''),
            'RELEASE_DATE': doc_control_data.get('release_date', ''),
            'DOCUMENT_CHANGE_HISTORY': doc_control_data.get('document_change_history', []),
            'DISTRIBUTION_LIST': doc_control_data.get('distribution_list', []),
            'ENGAGEMENT_SCOPE': doc_control_data.get('engagement_scope', []),
            'AUDITING_TEAM': doc_control_data.get('auditing_team', []),
            'TOOLS_SOFTWARE': doc_control_data.get('tools_software', []),
            'INTRODUCTION': doc_control_data.get('introduction', ''),  # Not required, can be empty
        }
        
        tpl.render(context)
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tpl.save(tmp.name)
            tmp_path = tmp.name
        
        doc = Document(tmp_path)
        
        # Insert overall findings table and charts (for earlier pages)
        insert_mitkat_overall_findings(doc, vulnerabilities)
        
        # Group vulnerabilities by affected_asset (URL/IP)
        assets_dict = {}
        for vuln in vulnerabilities:
            asset = vuln.get('affected_asset', 'Unknown')
            if asset not in assets_dict:
                assets_dict[asset] = []
            assets_dict[asset].append(vuln)
        
        # On page 9, after test case table, create tables for each unique asset
        doc.add_page_break()
        observations_heading = doc.add_paragraph("Observations")
        observations_heading.runs[0].font.name = 'Altone Trial'
        observations_heading.runs[0].font.size = Pt(18)
        observations_heading.runs[0].font.bold = True
        observations_heading.runs[0].font.color.rgb = RGBColor(106, 68, 154)
        observations_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # For each unique asset, create Overall Findings table and Observations table
        for sr_no, (asset, asset_vulns) in enumerate(assets_dict.items(), 1):
            # Extract purpose and status from first vulnerability if available
            purpose = asset_vulns[0].get('purpose', 'Web Application')
            status = asset_vulns[0].get('status', 'Completed')
            
            # Create Overall Findings table (dark blue header)
            create_mitkat_overall_findings_table_per_asset(doc, asset, asset_vulns, sr_no, purpose, status)
            
            # Add spacing
            doc.add_paragraph()
            
            # Create Observations/Vulnerabilities table
            create_mitkat_observations_vulnerabilities_table(doc, asset_vulns)
            
            # Add spacing between assets
            doc.add_paragraph()
            
            # Add page break between assets (except last one)
            if sr_no < len(assets_dict):
                doc.add_page_break()
        
        # After all asset tables, add detailed vulnerability tables (pages 10, 11, etc.)
        doc.add_page_break()
        create_mitkat_observations_table(doc, vulnerabilities, poc_image_map)
        
        doc.save(output_path)
        os.remove(tmp_path)
        
        return output_path
        
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating MitKat report: {str(e)}")

def create_mitkat_overall_findings_table_per_asset(doc, asset, asset_vulns, sr_no, purpose="Web Application", status="Completed"):
    """
    Create the dark blue "Overall Findings" table for a specific asset.
    This table shows: Sr. No., Hostname/Address, purpose, Status, and severity counts.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    
    # Count vulnerabilities by severity for this asset
    severity_counts = {
        'Critical': 0,
        'High': 0,
        'Medium': 0,
        'Low': 0,
        'Informational': 0
    }
    
    for vuln in asset_vulns:
        severity = vuln.get('severity', 'Medium')
        # Normalize severity names
        if severity.lower() in ['info', 'informational']:
            severity = 'Informational'
        if severity in severity_counts:
            severity_counts[severity] += 1
    
    total = sum(severity_counts.values())
    
    # Create table with 10 columns: Sr. No., Hostname/Address, purpose, Status, Critical, High, Medium, Low, Informational, Total
    # Structure: Header row, Data row, Overall Findings summary row
    table = doc.add_table(rows=3, cols=10)
    table.style = 'Table Grid'
    
    # Header row: All columns in one row
    header_row = table.rows[0]
    header_row.cells[0].text = 'Sr. No.'
    header_row.cells[1].text = 'Hostname/Address'
    header_row.cells[2].text = 'purpose'
    header_row.cells[3].text = 'Status'
    header_row.cells[4].text = 'Critical'
    header_row.cells[5].text = 'High'
    header_row.cells[6].text = 'Medium'
    header_row.cells[7].text = 'Low'
    header_row.cells[8].text = 'Informational'
    header_row.cells[9].text = 'Total'
    
    # Dark blue background for header row
    dark_blue = '1F4E78'  # Dark blue color
    for cell in header_row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{dark_blue}"/>')
        cell._element.get_or_add_tcPr().append(shading)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = 'Altone Trial'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Data row
    data_row = table.rows[1]
    data_row.cells[0].text = str(sr_no)
    data_row.cells[1].text = asset
    data_row.cells[2].text = purpose
    data_row.cells[3].text = status
    data_row.cells[4].text = str(severity_counts['Critical'])
    data_row.cells[5].text = str(severity_counts['High'])
    data_row.cells[6].text = str(severity_counts['Medium'])
    data_row.cells[7].text = str(severity_counts['Low'])
    data_row.cells[8].text = str(severity_counts['Informational'])
    data_row.cells[9].text = str(total)
    
    # Style the first 4 columns (no background color)
    for col_idx in range(4):
        cell = data_row.cells[col_idx]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)
    
    # Color code severity columns (columns 4-9)
    severity_colors = {
        4: ('Critical', 'C00000'),      # Dark red
        5: ('High', 'FF6600'),          # Orange/Red
        6: ('Medium', 'FFC000'),        # Yellow/Orange
        7: ('Low', '92D050'),           # Green
        8: ('Informational', '00B0F0'), # Light blue
        9: ('Total', '7030A0')          # Purple
    }
    
    for col_idx, (label, color) in severity_colors.items():
        cell = data_row.cells[col_idx]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Overall Findings summary row (row 2)
    summary_row = table.rows[2]
    # Merge first 4 cells for "Overall Findings"
    summary_row.cells[0].merge(summary_row.cells[1])
    summary_row.cells[0].merge(summary_row.cells[2])
    summary_row.cells[0].merge(summary_row.cells[3])
    summary_row.cells[0].text = 'Overall Findings'
    
    # Dark blue background for merged cell
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{dark_blue}"/>')
    summary_row.cells[0]._element.get_or_add_tcPr().append(shading)
    summary_row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in summary_row.cells[0].paragraphs[0].runs:
        run.font.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Add severity counts to summary row (same as data row)
    summary_row.cells[4].text = str(severity_counts['Critical'])
    summary_row.cells[5].text = str(severity_counts['High'])
    summary_row.cells[6].text = str(severity_counts['Medium'])
    summary_row.cells[7].text = str(severity_counts['Low'])
    summary_row.cells[8].text = str(severity_counts['Informational'])
    summary_row.cells[9].text = str(total)
    
    # Color code severity columns in summary row
    for col_idx, (label, color) in severity_colors.items():
        cell = summary_row.cells[col_idx]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    return table

def create_mitkat_observations_vulnerabilities_table(doc, asset_vulns):
    """
    Create the "Observations/Vulnerabilities" table listing all vulnerabilities for an asset.
    Columns: Sr. No., Affected Asset, Observation/Vulnerability Title, CVE/CWE, Severity, New or Repeat Observation
    """
    if not asset_vulns:
        return None
    
    # Create table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Header row
    headers = [
        'Sr. No.',
        'Affected Asset i.e. IP/URL/Application etc',
        'Observation/Vulnerability Title',
        'CVE/CWE',
        'Severity',
        'New or Repeat Observation'
    ]
    
    header_row = table.rows[0]
    dark_blue = '1F4E78'  # Dark blue color
    for idx, header_text in enumerate(headers):
        cell = header_row.cells[idx]
        p = cell.paragraphs[0]
        p.text = header_text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        # Dark blue background for header
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{dark_blue}"/>')
        cell._element.get_or_add_tcPr().append(shading)
    
    # Add data rows
    for idx, vuln in enumerate(asset_vulns, 1):
        row = table.add_row()
        row.cells[0].text = str(vuln.get('sr_no', idx))
        row.cells[1].text = vuln.get('affected_asset', 'N/A')
        row.cells[2].text = vuln.get('observation', 'N/A')
        row.cells[3].text = vuln.get('cve_cwe', 'N/A')
        
        # Severity with color coding
        severity = vuln.get('severity', 'Medium')
        row.cells[4].text = severity
        
        # Color code severity cell - match image colors
        severity_colors = {
            'Critical': 'C00000',      # Dark red
            'High': 'FF6600',          # Orange/Red
            'Medium': 'FFC000',        # Yellow/Orange
            'Low': '92D050',           # Green
            'Informational': '00B0F0', # Light blue
            'Info': '00B0F0'           # Also handle "Info" variant
        }
        
        # Normalize severity name
        severity_normalized = severity
        if severity.lower() in ['info', 'informational']:
            severity_normalized = 'Informational'
        
        severity_color = severity_colors.get(severity_normalized, severity_colors.get(severity, 'FFFFFF'))
        cell = row.cells[4]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.name = 'Altone Trial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 255, 255)  # White text
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{severity_color}"/>')
        cell._element.get_or_add_tcPr().append(shading)
        
        row.cells[5].text = vuln.get('new_or_repeat', 'New Observation')
        
        # Style all cells
        for cell in row.cells:
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.name = 'Altone Trial'
                    run.font.size = Pt(10)
    
    return table

def insert_mitkat_overall_findings(doc, vulnerabilities):
    """
    Insert overall findings table for MitKat report (legacy function, kept for compatibility).
    Chart generation removed as per requirements.
    """
    # This function is kept for compatibility but doesn't generate charts anymore
    pass

@mitkat_router.post("/mitkat/generate-report/")
async def generate_mitkat_report_endpoint(
    request: Request,
    vulnerability_file: UploadFile = File(...),
    poc_zip: UploadFile = File(None),
    front_page_data: str = Form(...),
    doc_control_data: str = Form(...),
    db: Session = Depends(get_db)
):
    """
    Generate MitKat report from uploaded files and form data.
    """
    try:
        print(f"Starting MitKat report generation")
        
        # Parse form data
        front_data = json.loads(front_page_data)
        doc_control = json.loads(doc_control_data)
        
        # Save vulnerability file
        vuln_file_path = os.path.join(UPLOAD_DIR, f"mitkat_vuln_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx")
        with open(vuln_file_path, "wb") as f:
            f.write(await vulnerability_file.read())
        
        # Parse vulnerabilities
        vulnerabilities = parse_mitkat_vulnerabilities_excel(vuln_file_path)
        
        # Process POC ZIP if provided
        poc_image_map = {}
        if poc_zip:
            tmp_zip_dir = tempfile.mkdtemp()
            zip_path = os.path.join(tmp_zip_dir, poc_zip.filename)
            with open(zip_path, 'wb') as f:
                f.write(await poc_zip.read())
            
            with ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(tmp_zip_dir)
            
            poc_image_map = index_mitkat_poc_images(tmp_zip_dir)
        
        # Template path - check multiple locations
        script_dir = get_script_dir()  # Automation/backend/
        automation_dir = os.path.dirname(script_dir)  # Automation/
        project_root = os.path.dirname(automation_dir)  # Report-Generator-IP-main/
        
        # Try multiple possible locations
        possible_paths = [
            # 1. In backend directory (where main.py is) - PRIMARY LOCATION
            os.path.join(script_dir, "Mitkat_Base_template.docx"),
            # 2. In Automation directory (parent of backend)
            os.path.join(automation_dir, "Mitkat_Base_template.docx"),
            # 3. In project root
            os.path.join(project_root, "Automation", "backend", "Mitkat_Base_template.docx"),
            os.path.join(project_root, "Automation", "Mitkat_Base_template.docx"),
            # 4. Windows local development paths
            r"C:\Users\jerio\Downloads\vapt-app-backup\Report-Generator-IP-main\Automation\backend\Mitkat_Base_template.docx",
            r"C:\Users\jerio\Downloads\vapt-app-backup\Report-Generator-IP-main\Automation\Mitkat_Base_template.docx",
            # 5. Fallback to old name (for backward compatibility)
            os.path.join(script_dir, "MitKat_Template.docx"),
            os.path.join(automation_dir, "MitKat_Template.docx"),
        ]
        
        template_path = None
        for path in possible_paths:
            if os.path.exists(path):
                template_path = path
                print(f"✅ Found MitKat template at: {template_path}")
                break
        
        if not template_path:
            error_msg = f"MitKat template not found. Checked locations:\n"
            for path in possible_paths:
                error_msg += f"  - {path}\n"
            error_msg += "\nPlease ensure Mitkat_Base_template.docx exists in one of these locations."
            raise HTTPException(status_code=404, detail=error_msg)
        
        # Generate report
        client_name = front_data.get('client_name', 'Client')
        output_filename = f"{client_name}_MitKat_Report_{datetime.now().strftime('%Y-%m-%d')}.docx"
        output_path = os.path.join(UPLOAD_DIR, output_filename)
        
        generate_mitkat_report(vulnerabilities, front_data, doc_control, template_path, output_path, poc_image_map)
        
        # Cleanup
        os.remove(vuln_file_path)
        
        # Audit log
        try:
            user = request.session.get('user') or {}
            db.add(AuditLog(
                user_email=user.get('email'),
                user_name=user.get('name'),
                action='generate-mitkat-report',
                metadata_json=json.dumps({
                    'client_name': client_name,
                    'vulnerability_count': len(vulnerabilities)
                }),
                ip_address=request.client.host if request.client else None,
                user_agent=request.headers.get('user-agent')
            ))
            db.commit()
        except Exception:
            pass
        
        return FileResponse(
            path=output_path,
            filename=output_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@mitkat_router.get("/mitkat/", response_class=HTMLResponse)
async def mitkat_form():
    """Serve the MitKat report generator form"""
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>MitKat Report Generator</title>
        <style>
            body { font-family: Calibri, sans-serif; margin: 20px; background-color: #f5f5f5; }
            .container { max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }
            h1 { color: #6923d0; text-align: center; margin-bottom: 30px; }
            .form-section { margin-bottom: 30px; padding: 20px; border: 1px solid #ddd; border-radius: 5px; }
            .form-section h3 { color: #6923d0; margin-top: 0; border-bottom: 2px solid #6923d0; padding-bottom: 10px; }
            .form-row { display: flex; gap: 20px; margin-bottom: 15px; }
            .form-group { flex: 1; }
            .form-group label { display: block; margin-bottom: 5px; font-weight: bold; color: #333; }
            .form-group input, .form-group select, .form-group textarea { width: 100%; padding: 8px; border: 1px solid #ddd; border-radius: 4px; box-sizing: border-box; }
            .form-group textarea { height: 80px; resize: vertical; }
            .btn { background-color: #6923d0; color: white; border: none; padding: 10px 20px; border-radius: 4px; cursor: pointer; margin: 5px; }
            .btn:hover { background-color: #5a1fb8; }
            .btn-success { background-color: #28a745; }
            .btn-success:hover { background-color: #218838; }
            .info-box { background-color: #e7f3ff; border-left: 4px solid #2196F3; padding: 15px; margin: 15px 0; }
            .info-box h4 { margin-top: 0; color: #2196F3; }
            .info-box code { background-color: #fff; padding: 2px 6px; border-radius: 3px; font-family: monospace; }
            .progress { display: none; margin: 20px 0; }
            .progress-bar { background-color: #6923d0; height: 20px; border-radius: 4px; width: 0%; transition: width 0.3s; }
            .result { display: none; margin: 20px 0; padding: 15px; border-radius: 4px; }
            .result.success { background-color: #d4edda; border: 1px solid #c3e6cb; color: #155724; }
            .result.error { background-color: #f8d7da; border: 1px solid #f5c6cb; color: #721c24; }
            .dynamic-item { display: flex; gap: 10px; margin-bottom: 10px; align-items: center; flex-wrap: wrap; }
            .dynamic-item input, .dynamic-item select { flex: 1; min-width: 150px; padding: 8px; border: 1px solid #ddd; border-radius: 4px; }
            .btn-danger { background-color: #dc3545; }
            .btn-danger:hover { background-color: #c82333; }
            .btn-secondary { background-color: #6c757d; }
            .btn-secondary:hover { background-color: #5a6268; }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>🔒 MitKat Report Generator</h1>
            <div style="text-align:center;margin-bottom:16px;">
                <a href="/report_formats.html" class="btn" style="text-decoration:none;display:inline-block;">← Back to Report Formats</a>
            </div>
            
            <form id="mitkatForm">
                <!-- Front Page Data -->
                <div class="form-section">
                    <h3>📄 Front Page Information</h3>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="client_name">Client Name *</label>
                            <input type="text" id="client_name" name="client_name" required>
                        </div>
                        <div class="form-group">
                            <label for="report_title">Report Title *</label>
                            <input type="text" id="report_title" name="report_title" value="PIPELINE MODULE VAPT FINAL REPORT" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="report_to">Report To *</label>
                            <input type="text" id="report_to" name="report_to" placeholder="e.g., RECEIVABLES EXCHANGE OF INDIA LIMITED (RXIL)" required>
                        </div>
                    </div>
                </div>

                <!-- Document Control Data -->
                <div class="form-section">
                    <h3>📋 Document Control</h3>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="report_release_date">Report Release Date *</label>
                            <input type="date" id="report_release_date" name="report_release_date" required>
                        </div>
                        <div class="form-group">
                            <label for="type_of_audit">Type of Audit *</label>
                            <input type="text" id="type_of_audit" name="type_of_audit" value="WebApplication Vulnerability Assessment & Penetration Testing" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="type_of_audit_report">Type of Audit Report *</label>
                            <input type="text" id="type_of_audit_report" name="type_of_audit_report" value="FinalReport" required>
                        </div>
                        <div class="form-group">
                            <label for="period">Period *</label>
                            <input type="text" id="period" name="period" placeholder="e.g., 07-07-2025 to 18-07-2025" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="document_title">Document Title *</label>
                            <input type="text" id="document_title" name="document_title" value="Webapplication VAPT - Final Report" required>
                        </div>
                        <div class="form-group">
                            <label for="document_id">Document ID *</label>
                            <input type="text" id="document_id" name="document_id" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="document_version">Document Version *</label>
                            <input type="text" id="document_version" name="document_version" value="1.0" required>
                        </div>
                        <div class="form-group">
                            <label for="prepared_by">Prepared by *</label>
                            <input type="text" id="prepared_by" name="prepared_by" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="reviewed_by">Reviewed by *</label>
                            <input type="text" id="reviewed_by" name="reviewed_by" required>
                        </div>
                        <div class="form-group">
                            <label for="approved_by">Approved by *</label>
                            <input type="text" id="approved_by" name="approved_by" required>
                        </div>
                    </div>
                    <div class="form-row">
                        <div class="form-group">
                            <label for="released_by">Released by *</label>
                            <input type="text" id="released_by" name="released_by" required>
                        </div>
                        <div class="form-group">
                            <label for="release_date">Release Date *</label>
                            <input type="date" id="release_date" name="release_date" required>
                        </div>
                    </div>
                </div>

                <!-- Document Change History -->
                <div class="form-section">
                    <h3>📝 Document Change History</h3>
                    <div id="changeHistoryList">
                        <div class="dynamic-item">
                            <input type="text" name="change_version[]" placeholder="Version" required>
                            <input type="date" name="change_date[]" required>
                            <input type="text" name="change_remarks[]" placeholder="Remarks/Reason of change" required>
                            <button type="button" class="btn btn-danger" onclick="removeChangeHistory(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addChangeHistory()">+ Add Change History</button>
                </div>

                <!-- Document Distribution List -->
                <div class="form-section">
                    <h3>📧 Document Distribution List</h3>
                    <div id="distributionList">
                        <div class="dynamic-item">
                            <input type="text" name="dist_name[]" placeholder="Name" required>
                            <input type="text" name="dist_organization[]" placeholder="Organization" required>
                            <input type="text" name="dist_designation[]" placeholder="Designation" required>
                            <input type="email" name="dist_email[]" placeholder="Email ID" required>
                            <button type="button" class="btn btn-danger" onclick="removeDistribution(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addDistribution()">+ Add Distribution</button>
                </div>

                <!-- Engagement Scope -->
                <div class="form-section">
                    <h3>🎯 Engagement Scope</h3>
                    <div id="engagementScopeList">
                        <div class="dynamic-item">
                            <input type="number" name="scope_sr_no[]" placeholder="S. No">
                            <input type="text" name="scope_asset[]" placeholder="Asset Description">
                            <input type="text" name="scope_criticality[]" placeholder="Criticality">
                            <input type="text" name="scope_url[]" placeholder="URL or NA">
                            <input type="text" name="scope_location[]" placeholder="Location">
                            <input type="text" name="scope_hash[]" placeholder="Hash Value">
                            <input type="text" name="scope_version[]" placeholder="Version">
                            <input type="text" name="scope_other[]" placeholder="Other Details">
                            <button type="button" class="btn btn-danger" onclick="removeEngagementScope(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addEngagementScope()">+ Add Asset</button>
                </div>

                <!-- Details of Auditing Team -->
                <div class="form-section">
                    <h3>👥 Details of Auditing Team</h3>
                    <div id="auditingTeamList">
                        <div class="dynamic-item">
                            <input type="text" name="team_name[]" placeholder="Name" required>
                            <input type="text" name="team_designation[]" placeholder="Designation" required>
                            <input type="email" name="team_email[]" placeholder="Email ID" required>
                            <input type="text" name="team_qualifications[]" placeholder="Qualifications/Certifications" required>
                            <select name="team_certin_listed[]" required>
                                <option value="">CERT-In Listed?</option>
                                <option value="Yes">Yes</option>
                                <option value="No">No</option>
                            </select>
                            <button type="button" class="btn btn-danger" onclick="removeAuditingTeam(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addAuditingTeam()">+ Add Team Member</button>
                </div>

                <!-- Tools/Software Used -->
                <div class="form-section">
                    <h3>🛠️ Tools/Software Used</h3>
                    <div id="toolsSoftwareList">
                        <div class="dynamic-item">
                            <input type="number" name="tool_sr_no[]" placeholder="S. No" required>
                            <input type="text" name="tool_name[]" placeholder="Name of Tool/Software" required>
                            <input type="text" name="tool_version[]" placeholder="Version" required>
                            <select name="tool_type[]" required>
                                <option value="">Type</option>
                                <option value="Opensource">Opensource</option>
                                <option value="Licensed">Licensed</option>
                            </select>
                            <button type="button" class="btn btn-danger" onclick="removeToolSoftware(this)">Remove</button>
                        </div>
                    </div>
                    <button type="button" class="btn btn-secondary" onclick="addToolSoftware()">+ Add Tool</button>
                </div>

                <!-- Vulnerability Data -->
                <div class="form-section">
                    <h3>🔍 Vulnerability Data</h3>
                    <div class="info-box">
                        <h4>📋 Excel Format Requirements</h4>
                        <p><strong>Required columns:</strong></p>
                        <ul>
                            <li><code>Sr No</code> - Serial number</li>
                            <li><code>Observation/Vulnerability Title</code> - Vulnerability name</li>
                            <li><code>Affected Asset</code> - IP/URL/Application</li>
                            <li><code>CVE/CWE</code> - CVE or CWE identifier</li>
                            <li><code>Severity</code> - Critical, High, Medium, Low, Informational</li>
                            <li><code>Detailed Observation</code> - Detailed description</li>
                            <li><code>Recommendation</code> - Remediation steps</li>
                            <li><code>Reference</code> - External references</li>
                            <li><code>Steps</code> - Proof of concept steps</li>
                            <li><code>Revalidation Status</code> - Status of revalidation (NEW COLUMN)</li>
                            <li><code>Screenshot</code> - Screenshot filename (NEW COLUMN)</li>
                        </ul>
                    </div>
                    <div class="form-group">
                        <label for="vulnerability_file">Upload Vulnerability Excel File *</label>
                        <input type="file" id="vulnerability_file" name="vulnerability_file" accept=".xlsx,.xls" required>
                    </div>
                    
                    <div class="info-box">
                        <h4>📁 PoC Zip Structure</h4>
                        <p><strong>Recommended structure:</strong></p>
                        <code>POC.zip/VUL-001/step1.png, step2.png, ...</code>
                        <p><strong>OR:</strong></p>
                        <code>POC.zip/#1/step1.png, step2.png, ...</code>
                        <p>Where VUL-001 or #1 corresponds to the Sr No in Excel.</p>
                    </div>
                    <div class="form-group">
                        <label for="poc_zip">Upload PoC Zip File (Optional)</label>
                        <input type="file" id="poc_zip" name="poc_zip" accept=".zip">
                    </div>
                </div>

                <div style="text-align: center; margin-top: 30px;">
                    <button type="submit" class="btn btn-success" style="font-size: 18px; padding: 15px 30px;">
                        🚀 Generate MitKat Report
                    </button>
                </div>
            </form>

            <div class="progress" id="progress">
                <div style="background-color: #f0f0f0; border-radius: 4px; padding: 10px;">
                    <div id="progressText">Processing...</div>
                    <div class="progress-bar" id="progressBar"></div>
                </div>
            </div>

            <div class="result" id="result"></div>
        </div>

        <script>
            // Get CSRF token
            let CSRF_TOKEN = null;
            (async () => {
                try {
                    const res = await fetch('/csrf-token', { credentials: 'same-origin' });
                    if (res.ok) {
                        const data = await res.json().catch(() => ({}));
                        CSRF_TOKEN = data.csrf_token || data.token || data.csrf || null;
                    }
                } catch (_) {}
            })();

            // Functions to add/remove dynamic items
            function addChangeHistory() {
                const container = document.getElementById('changeHistoryList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="change_version[]" placeholder="Version" required>
                    <input type="date" name="change_date[]" required>
                    <input type="text" name="change_remarks[]" placeholder="Remarks/Reason of change" required>
                    <button type="button" class="btn btn-danger" onclick="removeChangeHistory(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeChangeHistory(button) {
                button.parentElement.remove();
            }

            function addDistribution() {
                const container = document.getElementById('distributionList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="dist_name[]" placeholder="Name" required>
                    <input type="text" name="dist_organization[]" placeholder="Organization" required>
                    <input type="text" name="dist_designation[]" placeholder="Designation" required>
                    <input type="email" name="dist_email[]" placeholder="Email ID" required>
                    <button type="button" class="btn btn-danger" onclick="removeDistribution(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeDistribution(button) {
                button.parentElement.remove();
            }

            function addEngagementScope() {
                const container = document.getElementById('engagementScopeList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="number" name="scope_sr_no[]" placeholder="S. No">
                    <input type="text" name="scope_asset[]" placeholder="Asset Description">
                    <input type="text" name="scope_criticality[]" placeholder="Criticality">
                    <input type="text" name="scope_url[]" placeholder="URL or NA">
                    <input type="text" name="scope_location[]" placeholder="Location">
                    <input type="text" name="scope_hash[]" placeholder="Hash Value">
                    <input type="text" name="scope_version[]" placeholder="Version">
                    <input type="text" name="scope_other[]" placeholder="Other Details">
                    <button type="button" class="btn btn-danger" onclick="removeEngagementScope(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeEngagementScope(button) {
                button.parentElement.remove();
            }

            function addAuditingTeam() {
                const container = document.getElementById('auditingTeamList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="text" name="team_name[]" placeholder="Name" required>
                    <input type="text" name="team_designation[]" placeholder="Designation" required>
                    <input type="email" name="team_email[]" placeholder="Email ID" required>
                    <input type="text" name="team_qualifications[]" placeholder="Qualifications/Certifications" required>
                    <select name="team_certin_listed[]" required>
                        <option value="">CERT-In Listed?</option>
                        <option value="Yes">Yes</option>
                        <option value="No">No</option>
                    </select>
                    <button type="button" class="btn btn-danger" onclick="removeAuditingTeam(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeAuditingTeam(button) {
                button.parentElement.remove();
            }

            function addToolSoftware() {
                const container = document.getElementById('toolsSoftwareList');
                const div = document.createElement('div');
                div.className = 'dynamic-item';
                div.innerHTML = `
                    <input type="number" name="tool_sr_no[]" placeholder="S. No" required>
                    <input type="text" name="tool_name[]" placeholder="Name of Tool/Software" required>
                    <input type="text" name="tool_version[]" placeholder="Version" required>
                    <select name="tool_type[]" required>
                        <option value="">Type</option>
                        <option value="Opensource">Opensource</option>
                        <option value="Licensed">Licensed</option>
                    </select>
                    <button type="button" class="btn btn-danger" onclick="removeToolSoftware(this)">Remove</button>
                `;
                container.appendChild(div);
            }

            function removeToolSoftware(button) {
                button.parentElement.remove();
            }

            document.getElementById('mitkatForm').addEventListener('submit', async function(e) {
                e.preventDefault();
                
                const formData = new FormData();
                
                // Front page data
                const frontPageData = {
                    client_name: document.getElementById('client_name').value,
                    report_title: document.getElementById('report_title').value,
                    report_to: document.getElementById('report_to').value
                };
                formData.append('front_page_data', JSON.stringify(frontPageData));
                
                // Collect Document Change History
                const changeVersions = document.querySelectorAll('input[name="change_version[]"]');
                const changeDates = document.querySelectorAll('input[name="change_date[]"]');
                const changeRemarks = document.querySelectorAll('input[name="change_remarks[]"]');
                const changeHistory = [];
                for (let i = 0; i < changeVersions.length; i++) {
                    changeHistory.push({
                        version: changeVersions[i].value,
                        date: changeDates[i].value,
                        remarks: changeRemarks[i].value
                    });
                }
                
                // Collect Distribution List
                const distNames = document.querySelectorAll('input[name="dist_name[]"]');
                const distOrgs = document.querySelectorAll('input[name="dist_organization[]"]');
                const distDesignations = document.querySelectorAll('input[name="dist_designation[]"]');
                const distEmails = document.querySelectorAll('input[name="dist_email[]"]');
                const distributionList = [];
                for (let i = 0; i < distNames.length; i++) {
                    distributionList.push({
                        name: distNames[i].value,
                        organization: distOrgs[i].value,
                        designation: distDesignations[i].value,
                        email: distEmails[i].value
                    });
                }
                
                // Collect Engagement Scope
                const scopeSrNos = document.querySelectorAll('input[name="scope_sr_no[]"]');
                const scopeAssets = document.querySelectorAll('input[name="scope_asset[]"]');
                const scopeCriticalities = document.querySelectorAll('input[name="scope_criticality[]"]');
                const scopeURLs = document.querySelectorAll('input[name="scope_url[]"]');
                const scopeLocations = document.querySelectorAll('input[name="scope_location[]"]');
                const scopeHashes = document.querySelectorAll('input[name="scope_hash[]"]');
                const scopeVersions = document.querySelectorAll('input[name="scope_version[]"]');
                const scopeOthers = document.querySelectorAll('input[name="scope_other[]"]');
                const engagementScope = [];
                for (let i = 0; i < scopeSrNos.length; i++) {
                    engagementScope.push({
                        sr_no: parseInt(scopeSrNos[i].value) || (i + 1),
                        asset_description: scopeAssets[i].value,
                        criticality: scopeCriticalities[i].value,
                        url: (scopeURLs[i].value && scopeURLs[i].value.toUpperCase() !== 'NA') ? scopeURLs[i].value : '',
                        location: scopeLocations[i].value,
                        hash_value: scopeHashes[i].value,
                        version: scopeVersions[i].value,
                        other_details: scopeOthers[i].value
                    });
                }
                
                // Collect Auditing Team
                const teamNames = document.querySelectorAll('input[name="team_name[]"]');
                const teamDesignations = document.querySelectorAll('input[name="team_designation[]"]');
                const teamEmails = document.querySelectorAll('input[name="team_email[]"]');
                const teamQualifications = document.querySelectorAll('input[name="team_qualifications[]"]');
                const teamCertinListed = document.querySelectorAll('select[name="team_certin_listed[]"]');
                const auditingTeam = [];
                for (let i = 0; i < teamNames.length; i++) {
                    auditingTeam.push({
                        sr_no: i + 1,
                        name: teamNames[i].value,
                        designation: teamDesignations[i].value,
                        email: teamEmails[i].value,
                        qualifications: teamQualifications[i].value,
                        certin_listed: teamCertinListed[i].value
                    });
                }
                
                // Collect Tools/Software
                const toolSrNos = document.querySelectorAll('input[name="tool_sr_no[]"]');
                const toolNames = document.querySelectorAll('input[name="tool_name[]"]');
                const toolVersions = document.querySelectorAll('input[name="tool_version[]"]');
                const toolTypes = document.querySelectorAll('select[name="tool_type[]"]');
                const toolsSoftware = [];
                for (let i = 0; i < toolSrNos.length; i++) {
                    toolsSoftware.push({
                        sr_no: parseInt(toolSrNos[i].value) || (i + 1),
                        name: toolNames[i].value,
                        version: toolVersions[i].value,
                        type: toolTypes[i].value
                    });
                }
                
                // Document control data
                const docControlData = {
                    report_release_date: document.getElementById('report_release_date').value,
                    type_of_audit: document.getElementById('type_of_audit').value,
                    type_of_audit_report: document.getElementById('type_of_audit_report').value,
                    period: document.getElementById('period').value,
                    document_title: document.getElementById('document_title').value,
                    document_id: document.getElementById('document_id').value,
                    document_version: document.getElementById('document_version').value,
                    prepared_by: document.getElementById('prepared_by').value,
                    reviewed_by: document.getElementById('reviewed_by').value,
                    approved_by: document.getElementById('approved_by').value,
                    released_by: document.getElementById('released_by').value,
                    release_date: document.getElementById('release_date').value,
                    introduction: '',  // Not required
                    document_change_history: changeHistory,
                    distribution_list: distributionList,
                    engagement_scope: engagementScope,
                    auditing_team: auditingTeam,
                    tools_software: toolsSoftware
                };
                formData.append('doc_control_data', JSON.stringify(docControlData));
                
                // Files
                formData.append('vulnerability_file', document.getElementById('vulnerability_file').files[0]);
                if (document.getElementById('poc_zip').files[0]) {
                    formData.append('poc_zip', document.getElementById('poc_zip').files[0]);
                }
                
                // Add CSRF token to form data as well (backup)
                if (CSRF_TOKEN) {
                    formData.append('csrf_token', CSRF_TOKEN);
                }
                
                try {
                    document.getElementById('progress').style.display = 'block';
                    document.getElementById('progressText').textContent = 'Generating report...';
                    document.getElementById('progressBar').style.width = '50%';
                    
                    // Add CSRF token to headers
                    const headers = {};
                    if (CSRF_TOKEN) {
                        headers['X-CSRF-Token'] = CSRF_TOKEN;
                    }
                    
                    const response = await fetch('/mitkat/generate-report/', {
                        method: 'POST',
                        headers: headers,
                        body: formData,
                        credentials: 'same-origin'
                    });
                    
                    if (!response.ok) {
                        const errorData = await response.json();
                        throw new Error(errorData.detail || 'Report generation failed');
                    }
                    
                    // Download the file
                    const blob = await response.blob();
                    const url = window.URL.createObjectURL(blob);
                    const a = document.createElement('a');
                    a.href = url;
                    a.download = response.headers.get('content-disposition')?.split('filename=')[1]?.replace(/"/g, '') || 'MitKat_Report.docx';
                    document.body.appendChild(a);
                    a.click();
                    window.URL.revokeObjectURL(url);
                    document.body.removeChild(a);
                    
                    document.getElementById('progressText').textContent = 'Report generated successfully!';
                    document.getElementById('progressBar').style.width = '100%';
                    
                    document.getElementById('result').innerHTML = `
                        <h3>✅ Report Generated Successfully!</h3>
                        <p>The report has been downloaded.</p>
                    `;
                    document.getElementById('result').className = 'result success';
                    document.getElementById('result').style.display = 'block';
                    
                } catch (error) {
                    document.getElementById('progress').style.display = 'none';
                    document.getElementById('result').innerHTML = `
                        <h3>❌ Error</h3>
                        <p>${error.message}</p>
                    `;
                    document.getElementById('result').className = 'result error';
                    document.getElementById('result').style.display = 'block';
                }
            });
        </script>
    </body>
    </html>
    """
    return HTMLResponse(content=html_content)