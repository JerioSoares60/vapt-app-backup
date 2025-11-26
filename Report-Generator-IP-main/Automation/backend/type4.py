from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import JSONResponse, FileResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
import os
import sys
import pandas as pd
from docxtpl import DocxTemplate, InlineImage
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from datetime import datetime
import traceback
import io
import shutil
import tempfile
import matplotlib.pyplot as plt
import numpy as np
from collections import Counter
import urllib.parse
from zipfile import ZipFile
import glob
import hashlib
import re
import base64

# Try to import PIL for image processing, but don't fail if it's not available
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: PIL/Pillow library not available. Logo resizing will be disabled.")

# Add parent directory to path to import excel_parser
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../..'))
from excel_parser import parse_excel_data, format_for_template

# Version tracking for code changes
# Format: MAJOR.MINOR.PATCH
# Increment MAJOR for significant structural changes
# Increment MINOR for feature additions or improvements
# Increment PATCH for bug fixes
VERSION = "1.0.7"  # Updated PATCH version for splitting parsing logic

app = FastAPI()

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Directory to store uploaded files
UPLOAD_DIR = "uploads"
# Define the fixed path for the document control file
DOC_CONTROL_FILE = os.path.join(UPLOAD_DIR, "Comprehensive_Vulnerability_Template (2).xlsx")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(os.path.join(UPLOAD_DIR, "screenshots"), exist_ok=True)

# Define the order of severities
severity_order = {"Critical": 0, "High": 1, "Medium": 2, "Low": 3, "Informational": 4, "Unknown": 5}

@app.post("/upload/")
async def upload_file(file: UploadFile = File(...)):
    try:
        file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(file_location, "wb") as f:
            f.write(await file.read())
        return JSONResponse(content={"filename": file.filename}, status_code=200)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/upload-screenshots/")
async def upload_screenshots(files: list[UploadFile] = File(...)):
    """
    Upload multiple screenshot files.
    Returns a dictionary of the saved filenames.
    """
    try:
        saved_files = []
        for file in files:
            # Only allow image files
            if not file.content_type.startswith('image/'):
                continue
                
            file_location = os.path.join(UPLOAD_DIR, "screenshots", file.filename)
            with open(file_location, "wb") as f:
                f.write(await file.read())
            saved_files.append(file.filename)
        
        return JSONResponse(
            content={"filenames": saved_files},
            status_code=200
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

# Function to determine severity from CVSS score
def get_severity_from_cvss(cvss):
    print(f"Determining severity for CVSS: {cvss}, type: {type(cvss)}")  # Debug output
    
    if cvss is None:
        print("CVSS is None, defaulting to Informational")
        return "Informational"
    
    try:
        cvss_float = float(cvss)
        if cvss_float >= 9.0 and cvss_float <= 10.0:
            print(f"CVSS {cvss_float} classified as Critical")
            return "Critical"
        elif cvss_float >= 7.0 and cvss_float < 9.0:
            print(f"CVSS {cvss_float} classified as High")
            return "High"
        elif cvss_float >= 4.0 and cvss_float < 7.0:
            print(f"CVSS {cvss_float} classified as Medium")
            return "Medium"
        elif cvss_float > 0.0 and cvss_float < 4.0:
            print(f"CVSS {cvss_float} classified as Low")
            return "Low"
        else:
            print(f"CVSS {cvss_float} classified as Informational")
            return "Informational"
    except (ValueError, TypeError) as e:
        print(f"Error converting CVSS to float: {e}, defaulting to Informational")
        return "Informational"

# Function to get severity colors based on severity level
def get_severity_colors(severity):
    colors = {
        "Critical": {"row1_bg": "#990000", "row2_bg": "#FF3333", "font": "#990000"},
        "High": {"row1_bg": "#FF0000", "row2_bg": "#FF6666", "font": "#FF0000"},
        "Medium": {"row1_bg": "#FFCC00", "row2_bg": "#FFCC66", "font": "#FFCC00"},
        "Low": {"row1_bg": "#009933", "row2_bg": "#99CC33", "font": "#009933"},
        "Informational": {"row1_bg": "#3399CC", "row2_bg": "#66CCFF", "font": "#3399CC"}
    }
    return colors.get(severity, colors["Informational"])

# Function to generate more detailed vulnerability content based on vulnerability name
def get_vulnerability_description(vuln_name):
    """
    Generate a detailed description paragraph for the vulnerability based on OWASP guidance.
    """
    descriptions = {
        "SQL Injection": """SQL Injection is a code injection technique that exploits vulnerabilities in the interface between web applications and database servers. The vulnerability occurs when user input is incorrectly filtered and directly included in SQL statements, allowing attackers to manipulate the structure of the SQL query. This can lead to unauthorized access to sensitive data, bypassing authentication mechanisms, and in severe cases, compromising the underlying server or performing denial of service attacks. The primary cause is insufficient validation and sanitization of user-supplied input used in database queries.""",
        
        "Cross-Site Scripting": """Cross-Site Scripting (XSS) is a security vulnerability that allows attackers to inject malicious client-side scripts into web pages viewed by other users. These attacks occur when an application includes untrusted data in a new web page without proper validation or encoding. XSS vulnerabilities can be found in various forms: Reflected XSS where malicious script is reflected off the web server in an error message or search result; Stored XSS where the malicious script is permanently stored on target servers; and DOM-based XSS where vulnerability exists in client-side code. The impact of XSS can range from session theft to complete compromise of user accounts.""",
        
        "File Upload": """Unrestricted File Upload vulnerabilities allow attackers to upload files of dangerous types or content to the server. This vulnerability occurs when an application fails to properly validate the type, content, size, or naming convention of uploaded files. Attackers can exploit this to upload web shells, malware, or other malicious files that allow them to execute arbitrary code on the server. This may lead to complete system compromise, data theft, or defacement of websites. The vulnerability is particularly dangerous when uploaded files are stored in web-accessible locations, allowing direct execution through browser requests.""",
        
        "Insecure Direct Object References": """Insecure Direct Object References (IDOR) occur when an application provides direct access to objects based on user-supplied input without appropriate authorization checks. This vulnerability allows attackers to bypass authorization and directly access resources in the system, potentially exposing sensitive data that should be protected. IDOR vulnerabilities commonly manifest as manipulated parameters within requests that reveal sensitive objects like files, database records, or user accounts. This can lead to unauthorized access to sensitive information, modification of data belonging to other users, or privilege escalation allowing attackers to gain administrative access to the application.""",
        
        "Sensitive Data Exposure": """Sensitive Data Exposure occurs when an application fails to adequately protect sensitive information such as financial data, healthcare information, personally identifiable information (PII), or authentication credentials. This vulnerability can manifest through various channels including insufficient encryption of data at rest or in transit, improper handling of sensitive information in logs or error messages, and inadequate access controls. Attackers can exploit these weaknesses to access sensitive data through packet sniffing, man-in-the-middle attacks, or by directly accessing poorly protected storage. The impact can include identity theft, financial fraud, privacy violations, and regulatory compliance failures."""
    }
    
    # Try to find a match in our predefined descriptions
    for key in descriptions:
        if key.lower() in vuln_name.lower():
            return descriptions[key]
    
    # Default description for any vulnerability not specifically defined
    return f"""This vulnerability ({vuln_name}) allows attackers to potentially compromise the application security through improper input validation and processing mechanisms. The vulnerability exists due to insufficient security controls in the application code, which could allow malicious actors to manipulate system behavior in unintended ways. Such vulnerabilities often arise from coding practices that do not account for all possible edge cases or user inputs, creating potential attack vectors. When discovered, these issues should be addressed promptly to prevent potential exploitation."""

def get_vulnerability_impact(vuln_name):
    """
    Generate a detailed impact paragraph for the vulnerability based on OWASP guidance.
    """
    impacts = {
        "SQL Injection": """The impact of SQL Injection vulnerabilities can be severe and far-reaching for an organization. Successful exploitation can lead to unauthorized access to sensitive database information, including personal user data, financial records, and intellectual property. Attackers may bypass authentication mechanisms, gaining unauthorized administrative access to the application. In more severe cases, attackers might execute commands on the database server, potentially leading to complete server compromise. This can result in significant financial losses, regulatory penalties for data breaches, damage to company reputation, and loss of customer trust. Business operations may be disrupted if data is stolen, manipulated, or deleted entirely.""",
        
        "Cross-Site Scripting": """The exploitation of Cross-Site Scripting vulnerabilities can have significant consequences for both users and the organization. Attackers can hijack user sessions, steal authentication cookies, and impersonate legitimate users to access sensitive information or perform unauthorized actions. They may also inject malicious content that defaces websites, leads to redirection to malicious sites, or installs malware on user devices. For organizations, this can lead to reputational damage, loss of customer trust, potential data breaches, and regulatory compliance issues. XSS attacks targeting administrative users can result in complete application compromise, allowing attackers to access backend systems, modify server-side logic, or obtain sensitive business data.""",
        
        "File Upload": """Unrestricted file upload vulnerabilities can have devastating impacts on organization security. Successfully exploited vulnerabilities allow attackers to gain remote code execution capabilities on the server, potentially leading to complete system compromise. This can result in unauthorized access to databases, internal networks, and sensitive enterprise systems. Attackers may establish persistence through backdoors, maintain long-term access to compromised systems, and use the server as a platform to attack other internal systems. Organizations face significant risks including data theft, intellectual property loss, financial loss, service disruption due to regulatory penalties. The reputational damage from such breaches can lead to loss of customer confidence and business opportunities.""",
        
        "Insecure Direct Object References": """The business impact of Insecure Direct Object References can be substantial. Exploitation allows attackers to access unauthorized data across user boundaries, potentially exposing sensitive personal information, financial data, or proprietary business information. This can lead to privacy violations, identity theft, financial fraud, or industrial espionage. Organizations may face significant regulatory penalties for data protection failures, particularly under frameworks like GDPR, HIPAA, or PCI-DSS. The reputational damage from such data breaches can severely impact customer trust and result in customer churn. Additionally, competitors may gain access to confidential business information, leading to competitive disadvantage and potential revenue loss.""",
        
        "Sensitive Data Exposure": """Exposure of sensitive data can have catastrophic consequences for both individuals and organizations. When personal information is compromised, individuals may become victims of identity theft, financial fraud, or reputational damage. For organizations, the impacts include regulatory fines and penalties under data protection laws such as GDPR, HIPAA, or PCI-DSS, which can amount to millions of dollars. Legal costs may escalate through lawsuits from affected parties seeking damages. Business reputation suffers significantly, leading to loss of customer trust, decreased market share, and reduced stock value for public companies. Operational disruptions can occur while addressing the breach, and long-term business relationships may be damaged if partner data is compromised."""
    }
    
    # Try to find a match in our predefined impacts
    for key in impacts:
        if key.lower() in vuln_name.lower():
            return impacts[key]
    
    # Default impact for any vulnerability not specifically defined
    return f"""Exploitation of this vulnerability ({vuln_name}) could lead to unauthorized access, data breaches, reputation damage, and potential regulatory compliance violations. Customers' trust may be eroded, resulting in business loss. The organization may face financial penalties from regulatory bodies for failing to protect sensitive information adequately. Remediation costs can be substantial, including not only the technical fixes but also potential legal fees, customer notification expenses, and increased security monitoring. In severe cases, business operations might be disrupted, leading to revenue loss and long-term damage to the organization's brand and market position."""

def get_vulnerability_recommendation(vuln_name):
    """
    Generate detailed remediation recommendations based on OWASP guidance.
    """
    recommendations = {
        "SQL Injection": """To remediate SQL Injection vulnerabilities, implement parameterized queries (prepared statements) for all database operations, which ensures that user input is treated as data rather than executable code. Apply the principle of least privilege to database accounts used by applications, limiting their capabilities to only necessary operations. Implement input validation on both client and server sides, with a whitelist approach that accepts only known-good input. Utilize stored procedures and ORM (Object-Relational Mapping) frameworks that handle SQL escaping automatically. Implement a WAF (Web Application Firewall) as an additional layer of protection to filter malicious SQL injection attempts. Regularly scan application code for SQL injection vulnerabilities and conduct penetration testing to verify the effectiveness of implemented controls.""",
        
        "Cross-Site Scripting": """To effectively mitigate Cross-Site Scripting vulnerabilities, implement context-sensitive output encoding for all user-supplied data before rendering it in HTML, JavaScript, CSS, or URL contexts. Use Content Security Policy (CSP) headers to restrict sources of executable scripts and prevent execution of inline JavaScript. Apply input validation with a whitelist approach to filter potentially dangerous inputs before processing. Utilize modern web frameworks that automatically handle XSS protection through built-in template escaping mechanisms. For cookie security, implement the HttpOnly flag to prevent JavaScript access and use the Secure flag for transmission over HTTPS only. Keep all frameworks and libraries updated to ensure known XSS vulnerabilities are patched. Regularly conduct security testing including both automated scanning and manual code review to identify potential XSS vulnerabilities.""",
        
        "File Upload": """To securely handle file uploads, implement strict file type validation using content inspection (checking file headers) rather than relying only on file extensions. Enforce file size limits to prevent denial of service attacks through oversized file uploads. Store uploaded files outside the web root or application directory to prevent direct execution, and if files must be publicly accessible, store them in a different domain or subdomain. Generate random file names to prevent directory traversal attempts and consistent naming patterns that could be guessed. Scan uploaded files with antivirus software or integrate with security services that can detect malicious content. Implement proper access controls to ensure uploaded files are only accessible by authorized users. Consider using a CDN or dedicated file storage service that specializes in secure file handling for high-risk environments.""",
        
        "Insecure Direct Object References": """To remediate Insecure Direct Object References, implement proper access control checks for every request that accesses data objects, verifying that the authenticated user has appropriate permissions for the requested resource. Replace direct references to objects with indirect references or tokens that are mapped server-side to actual resource identifiers. Use server-side session objects to store and verify user authorization to access specific resources. Create a centralized authorization mechanism that provides consistent enforcement across the application. Implement the principle of least privilege, ensuring users only have access to resources necessary for their legitimate functions. Log access attempts to sensitive resources and configure alerts for potential access control violations. Regularly review and test access control mechanisms through security assessments and penetration testing.""",
        
        "Sensitive Data Exposure": """To protect sensitive data effectively, implement strong encryption for data both at rest and in transit using industry-standard algorithms and appropriate key management. Use TLS (at least version 1.2 or higher) for all connections that transmit sensitive information, and enforce this through HSTS headers. Store passwords using strong adaptive hashing algorithms with salts (such as bcrypt, Argon2, or PBKDF2). Disable caching for responses containing sensitive data using appropriate cache-control headers. Minimize the collection and retention of sensitive data – only collect what is absolutely necessary for business functions and purge unnecessary data regularly. Implement proper key management practices including secure key generation, storage, distribution, and rotation. Classify data based on sensitivity and apply appropriate controls based on classification. Regularly audit systems for sensitive data handling compliance and conduct penetration testing focused on data protection controls."""
    }
    
    # Try to find a match in our predefined recommendations
    for key in recommendations:
        if key.lower() in vuln_name.lower():
            return recommendations[key]
    
    # Default recommendation for any vulnerability not specifically defined
    return f"""To remediate this {vuln_name} vulnerability, implement proper input validation with a strong allowlist approach that verifies all user inputs against expected formats and values. Follow the principle of least privilege throughout the application, ensuring components only have access to resources necessary for their legitimate function. Develop and maintain secure coding practices including regular code reviews and security-focused testing. Keep all software components, including frameworks and libraries, updated with security patches. Implement defense in depth with multiple layers of security controls, so that if one layer fails, others will still provide protection. Conduct regular security assessments and penetration testing to identify vulnerabilities before they can be exploited. Create a security-focused development lifecycle that incorporates security at every stage from design to deployment and maintenance."""

def normalize_filename(fname):
    """Normalize filenames for robust matching: strip, lower, keep extension."""
    return fname.strip().lower() if fname else ''

# Function to parse the uploaded Excel file for vulnerability data
def parse_vulnerabilities_excel(file_path):
    """
    Parse the uploaded Excel file and return vulnerability data.
    Supports the standardized format with headers in row 1.
    Returns (vulnerabilities_list, metadata_dict)
    """
    print(f"Parsing vulnerability data from: {file_path}")
    
    engagement_url = ""
    
    try:
        # Read Excel with header in row 1 (index 0)
        df = pd.read_excel(file_path, header=0)
        print(f"Vulnerability Excel file read successfully. Columns: {df.columns.tolist()}")
        print(f"Total rows: {len(df)}")
        
        # Clean column names by stripping whitespace
        df.columns = df.columns.str.strip()
        # Drop columns that are completely empty
        df = df.dropna(axis=1, how='all')
        
        # Remove the Tester Name column if present (ignored by generator)
        tester_cols = [col for col in df.columns if col.lower() == 'tester name']
        if tester_cols:
            df = df.drop(columns=tester_cols)
        
        vulnerabilities = []
        poc_columns = [col for col in df.columns if col.strip().upper().startswith("POC")]
        
        def clean_value(value):
            if isinstance(value, float) and pd.isna(value):
                return ""
            return str(value).strip()
        
        def value_or_empty(value):
            val = clean_value(value)
            return val if val else ""
        
        def extract_cvss_score(value):
            if value is None or (isinstance(value, float) and pd.isna(value)):
                return None
            try:
                # Value might look like "CVSS:7.5 ..." or just "7.5"
                match = re.search(r"(\d+\.\d+|\d+)", str(value))
                if match:
                    return float(match.group(1))
            except Exception:
                pass
            return None
        
        def has_valid_sr_no(value):
            text = clean_value(value)
            if not text:
                return False
            return bool(re.fullmatch(r'\d+(\.\d*)?', text))
        
        def parse_poc_value(value, default_label):
            if value is None or (isinstance(value, float) and pd.isna(value)):
                return None
            text = str(value).strip()
            if not text or text.lower() == 'nan':
                return None
            lines = [line.strip() for line in re.split(r'[\r\n]+', text) if line.strip()]
            if not lines:
                return None
            screenshot_filename = ""
            description_parts = []
            for line in lines:
                if re.search(r'\.(png|jpg|jpeg)$', line.lower()):
                    screenshot_filename = line.strip()
                else:
                    description_parts.append(line.strip())
            description = " ".join(description_parts).strip() or default_label
            screenshot_filename = screenshot_filename.strip() if screenshot_filename else ""
            return {
                'text': description,
                'screenshot': screenshot_filename or ""
            }
        
        # Helper to find column by various names (case-insensitive, partial match)
        def find_col(names_list):
            for name in names_list:
                for col in df.columns:
                    col_clean = col.lower().replace(' ', '').replace('_', '').replace('-', '').replace('/', '')
                    name_clean = name.lower().replace(' ', '').replace('_', '').replace('-', '').replace('/', '')
                    if col_clean == name_clean or col_clean.startswith(name_clean) or name_clean.startswith(col_clean):
                        return col
            return None
        
        # Map column names
        col_sr_no = find_col(['Sr no', 'Sr No', 'Sr.No', 'Sr. No.', 'S.No', 'Serial No', 'VUL ID'])
        col_observation = find_col(['Observation', 'Name of Vulnerability', 'Vulnerability Name', 'Finding', 'Title'])
        col_severity = find_col(['Severity', 'Risk(Severity)', 'Risk', 'Risk Level'])
        col_cvss = find_col(['CVSS', 'CVSS Score', 'Score'])
        col_affected_url = find_col(['Affected A', 'Affected Asset', 'IP/URL/App', 'Affected URL', 'Vulnerable URL', 'URL'])
        col_description = find_col(['Detailed Observation', 'Detailed O', 'Observation/Vulnerability', 'Audit Observation', 'Description'])
        col_recommendation = find_col(['Recommendation', 'Recomme', 'Remediation', 'Fix'])
        col_reference = find_col(['Reference', 'Reference Link', 'References'])
        col_cwe = find_col(['CVE/CWE', 'CVE CWE', 'CWE ID', 'CWE', 'CVE'])
        col_evidence = find_col(['Evidence / Proof of Concept', 'Evidence', 'Proof of Concept', 'POC'])
        
        # Find screenshot columns
        screenshot_cols = [col for col in df.columns if 'screenshot' in col.lower()]
        
        print(f"Column mapping:")
        print(f"  Sr No: {col_sr_no}")
        print(f"  Observation: {col_observation}")
        print(f"  Severity: {col_severity}")
        print(f"  CVSS: {col_cvss}")
        print(f"  Affected URL: {col_affected_url}")
        print(f"  Description: {col_description}")
        print(f"  Recommendation: {col_recommendation}")
        print(f"  Reference: {col_reference}")
        print(f"  CVE/CWE: {col_cwe}")
        print(f"  Evidence: {col_evidence}")
        print(f"  Screenshot columns: {screenshot_cols}")
        
        for index, row in df.iterrows():
            if row.dropna(how='all').empty:
                continue  # Skip completely empty rows
            try:
                # Get Sr No
                sr_raw = row.get(col_sr_no, '') if col_sr_no else ''
                sr_no_value = clean_value(sr_raw)
                if not sr_no_value:
                    sr_no_value = f"VUL-{index+1:03d}"
                
                # Get vulnerability name/observation
                vuln_name = clean_value(row.get(col_observation, '')) if col_observation else ''
                if not vuln_name:
                    print(f"Skipping row {index} due to missing vulnerability name.")
                    continue
                
                # Get CVSS score
                cvss_score = extract_cvss_score(row.get(col_cvss, '')) if col_cvss else None
                
                # Get severity
                severity_text = clean_value(row.get(col_severity, '')) if col_severity else ''
                if severity_text:
                    severity = severity_text.title()
                    severity_display = severity
                elif cvss_score is not None:
                    severity = get_severity_from_cvss(cvss_score)
                    severity_display = severity
                else:
                    severity = "Medium"
                    severity_display = "Medium"
                
                # Get affected URL
                vulnerable_url = value_or_empty(row.get(col_affected_url, '')) if col_affected_url else ''
                
                vulnerable_parameter = ''
                
                # Get description
                description = value_or_empty(row.get(col_description, '')) if col_description else ''
                
                impact = ''
                
                # Get recommendation
                remediation = value_or_empty(row.get(col_recommendation, '')) if col_recommendation else ''
                
                # Get reference
                reference_link = value_or_empty(row.get(col_reference, '')) if col_reference else ''
                
                # Get CWE/CVE
                cwe_id = value_or_empty(row.get(col_cwe, '')) if col_cwe else ''
                
                # Parse steps from Evidence column or POC columns
                steps_with_screenshots = []
                
                # First try Evidence column
                evidence_text = value_or_empty(row.get(col_evidence, '')) if col_evidence else ''
                if evidence_text:
                    # Parse "Step 1: ...", "Step 2: ..." from evidence text
                    step_pattern = re.compile(r'Step\s*(\d+)\s*[:\-\.]?\s*(.*?)(?=Step\s*\d+|$)', re.IGNORECASE | re.DOTALL)
                    step_matches = step_pattern.findall(str(evidence_text))
                    for step_num_str, step_content in step_matches:
                        content = step_content.strip()
                        if content:
                            steps_with_screenshots.append({
                                'text': content,
                                'screenshot': ''
                            })
                
                # If no steps from evidence, try POC columns
                if not steps_with_screenshots:
                    for step_idx, poc_col in enumerate(poc_columns, 1):
                        poc_entry = parse_poc_value(row.get(poc_col, ''), "")
                        if poc_entry:
                            steps_with_screenshots.append(poc_entry)
                
                # Get screenshots from screenshot columns
                for idx, screenshot_col in enumerate(screenshot_cols):
                    screenshot_val = clean_value(row.get(screenshot_col, ''))
                    if screenshot_val and idx < len(steps_with_screenshots):
                        steps_with_screenshots[idx]['screenshot'] = screenshot_val
                    elif screenshot_val:
                        steps_with_screenshots.append({
                            'text': f'Step {len(steps_with_screenshots)+1}',
                            'screenshot': screenshot_val
                        })
                
                vulnerability = {
                    'name': vuln_name,
                    'description': description,
                    'impact': impact,
                    'severity': severity,
                    'severity_display': severity_display,
                    'cvss': cvss_score,
                    'vulnerable_url': vulnerable_url,
                    'vulnerable_parameter': vulnerable_parameter,
                    'remediation': remediation,
                    'steps_with_screenshots': steps_with_screenshots,
                    'reference_link': reference_link,
                    'cwe_id': cwe_id or "",
                    'sr_no': sr_no_value.replace('VULN-', 'VUL-'),
                }
                
                vulnerabilities.append(vulnerability)
                
                print(f"Parsed vulnerability Sr No: {vulnerability['sr_no']} at index {index}") # Debug print
                
            except Exception as e:
                print(f"Error processing row {index}: {e}")
                continue
        
        print(f"Total vulnerabilities extracted: {len(vulnerabilities)}")
        
        metadata = {'engagement_url': engagement_url}
        return vulnerabilities, metadata
        
    except Exception as e:
        print(f"Error parsing vulnerability Excel: {e}")
        return [], {'engagement_url': engagement_url}

# Function to parse the fixed document control file
def parse_doc_control_excel(file_path):
    """
    Parse the fixed document control Excel file and return document control details
    including assessment scope, dates, note, and amendment log.
    """
    print(f"Parsing document control data from: {file_path}")
    doc_control_data = {
        'assessment_scope': '',
        'assessment_start_date': '',
        'assessment_end_date': '',
        'assessment_note': '',  # Add assessment_note
        'amendment_log': [],
        'reviewed_by': '',
        'authorized_by': ''
    }
    try:
        df = pd.read_excel(file_path)
        print(f"Document Control Excel file read successfully. Columns: {df.columns.tolist()}")

        # Clean column names by stripping whitespace
        df.columns = df.columns.str.strip()
        
        # Find the header row containing "Version", "Date", "Reviewed By", "Brief description of the change"
        header_row_index = None
        for index, row in df.iterrows():
            # Check if key headers are present in this row
            if all(col in row.index for col in ["Version", "Date", "Reviewed By", "Brief description of the change"]):
                header_row_index = index
                print(f"Found document control header row at index: {header_row_index}")
                break
        
        if header_row_index is None:
            print("Document control header row not found. Cannot extract document control details or amendment log.")
            return doc_control_data # Return empty data

        # Assuming assessment details are in the row *after* the header row (first data row)
        first_data_row_index = header_row_index + 1
        if first_data_row_index < len(df):
            first_data_row = df.iloc[first_data_row_index]
            
            # Handle Assessment Scope
            if 'Assessment Scope' in first_data_row:
                doc_control_data['assessment_scope'] = str(first_data_row['Assessment Scope']).strip()
                print(f"Found Assessment Scope: {doc_control_data['assessment_scope']}")
            # Handle Assessment Note
            if 'Assessment Note' in first_data_row:
                doc_control_data['assessment_note'] = str(first_data_row['Assessment Note']).strip()
                print(f"Found Assessment Note: {doc_control_data['assessment_note']}")
            # Handle Assessment Start Date
            if 'Assessment Start Date' in first_data_row:
                start_date = first_data_row['Assessment Start Date']
                print(f"Raw Start Date (Doc Control): {start_date}, Type: {type(start_date)}")
                if isinstance(start_date, pd.Timestamp):
                    doc_control_data['assessment_start_date'] = start_date.strftime('%dth %B %Y')
                else:
                    # Try to parse the date string if it's in the format "10th March 2025"
                    try:
                         # Remove 'th', 'st', 'nd', 'rd' from the day
                        date_str = str(start_date).strip()
                        date_str = date_str.replace('th', '').replace('st', '').replace('nd', '').replace('rd', '')
                        # Parse the date
                        parsed_date = pd.to_datetime(date_str, format='%d %B %Y')
                        doc_control_data['assessment_start_date'] = parsed_date.strftime('%dth %B %Y')
                    except:
                        doc_control_data['assessment_start_date'] = str(start_date).strip()
                print(f"Formatted Start Date (Doc Control): {doc_control_data['assessment_start_date']}")

            # Handle Assessment End Date
            if 'Assessment End Date' in first_data_row:
                end_date = first_data_row['Assessment End Date']
                print(f"Raw End Date (Doc Control): {end_date}, Type: {type(end_date)}")
                if isinstance(end_date, pd.Timestamp):
                    doc_control_data['assessment_end_date'] = end_date.strftime('%dth %B %Y')
                else:
                     # Try to parse the date string if it's in the format "30th March 2025"
                    try:
                         # Remove 'th', 'st', 'nd', 'rd' from the day
                        date_str = str(end_date).strip()
                        date_str = date_str.replace('th', '').replace('st', '').replace('nd', '').replace('rd', '')
                        # Parse the date
                        parsed_date = pd.to_datetime(date_str, format='%d %B %Y')
                        doc_control_data['assessment_end_date'] = parsed_date.strftime('%dth %B %Y')
                    except:
                        doc_control_data['assessment_end_date'] = str(end_date).strip()
                print(f"Formatted End Date (Doc Control): {doc_control_data['assessment_end_date']}")
            
            # Extract "Reviewed By" and "Authorized By" from the first data row
            if first_data_row_index < len(df):
                first_data_row = df.iloc[first_data_row_index]
                doc_control_data['reviewed_by'] = str(first_data_row.get('Reviewed By', '')).strip()
                doc_control_data['authorized_by'] = str(first_data_row.get('Authorized By', '')).strip()
                print(f"Found Reviewed By: {doc_control_data['reviewed_by']}")
                print(f"Found Authorized By: {doc_control_data['authorized_by']}")

            print(f"Final assessment details from doc control:")
            print(f"Scope: {doc_control_data['assessment_scope']}")
            print(f"Start Date: {doc_control_data['assessment_start_date']}")
            print(f"End Date: {doc_control_data['assessment_end_date']}")
            print(f"Note: {doc_control_data['assessment_note']}")

        # Extract amendment log entries from rows below the header row
        amendment_log_start_index = header_row_index + 1
        for index in range(amendment_log_start_index, len(df)):
            row = df.iloc[index]
            # Stop if the "Version" cell is empty/NaN
            if pd.isna(row.get('Version', None)) or str(row.get('Version', '')).strip() == '':
                break # Stop reading amendment log
            
            # Handle Date formatting to remove time
            log_date = row.get('Date', '')
            formatted_date = str(log_date).strip()
            if isinstance(log_date, pd.Timestamp):
                formatted_date = log_date.strftime('%Y-%m-%d') # Format as YYYY-MM-DD
            
            log_entry = {
                'Version': str(row.get('Version', '')).strip(),
                'Date': formatted_date,
                'Reviewed By': str(row.get('Reviewed By', '')).strip(),
                'Brief description of the change': str(row.get('Brief description of the change', '')).strip()
            }
            doc_control_data['amendment_log'].append(log_entry)
            print(f"Added amendment log entry: {log_entry}")

        print(f"Total amendment log entries extracted: {len(doc_control_data['amendment_log'])}")
        
        return doc_control_data

    except FileNotFoundError:
        print(f"Error: Document control file not found at {file_path}")
        # Continue without document control data if file is missing
        return doc_control_data
    except Exception as e:
        print(f"Error parsing document control Excel file: {e}")
        # Continue without document control data if parsing fails
        return doc_control_data

def create_summary_table(doc, vulnerabilities):
    """
    Creates and inserts a color-coded vulnerability summary table.
    This function replaces a placeholder in the document.
    """
    placeholder_text = '_VULNERABILITY_SUMMARY_TABLE_PLACEHOLDER_'
    
    # Find the placeholder paragraph
    placeholder_paragraph = None
    for p in doc.paragraphs:
        if placeholder_text in p.text:
            placeholder_paragraph = p
            break
            
    if not placeholder_paragraph:
        print("Warning: Placeholder '_VULNERABILITY_SUMMARY_TABLE_PLACEHOLDER_' not found in the template.")
        return

    # Headers and table creation
    headers = ['Sl.no', 'Vulnerability Name', 'Vulnerability Risk Type']
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Set column widths for better layout
    table.columns[0].width = Inches(0.5)
    table.columns[1].width = Inches(4.0)
    table.columns[2].width = Inches(1.5)

    # Style header row
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        cell = hdr_cells[i]
        # Clear the default paragraph and add new text
        p = cell.paragraphs[0]
        p.text = header_text
        run = p.runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255) # White font
        # Set header background color (purple)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="6A449A"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)

    # Add data rows with color coding
    for idx, vuln in enumerate(vulnerabilities, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = vuln.get('name', '')
        
        severity_logic = vuln.get('severity', 'Unknown')
        severity_display = vuln.get('severity_display', severity_logic)
        severity_cell = row_cells[2]
        severity_cell.text = severity_display
        
        # Get severity colors and apply them
        colors = get_severity_colors(severity_logic)
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"].lstrip("#")}"/>')
        severity_cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Set font to white for dark backgrounds for better readability
        if severity_logic in ["Critical", "High"]:
            p = severity_cell.paragraphs[0]
            if p.runs:
                run = p.runs[0]
                run.font.color.rgb = RGBColor(255, 255, 255)

    # Replace placeholder paragraph with the new table
    parent = placeholder_paragraph._p.getparent()
    parent.insert(parent.index(placeholder_paragraph._p), table._element)
    parent.remove(placeholder_paragraph._p)

def set_cell_text_preserve_format(cell, text):
    """
    Update cell text while retaining existing paragraph/run formatting as much as possible.
    """
    if not cell.paragraphs:
        cell.text = text
        return
    first_para = cell.paragraphs[0]
    if not first_para.runs:
        first_para.add_run(text)
    else:
        first_para.runs[0].text = text
        for run in first_para.runs[1:]:
            run.text = ''
    for para in cell.paragraphs[1:]:
        for run in para.runs:
            run.text = ''

def update_risk_assessment_table(doc, severity_counts):
    """
    Locate the 'Risk Assessment Analysis of the Entire Framework' table
    and update severity/total counts dynamically without altering formatting.
    """
    severity_order_display = ["Critical", "High", "Medium", "Low", "Informational"]
    required_headers = {"sr. no.", "hostname", "instant purpose", "vapt status",
                        "critical", "high", "medium", "low", "informational", "total"}
    total_vulns = sum(severity_counts.values())
    
    for table in doc.tables:
        if not table.rows:
            continue
        header_texts = [cell.text.strip().lower() for cell in table.rows[0].cells]
        header_set = set(header_texts)
        if not required_headers.issubset(header_set):
            continue
        header_index = {}
        for idx, text in enumerate(header_texts):
            if text and text not in header_index:
                header_index[text] = idx
        severity_indices = {}
        for sev in severity_order_display:
            idx = header_index.get(sev.lower())
            if idx is None:
                break
            severity_indices[sev] = idx
        else:
            total_idx = header_index.get("total")
            for row in table.rows[1:]:
                for sev in severity_order_display:
                    count = severity_counts.get(sev, 0)
                    set_cell_text_preserve_format(row.cells[severity_indices[sev]], str(count))
                if total_idx is not None:
                    set_cell_text_preserve_format(row.cells[total_idx], str(total_vulns))
            return True
    print("Warning: Risk assessment table with severity columns not found.")
    return False

def remove_trailing_empty_paragraphs(doc):
    """
    Remove trailing empty paragraphs at the end of the document to avoid extra spacing
    before inserting dynamic content such as vulnerability tables.
    """
    removed = 0
    while doc.paragraphs:
        last_para = doc.paragraphs[-1]
        if last_para.text.strip():
            break
        parent = last_para._element.getparent()
        if parent is None:
            break
        parent.remove(last_para._element)
        removed += 1
    if removed:
        print(f"Removed {removed} trailing empty paragraphs before inserting vulnerabilities.")

# Function to create a styled vulnerability table in Word document
def create_vulnerability_table(doc, vulnerability, display_sr_no=None, image_map=None):
    print(f"Creating table for vulnerability: {vulnerability['name']}, Severity: {vulnerability['severity']}")  # Debug output
    
    # Row titles in exact order requested
    row_labels = [
        "",  # Row 1: Vulnerability Name (handled specially)
        "",  # Row 2: Vulnerable URL (handled specially)
        "Vulnerable Parameter",  # Row 3
        "CWE ID",  # Row 4
        "CVSS Score",  # Row 4
        "Severity",  # Row 5
        "Vulnerability Description",  # Row 6
        "Vulnerability Impact",  # Row 7
        "Recommendation",  # Row 8
        "Reference Link",  # Row 9
        "Proof of Concept / Steps to Reproduce"  # Row 9
    ]
    
    # Create a table with the appropriate number of rows
    # We need 9 rows in total as per the specified order
    table = doc.add_table(rows=len(row_labels), cols=1)
    table.style = 'Table Grid'  # Add borders
    
    # Get severity and colors
    severity = vulnerability["severity"]
    colors = get_severity_colors(severity)
    
    print(f"Using colors: {colors}")  # Debug output
    
    # Extract RGB components from hex colors
    row1_bg_hex = colors["row1_bg"].lstrip('#')
    row2_bg_hex = colors["row2_bg"].lstrip('#')
    font_hex = colors["font"].lstrip('#')
    
    row1_r, row1_g, row1_b = tuple(int(row1_bg_hex[i:i+2], 16) for i in (0, 2, 4))
    row2_r, row2_g, row2_b = tuple(int(row2_bg_hex[i:i+2], 16) for i in (0, 2, 4))
    font_r, font_g, font_b = tuple(int(font_hex[i:i+2], 16) for i in (0, 2, 4))
    
    # Row 1: Vulnerability Name
    row = table.rows[0]
    cell = row.cells[0]
    # Apply background color based on severity
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row1_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add text with formatting - use the display_sr_no if provided
    sr_display = display_sr_no if display_sr_no else vulnerability.get('sr_no', '')
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(f"{sr_display}: {vulnerability['name']}")
    run.font.name = 'Altone Trial'  # Change to Altone Trial font
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Row 2: Vulnerable URL
    row = table.rows[1]
    cell = row.cells[0]
    # Apply background color based on severity
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{colors["row2_bg"]}"/>')
    cell._element.get_or_add_tcPr().append(shading_elm)
    
    # Add text with formatting
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Vulnerable URL: ")
    run.font.name = 'Altone Trial'  # Change to Altone Trial font
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    run = paragraph.add_run(vulnerability.get('vulnerable_url', ""))
    run.font.name = 'Altone Trial'  # Change to Altone Trial font
    run.font.size = Pt(12)
    run.font.bold = False
    run.font.color.rgb = RGBColor(255, 255, 255)  # White text
    
    # Rows 3-9: Content rows with white background and proper formatting
    # Prepare content for each row - ENSURE ALL VALUES ARE STRINGS
    vuln_param = vulnerability.get('vulnerable_parameter', "")
        
    row_contents = [
        str(vuln_param),  # Vulnerable Parameter
        str(vulnerability.get('cwe_id', "")),  # CWE ID
        str(vulnerability['cvss']) if vulnerability['cvss'] is not None else "",  # CVSS Score
        str(vulnerability.get('severity_display', vulnerability.get('severity', ""))) or "",  # Severity
        str(vulnerability.get('description', "")),  # Vulnerability Description
        str(vulnerability.get('impact', "")),  # Vulnerability Impact
        str(vulnerability.get('remediation', "")),  # Recommendation
        str(vulnerability.get('reference_link', "")) or "",  # Reference Link
        ""  # Steps will be handled specially
    ]
    
    # Process content rows with white background and proper formatting
    for i in range(2, len(row_labels)):  # Start from row index 2 (third row)
        row_idx = i  # Already 0-indexed
        row = table.rows[row_idx]
        cell = row.cells[0]
        
        # Apply white background
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="#FFFFFF"/>')
        cell._element.get_or_add_tcPr().append(shading_elm)
        
        # Add text with formatting
        paragraph = cell.paragraphs[0]
        
        # Label with color based on severity
        run = paragraph.add_run(f"{row_labels[i]}: ")
        run.font.name = 'Altone Trial'  # Change to Altone Trial font
        run.font.size = Pt(12)
        run.font.bold = True
        # Set font color using the RGB values from the severity color
        run.font.color.rgb = RGBColor(font_r, font_g, font_b)
        
        # Content in normal formatting (except for steps which we'll handle specially)
        if i < len(row_labels) - 1:  # Not the steps row
            run = paragraph.add_run(row_contents[i-2])
            run.font.name = 'Altone Trial'  # Change to Altone Trial font
            run.font.size = Pt(11)
            run.font.bold = False
        else:  # This is the Steps row - handle steps with screenshots
            steps = vulnerability.get('steps_with_screenshots', [])
            if not steps:
                run = paragraph.add_run("")
                run.font.name = 'Altone Trial'
                run.font.size = Pt(11)
            else:
                for step_idx, step in enumerate(steps):
                    step_para = cell.add_paragraph()
                    step_para.paragraph_format.left_indent = Pt(10)
                    step_text = step['text']
                    run = step_para.add_run(f"{step_text}")
                    run.font.name = 'Altone Trial'
                    run.font.size = Pt(11)
                    run.font.bold = False
                    screenshot_path = None
                    if image_map and step.get('screenshot'):
                        screenshot_path = find_image_for_step(step['screenshot'], image_map, display_sr_no, step_idx+1)
                    if screenshot_path and os.path.exists(screenshot_path):
                        print(f"[DEBUG] Adding image from ZIP: {screenshot_path}")
                        screenshot_para = cell.add_paragraph()
                        screenshot_para.paragraph_format.left_indent = Pt(20)
                        screenshot_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        try:
                            screenshot_para.add_run().add_picture(screenshot_path, width=Inches(5))
                        except Exception as img_err:
                            print(f"[ERROR] Failed to insert screenshot {screenshot_path}: {img_err}")
                            missing_para = cell.add_paragraph()
                            missing_para.paragraph_format.left_indent = Pt(20)
                            missing_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = missing_para.add_run(f"[Screenshot unavailable: {step.get('screenshot')}]")
                            run.font.name = 'Altone Trial'
                            run.font.size = Pt(10)
                            run.font.italic = True
                    else:
                        print(f"  -> No image found for step {step_idx+1}, inserting missing message.")
                        missing_para = cell.add_paragraph()
                        missing_para.paragraph_format.left_indent = Pt(20)
                        missing_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = missing_para.add_run(f"[Screenshot missing: {step.get('screenshot')}]")
                        run.font.name = 'Altone Trial'
                        run.font.size = Pt(10)
                        run.font.italic = True
                    if step_idx < len(steps) - 1:
                        spacer = cell.add_paragraph()
                        spacer.paragraph_format.space_after = Pt(6)
    
    return table

def create_vulnerability_table_with_poc(doc, vulnerability, display_sr_no=None, image_map=None):
    # Create the main vulnerability table as before
    table = create_vulnerability_table(doc, vulnerability, display_sr_no, image_map)
    # After the table, add PoC steps and images
    para = doc.add_paragraph(f"Proof of Concept / Steps to Reproduce for {vulnerability['name']}:")
    para.runs[0].bold = True
    para.runs[0].font.size = Pt(14)
    steps = vulnerability.get('steps_with_screenshots', [])
    for step_idx, step in enumerate(steps, 1):
        step_para = doc.add_paragraph(f"Step {step_idx}: {step['text']}")
        step_para.paragraph_format.left_indent = Pt(10)
        screenshot_path = find_image_for_step(step['screenshot'], image_map, display_sr_no, step_idx)
        if screenshot_path and os.path.exists(screenshot_path):
            print(f"[DEBUG] Adding image from ZIP: {screenshot_path}")
            img_run = doc.add_paragraph().add_run()
            try:
                img_run.add_picture(screenshot_path, width=Inches(4))
            except Exception as img_err:
                print(f"[ERROR] Failed to insert screenshot {screenshot_path}: {img_err}")
                doc.add_paragraph(f"[Screenshot unavailable: {step['screenshot']}]")
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]")
    return table

def calculate_optimal_logo_size(image_path, max_width=2.3, max_height=2.0):
    """
    Calculate the optimal logo size to fit within the template while maintaining aspect ratio.
    
    Args:
        image_path (str): Path to the logo image
        max_width (float): Maximum width in inches (increase this value if logos appear too small)
        max_height (float): Maximum height in inches (increase this value for taller logos)
        
    Returns:
        float: Optimal width in inches
        
    Note:
        - The default parameters (2.5 x 1.2 inches) are suitable for most company logos
        - For wider logos, increase max_width
        - For taller logos, increase max_height
        - The minimum width is set to 1.5 inches to ensure logos are clearly visible
    """
    # Check if PIL is available
    if not PIL_AVAILABLE:
        print("PIL not available, using default logo width")
        return max_width * 0.9  # Use a slightly smaller width as a safe default
    
    try:
        # Open the image to get its dimensions
        with Image.open(image_path) as img:
            img_width, img_height = img.size
            
        # Convert pixels to inches (assuming 96 DPI which is standard for many screens)
        # This is a rough estimate, as exact DPI may vary based on the image
        width_in_inches = img_width / 96
        height_in_inches = img_height / 96
            
        print(f"Original logo dimensions: {img_width}x{img_height}px ({width_in_inches:.2f}x{height_in_inches:.2f} inches)")
        
        # If the image is already within our size limits, use its actual size,
        # but ensure it's not too small (at least 1.5 inches wide for good visibility)
        if width_in_inches <= max_width and height_in_inches <= max_height:
            if width_in_inches < 1.5:
                print(f"Logo is too small, scaling up to minimum width: 1.5 inches")
                return 1.5
            print(f"Logo is already within size limits, using original size: {width_in_inches:.2f} inches wide")
            return width_in_inches
            
        # Calculate aspect ratio
        aspect_ratio = img_width / img_height
        
        # Calculate potential width and height based on max dimensions
        width_from_height = max_height * aspect_ratio
        height_from_width = max_width / aspect_ratio
        
        # Determine which dimension is the limiting factor
        if width_from_height <= max_width:
            # Height is the limiting factor
            optimal_width = width_from_height
        else:
            # Width is the limiting factor
            optimal_width = max_width
            
        # Ensure minimum size of 1.5 inches width
        if optimal_width < 1.5:
            optimal_width = 1.5
            
        print(f"Resizing logo to: {optimal_width:.2f} inches wide")
        print(f"Logo display size will be approximately {optimal_width:.2f} x {optimal_width/aspect_ratio:.2f} inches")
        return optimal_width
    except Exception as e:
        print(f"Error calculating optimal logo size: {e}")
        return max_width  # Default to max_width if there's an error

# Updated function signature to accept amendment_log and assessment details
def generate_word_report(
    vulnerabilities,
    doc_control_data,
    template_path,
    output_path,
    client_name="Client",
    image_map=None,
    logo_path=None,
    engagement_url=""
):
    try:
        tpl = DocxTemplate(template_path)
        # Generate document ID in the format C-XX-WEB-Q1-2025
        current_year = datetime.now().year
        current_quarter = (datetime.now().month - 1) // 3 + 1
        client_code = client_name.upper()[:2]  # Get first two letters of client name
        doc_id = f"C-{client_code}-WEB-Q{current_quarter}-{current_year}"
        # Get assessment details and amendment log from parsed doc_control_data
        assessment_scope = doc_control_data.get('assessment_scope', '')
        assessment_start_date = doc_control_data.get('assessment_start_date', '')
        assessment_end_date = doc_control_data.get('assessment_end_date', '')
        assessment_note = doc_control_data.get('assessment_note', '')
        amendment_log = doc_control_data.get('amendment_log', [])
        # Format the date range string
        date_range = f"{assessment_start_date} to {assessment_end_date}"
        print("Template variables being passed:")
        print(f"assessment_scope: {assessment_scope}")
        print(f"date_range: {date_range}")
        print(f"document_id: {doc_id}")
        print(f"amendment_log (count): {len(amendment_log)}") # Debug print
        
        # Handle the company logo
        company_logo = None
        default_logo_path = os.path.join(os.path.dirname(__file__), "default_logo.png")
        
        try:
            if logo_path and os.path.exists(logo_path):
                print(f"Using custom company logo: {logo_path}")
                optimal_width = calculate_optimal_logo_size(logo_path)
                company_logo = InlineImage(tpl, logo_path, width=Inches(optimal_width))
            elif os.path.exists(default_logo_path):
                print(f"Using default company logo: {default_logo_path}")
                optimal_width = calculate_optimal_logo_size(default_logo_path)
                company_logo = InlineImage(tpl, default_logo_path, width=Inches(optimal_width))
            else:
                print("No valid logo found, logo will be blank in the report")
        except Exception as e:
            print(f"Error processing logo: {e}. Continuing without logo.")
            company_logo = None
            
        # Build steps for each vulnerability (text only, not InlineImage)
        for vuln in vulnerabilities:
            vuln['steps'] = [
                {'text': step['text'], 'image': f"[Screenshot missing: {step['screenshot']}]"}
                for step in vuln['steps_with_screenshots']
            ]
        context = {
            'client_name': client_name,
            'assessment_scope': assessment_scope,  # Exact match for {{assessment_scope}}
            'date_range': date_range,             # Exact match for {{date_range}}
            'document_id': doc_id,                # Add document ID to context
            'amendment_log': amendment_log,       # Add amendment log to context
            'reviewed_by': doc_control_data.get('reviewed_by', ''), # Add reviewed_by to context
            'authorized_by': doc_control_data.get('authorized_by', ''), # Add authorized_by to context
            'report_date': datetime.now().strftime('%d %B %Y'),  # Add current date for {{report_date}}
            'assessment_note': assessment_note,  # Add assessment_note to context
            'vulnerabilities': vulnerabilities,
            'company_logo': company_logo,         # Add company logo to context
            'engagement_url': engagement_url,
        }
        print("Full template context:")
        print(context)
        # Render template with all variables (static content only)
        tpl.render(context)
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tpl.save(tmp.name)
            tmp_path = tmp.name
        # Open with python-docx for dynamic PoC steps/images
        doc = Document(tmp_path)
        # Update the risk assessment table counts
        severity_counts = count_vulnerabilities_by_severity(vulnerabilities)
        update_risk_assessment_table(doc, severity_counts)
        # Sort vulnerabilities by severity before creating any tables
        vulnerabilities.sort(key=lambda x: float(x.get('cvss') or 0), reverse=True)
        # Programmatically create and insert the summary table
        create_summary_table(doc, vulnerabilities)
        # Attempt to set default font for the document body
        try:
            style = doc.styles['Normal']
            style.font.name = 'Altone Trial'
        except KeyError:
            print("Warning: 'Normal' style not found in template. Cannot set default font.")
        except Exception as e:
            print(f"Warning: Could not set default font on 'Normal' style: {e}")
        # Apply Altone Trial font to all paragraphs and runs in the document body (additional pass)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = 'Altone Trial'
        for table in doc.tables:
             for row in table.rows:
                 for cell in row.cells:
                     for paragraph in cell.paragraphs:
                         for run in paragraph.runs:
                             run.font.name = 'Altone Trial'
        # Assign display_sr_no to each vulnerability BEFORE generating tables/PoC steps
        for idx, vuln in enumerate(vulnerabilities, 1):
            vuln['display_sr_no'] = f"VUL-{idx:03d}"
        # Insert vulnerability table and PoC steps/images for each vulnerability using python-docx only
        remove_trailing_empty_paragraphs(doc)
        if vulnerabilities:
            doc.add_page_break()
        for idx, vuln in enumerate(vulnerabilities):
            create_vulnerability_table(doc, vuln, vuln['display_sr_no'], image_map)
            if idx < len(vulnerabilities) - 1:
                doc.add_page_break()
        # Add a page break and the SUMMARY OF FINDINGS & CONCLUSION section as the last page
        doc.add_page_break()
        # Add two blank lines after the SUMMARY OF FINDINGS & CONCLUSION: heading
        doc.add_paragraph()
        doc.add_paragraph()
        # Heading as a single paragraph with a line break for TOC compatibility
        summary_heading = doc.add_paragraph()
        run = summary_heading.add_run("SUMMARY OF FINDINGS &\nCONCLUSION:")
        doc.add_paragraph()
        run.bold = True
        run.font.size = Pt(28)
        run.font.name = 'Altone Trial'
        run.font.color.rgb = RGBColor(106, 68, 154)
        summary_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Set style to Heading 1 if available for TOC
        try:
            summary_heading.style = doc.styles['Heading 1']
        except Exception:
            pass
        # Paragraph 1
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p1.paragraph_format.space_after = Pt(24)  # Double spacing after
        run = p1.add_run(
            "Finally, it must be remembered that security is an ongoing process, and that this report will provide an idea of the current vulnerabilities we were able to detect. There is no guarantee that new vulnerabilities will not be found and exploited in the future."
        )
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Paragraph 2
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.space_after = Pt(24)
        run = p2.add_run("The assessment was only possible because ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("joint support & coordination")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" from the information security team of organization for ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("sharing & coordinating")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" during the assessment period. It is advised to refer the ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("Technical Report")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run(" for understanding in-depth of vulnerabilities that were discovered by technical team of ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p2.add_run("CyberSmithSECURE Pvt. Ltd.")
        run.bold = True
        run.underline = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Paragraph 3
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p3.paragraph_format.space_after = Pt(24)
        run = p3.add_run("The Security Researchers of the ")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p3.add_run("CyberSmithSECURE")
        run.bold = True
        run.underline = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run = p3.add_run(" performed Vulnerability Testing. We jointly recommend that all suggested measures in this document be performed to ensure the overall security of the target device. The following targeted sectors were identified by the security researchers for the scope of this testing.")
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
        # Paragraph 4 (last line, bold and purple)
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p4.paragraph_format.space_after = Pt(24)
        run = p4.add_run("We thank internal Information Security team for their support & cooperation during the time of assessment.")
        run.bold = True
        run.font.name = 'Altone Trial'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(106, 68, 154)
        doc.save(output_path)
        return output_path
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error generating Word report: {str(e)}")

# Function to check if file exists in uploads directory
def check_screenshots_directory(vulnerability, screenshots_dir="uploads/screenshots"):
    """
    Check if screenshot directory exists for a vulnerability, create it if not
    Returns the path to the screenshots directory
    """
    try:
        # Create directory structure based on severity and vulnerability name
        severity = vulnerability.get('severity', 'Unknown')
        vuln_name = vulnerability.get('name', 'Unknown').replace(' ', '_').replace('/', '_')
        
        # Create path: uploads/screenshots/[Severity]/[Vulnerability_Name]
        screenshot_path = os.path.join(screenshots_dir, severity, vuln_name)
        os.makedirs(screenshot_path, exist_ok=True)
        
        return screenshot_path
    except Exception as e:
        print(f"Error creating screenshots directory: {e}")
        return None

def find_image_case_insensitive(directory, filename):
    """
    Search for an image file in the directory, case-insensitive, and try common extensions.
    """
    base, ext = os.path.splitext(filename)
    possible_exts = ['.png', '.jpg', '.jpeg', '.PNG', '.JPG', '.JPEG']
    candidates = [base + e for e in possible_exts]
    files = os.listdir(directory)
    for candidate in candidates:
        for f in files:
            if f.lower() == candidate.lower():
                return os.path.join(directory, f)
    return None

def index_images_recursively(root_dir):
    image_map = {}
    for dirpath, _, filenames in os.walk(root_dir):
        for fname in filenames:
            if fname.lower().endswith((".png", ".jpg", ".jpeg")):
                norm = normalize_filename(fname)
                full_path = os.path.join(dirpath, fname)
                with open(full_path, 'rb') as f:
                    data = f.read()
                    h = hashlib.md5(data).hexdigest()
                    print(f"[IMAGE MAP] {fname} => {norm} | Size: {len(data)} | Hash: {h}")
                    f.seek(0)
                image_map[norm] = full_path
    print(f"[DEBUG] Final image map: {list(image_map.keys())}")
    return image_map

def insert_steps_with_images(doc, vulnerability, image_map):
    doc.add_paragraph(f"Proof of Concept / Steps to Reproduce for {vulnerability['name']}").runs[0].bold = True
    for idx, step in enumerate(vulnerability.get("steps_with_screenshots", []), 1):
        doc.add_paragraph(f"Step {idx}: {step['text']}")
        norm = normalize_filename(step['screenshot'])
        img_path = image_map.get(norm)
        print(f"[PoC] Step {idx} for '{vulnerability['name']}' → Looking for '{norm}' → {img_path}")
        if img_path and os.path.exists(img_path):
            try:
                with open(img_path, 'rb') as f:
                    img_data = f.read()
                    print(f"[PoC] FOUND: {img_path} | Hash: {hashlib.md5(img_data).hexdigest()} | Size: {len(img_data)}")
                    f.seek(0)
                    doc.add_paragraph().add_run().add_picture(f, width=Inches(4))
            except Exception as img_err:
                print(f"[ERROR] Failed to insert screenshot {img_path}: {img_err}")
                doc.add_paragraph(f"[Screenshot unavailable: {step['screenshot']}]")
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]")

@app.post("/generate-report/")
async def generate_report(
    file: UploadFile = File(...),  # Vulnerability file
    assessment_file: UploadFile = File(...),  # Assessment details file
    poc_images: UploadFile = File(None),  # POC screenshots ZIP file (optional)
    company_logo: UploadFile = File(None),  # Company logo file (optional)
    client_name: str = Form("Client")
):
    try:
        print(f"Starting report generation - Version {VERSION}")
        
        # Create directories if they don't exist
        if not os.path.exists(UPLOAD_DIR):
            os.makedirs(UPLOAD_DIR)
        
        screenshots_dir = os.path.join(UPLOAD_DIR, "screenshots")
        if not os.path.exists(screenshots_dir):
            os.makedirs(screenshots_dir)
        
        # Create logos directory if it doesn't exist
        logos_dir = os.path.join(UPLOAD_DIR, "logos")
        if not os.path.exists(logos_dir):
            os.makedirs(logos_dir)
            
        # Save the uploaded vulnerability file
        vulnerability_file_location = os.path.join(UPLOAD_DIR, file.filename)
        with open(vulnerability_file_location, "wb") as f:
            f.write(await file.read())
        print(f"Saved uploaded vulnerability file to: {vulnerability_file_location}")

        # Save the uploaded assessment file
        assessment_file_location = os.path.join(UPLOAD_DIR, assessment_file.filename)
        with open(assessment_file_location, "wb") as f:
            f.write(await assessment_file.read())
        print(f"Saved uploaded assessment file to: {assessment_file_location}")
        
        # Handle company logo if provided
        logo_path = None
        if company_logo is not None:
            logo_filename = f"{client_name.replace(' ', '_')}_logo{os.path.splitext(company_logo.filename)[-1]}"
            logo_path = os.path.join(logos_dir, logo_filename)
            with open(logo_path, "wb") as f:
                f.write(await company_logo.read())
            print(f"Saved company logo to: {logo_path}")

        # Handle POC images ZIP
        poc_images_dir = None
        if poc_images is not None:
            tmpdirname = tempfile.mkdtemp()
            zip_path = os.path.join(tmpdirname, poc_images.filename)
            with open(zip_path, 'wb') as f:
                f.write(await poc_images.read())
            with ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(tmpdirname)
            # Debug: Print the extracted directory tree
            print('--- Extracted ZIP Directory Tree ---')
            for dirpath, dirnames, filenames in os.walk(tmpdirname):
                print(f"DIR: {dirpath}")
                print(f"  DIRNAMES: {dirnames}")
                print(f"  FILENAMES: {filenames}")
            print('--- End of Directory Tree ---')
            # Find the first folder inside the extracted directory
            for entry in os.listdir(tmpdirname):
                entry_path = os.path.join(tmpdirname, entry)
                if os.path.isdir(entry_path):
                    poc_images_dir = entry_path
                    break
            # If no folder found, fallback to tmpdirname
            if poc_images_dir is None:
                poc_images_dir = tmpdirname

        # Parse the uploaded vulnerability file (data + metadata like engagement URL)
        vulnerabilities, vuln_metadata = parse_vulnerabilities_excel(vulnerability_file_location)
        
        # Parse the assessment file for document control data
        doc_control_data = parse_doc_control_excel(assessment_file_location)
        
        # Use the CSS template v2.1 - use the template in the same directory as type4.py
        template_path = os.path.join(os.path.dirname(__file__), "CSS-NEW_Technical_Sample_Report_Automated_v2.1.docx")
        report_filename = f"{client_name} VAPT Report {datetime.now().strftime('%Y-%d-%m')}.docx"
        output_path = os.path.join(UPLOAD_DIR, report_filename)
        
        # Generate the report, passing both vulnerability and document control data
        image_map = None
        if poc_images_dir:
            image_map = index_images_from_poc_zip(poc_images_dir)
            if image_map:
                print('Available images in ZIP:')
                for k in image_map.keys():
                    print('  ', k)
        engagement_url = vuln_metadata.get('engagement_url', '')
        generate_word_report(
            vulnerabilities,
            doc_control_data,
            template_path,
            output_path,
            client_name,
            image_map,
            logo_path,
            engagement_url=engagement_url
        )
        
        print(f"Report generation completed successfully - Version {VERSION}")
        
        # Return the report
        return FileResponse(
            path=output_path,
            filename=report_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

def count_vulnerabilities_by_severity(vulnerabilities):
    """
    Counts the number of vulnerabilities for each severity level.
    """
    severity_counts = Counter()
    for vuln in vulnerabilities:
        severity_counts[vuln.get('severity', 'Unknown')] += 1
        
    return severity_counts

def create_severity_bar_chart(severity_counts, output_path):
    """
    Generates a bar chart showing the count of vulnerabilities by severity.
    Saves the chart as an image file.
    """
    severities = list(severity_counts.keys())
    counts = list(severity_counts.values())
    
    # Define colors matching the order (approximate hex values)
    colors = [
        '#990000' if s == 'Critical' else
        '#FF0000' if s == 'High' else
        '#FFCC00' if s == 'Medium' else
        '#009933' if s == 'Low' else
        '#3399CC' # Informational/Unknown
        for s in severities
    ]
    
    # Sort data by severity order for consistency display
    sorted_severities, sorted_counts, sorted_colors = zip(*sorted(zip(severities, counts, colors), key=lambda x: severity_order[x[0]]))
    
    x_pos = np.arange(len(sorted_severities))
    
    plt.figure(figsize=(8, 6))
    plt.bar(x_pos, sorted_counts, color=sorted_colors)
    plt.xticks(x_pos, sorted_severities, rotation=10, ha='right', fontsize=10)
    plt.ylabel('Number of Vulnerabilities', fontsize=12)
    plt.title('Total Vulnerabilities Found', fontsize=14)
    
    # Add counts on top of bars
    for i, count in enumerate(sorted_counts):
        plt.text(x_pos[i], count + 0.1, str(count), ha='center', va='bottom', fontsize=10, weight='bold')
    
    plt.tight_layout() # Adjust layout to prevent labels overlapping
    plt.savefig(output_path)
    plt.close()
    
    return output_path

def create_severity_donut_chart(severity_counts, output_path):
    """
    Generates a donut chart showing the distribution of vulnerabilities by severity.
    Saves the chart as an image file with improved styling to match a typical cybersecurity report.
    """
    severities = list(severity_counts.keys())
    counts = list(severity_counts.values())
    total_findings = sum(counts)
    
    # Use colors from get_severity_colors for consistency
    color_map = {severity: get_severity_colors(severity)['row1_bg'] for severity in severity_order.keys()}
    # Default color for 'Unknown' if present
    color_map['Unknown'] = '#808080'  # Grey for unknown
    
    # Sort data by severity order
    sorted_severities, sorted_counts = zip(*sorted(severity_counts.items(), key=lambda item: severity_order.get(item[0], 5)))
    sorted_colors = [color_map.get(s) for s in sorted_severities]
    
    # Create the donut chart
    fig, ax = plt.subplots(figsize=(8, 6))  # Adjusted figure size for better proportions
    
    # Create the pie chart slices with a thinner donut (width=0.4 instead of 0.3)
    wedges, texts, autotexts = ax.pie(
        sorted_counts,
                                      colors=sorted_colors, 
        wedgeprops=dict(width=0.4, edgecolor='white'),  # Thinner donut with white edges for separation
        startangle=90,  # Start at the top
        autopct=lambda pct: f'{pct:.0f}%',  # Percentage format using a lambda function
        pctdistance=0.75,  # Move percentage labels closer to the center
        textprops={'fontsize': 10, 'weight': 'bold', 'color': 'white'}  # Style percentage labels
    )

    # Draw a circle in the center to make it a donut
    centre_circle = plt.Circle((0, 0), 0.60, fc='white')  # Adjusted center circle size
    fig.gca().add_artist(centre_circle)
    
    # Add total findings text in the center
    ax.text(
        0, 0,
        f'Total\nFindings\n{total_findings}',
        ha='center',
        va='center',
        fontsize=12,
        weight='bold',
        color='black'
    )
    
    # Add legend with counts, positioned below the chart to avoid overlap
    legend_labels = [f'{s} ({c})' for s, c in zip(sorted_severities, sorted_counts)]
    ax.legend( # type: ignore
        wedges,
        legend_labels,
        title="Severity Levels",
        loc="center",
        bbox_to_anchor=(0.5, -0.1),  # Place legend below the chart
        ncol=len(sorted_severities),  # Display in one row
        fontsize=10,
        title_fontsize=12,
        frameon=False  # Remove the legend frame for a cleaner look
    )
    
    # Set the title with consistent styling
    plt.title(
        'Overall Vulnerabilities Identified',
        fontsize=14,
        weight='bold',
        pad=20,
        color='#660099'  # Match the heading color used in the Word document
    )
    
    ax.axis('equal')  # Equal aspect ratio to ensure that pie is drawn as a circle
    plt.tight_layout()
    plt.savefig(output_path, bbox_inches='tight')  # Ensure the chart legend is not cut off
    plt.close()
    
    return output_path 

def index_images_from_poc_zip(root_dir):
    import re
    image_map = {}
    screenshots_root = None

    # If the current directory is 'POC Screenshots', use it directly
    if os.path.basename(root_dir).strip().lower() == 'poc screenshots':
        screenshots_root = root_dir
    else:
        # Look for the 'POC Screenshots' directory inside the extracted ZIP
        for dirpath, dirnames, _ in os.walk(root_dir):
            for d in dirnames:
                if d.strip().lower() == 'poc screenshots':
                    screenshots_root = os.path.join(dirpath, d)
                    break
            if screenshots_root:
                break

    if not screenshots_root:
        print("[ERROR] 'POC Screenshots' folder not found in ZIP.")
        return image_map

    # Walk through severity folders (Critical, High, etc.)
    for severity in os.listdir(screenshots_root):
        severity_path = os.path.join(screenshots_root, severity)
        if os.path.isdir(severity_path):
            for vuln_id in os.listdir(severity_path):
                vuln_path = os.path.join(severity_path, vuln_id)
                if os.path.isdir(vuln_path) and vuln_id.lower().startswith("vul-"):
                    for fname in os.listdir(vuln_path):
                        if fname.lower().endswith(('.png', '.jpg', '.jpeg')):
                            match = re.search(r'step(\d+)', fname.lower())
                            if match:
                                step_num = int(match.group(1))
                                key = f"{vuln_id.strip().lower()}_step{step_num}"
                                full_path = os.path.join(vuln_path, fname)
                                image_map[key] = full_path
                                print(f"[MAP] {key} → {full_path}")
    print('--- Final Image Map Keys ---')
    for k in image_map:
        print(k)
    print('--- End Image Map Keys ---')
    return image_map

def build_steps_with_images(steps_with_screenshots, image_map, tpl, sr_no):
    steps = []
    sr_no_key = sr_no.strip().lower() if sr_no else ''
    for idx, step in enumerate(steps_with_screenshots, 1):
        screenshot_name = step.get('screenshot', '')
        match = re.search(r'step(\\d+)', screenshot_name.strip().lower())
        if match:
            step_num = int(match.group(1))
        else:
            step_num = idx
        key = f"{sr_no_key}_step{step_num}"
        img_path = image_map.get(key) if image_map else None
        print(f"[POC IMAGE DEBUG] Step {idx}, Key: {key}, Path: {img_path}")
        if img_path and os.path.exists(img_path):
            print(f"[POC IMAGE FOUND] Using image: {img_path}")
            image = InlineImage(tpl, img_path, width=Inches(5))
        else:
            print(f"[POC IMAGE MISSING] {img_path}")
            image = f"[Screenshot missing: {screenshot_name}]"
        steps.append({'text': step['text'], 'image': image})
    return steps

def find_image_for_step(screenshot_name, image_map, sr_no, step_idx):
    if not image_map or not screenshot_name:
        return None
    screenshot_name_key = screenshot_name.strip().lower()
    sr_no_key = sr_no.strip().lower() if sr_no else ''
    match = re.search(r'step(\d+)', screenshot_name_key)
    if match:
        step_num = int(match.group(1))
    else:
        step_num = step_idx
    key = f"{sr_no_key}_step{step_num}"
    print(f"Looking for key: {key} in image_map")
    if key in image_map:
        print(f"Found image for {key}: {image_map[key]}")
    else:
        print(f"Image for {key} NOT FOUND")
    return image_map.get(key)

# After parsing vulnerabilities, ensure each has a display_sr_no
def assign_display_sr_no(vulnerabilities):
    for idx, vuln in enumerate(vulnerabilities, 1):
        if not vuln.get('display_sr_no'):
            sr_no = str(vuln.get('sr_no', f"VUL-{idx:03d}")).replace('VULN-', 'VUL-').strip()
            vuln['display_sr_no'] = sr_no
    return vulnerabilities 

def insert_poc_steps_section(doc, vulnerability, image_map):
    # Always add a new section for PoC steps
    doc.add_paragraph("Proof of Concept / Steps to Reproduce:").runs[0].bold = True
    steps = vulnerability.get('steps_with_screenshots', [])
    sr_no = vulnerability.get('display_sr_no') or vulnerability.get('sr_no')
    for idx, step in enumerate(steps, 1):
        step_para = doc.add_paragraph(f"Step {idx}: {step['text']}")
        step_para.paragraph_format.left_indent = Pt(10)
        screenshot_path = find_image_for_step(step['screenshot'], image_map, sr_no, idx)
        if screenshot_path and os.path.exists(screenshot_path):
            print(f"[DEBUG] Adding image from ZIP: {screenshot_path}")
            img_run = doc.add_paragraph().add_run()
            try:
                img_run.add_picture(screenshot_path, width=Inches(5))
            except Exception as img_err:
                print(f"[ERROR] Failed to insert screenshot {screenshot_path}: {img_err}")
                doc.add_paragraph(f"[Screenshot unavailable: {step['screenshot']}]")
        else:
            doc.add_paragraph(f"[Screenshot missing: {step['screenshot']}]") 

@app.get("/", response_class=HTMLResponse)
async def root():
    """Serve the Web Report Generator interface"""
    html_content = """<!DOCTYPE html>

<html>

<head>

    <title>VAPT Report Generator</title>

    <style>

        body {

            font-family: Arial, sans-serif;

            background-color: #1a1a2e;

            color: white;

            margin: 0;

            padding: 20px;

        }

        .container {

            max-width: 800px;

            margin: 0 auto;

            padding: 20px;

            background-color: #16213e;

            border-radius: 8px;

        }

        h1 {

            color: #8a2be2;

            text-align: center;

        }

        .form-group {

            margin-bottom: 15px;

        }

        label {

            display: block;

            margin-bottom: 5px;

        }

        input[type="file"], input[type="text"] {

            width: 100%;

            padding: 8px;

            border: 1px solid #6923d0;

            border-radius: 4px;

            background-color: #0f3460;

            color: white;

        }

        button {

            background-color: #6923d0;

            color: white;

            border: none;

            padding: 10px 15px;

            border-radius: 4px;

            cursor: pointer;

            font-weight: bold;

            margin-top: 10px;

        }

        button:hover {

            background-color: #8a2be2;

        }

        .status {

            margin-top: 20px;

            padding: 10px;

            border-radius: 4px;

            background-color: rgba(0, 0, 0, 0.3);

        }

        .info-section {

            margin-top: 20px;

            padding: 10px;

            border: 1px dashed #6923d0;

            border-radius: 4px;

            background-color: rgba(0, 0, 0, 0.2);

        }

        .info-section h3 {

            margin-top: 0;

            color: #8a2be2;

        }

        .info-section ul {

            margin-left: 20px;

            padding-left: 0;

        }

        .info-section li {

            margin-bottom: 5px;

        }

        .error-list {

            color: #ff5555;

            margin-top: 15px;

            background-color: rgba(255, 0, 0, 0.1);

            padding: 10px;

            border-radius: 4px;

            max-height: 200px;

            overflow-y: auto;

        }

        .error-list h4 {

            margin-top: 0;

            color: #ff5555;

        }

        .error-list ul {

            margin-top: 5px;

            padding-left: 20px;

        }

        .error-list li {

            margin-bottom: 5px;

        }

        .back-link {

            color: #8a2be2;

            text-decoration: underline;

            display: inline-block;

            margin-bottom: 20px;

        }

    </style>

</head>

<body>

    <div class="container">

        <a href="/report_formats.html" class="back-link">← Back to Report Formats</a>

        <h1>VAPT Report Generator</h1>

        

        <form id="generateForm" enctype="multipart/form-data">

            <div class="form-group">

                <label for="excelFile">Excel File with Vulnerabilities </label>

                <input type="file" id="excelFile" name="file" accept=".xlsx, .xls">

            </div>

            <div class="form-group">

                <label for="assessmentFile">Excel File with Assessment Details:</label>

                <input type="file" id="assessmentFile" name="assessment_file" accept=".xlsx, .xls">

            </div>

            <div class="form-group">

                <label for="pocImages">POC Screenshots (.zip):</label>

                <input type="file" id="pocImages" name="poc_images" accept=".zip">

            </div>

            <div class="form-group">

                <label for="clientName">Client Name:</label>

                <input type="text" id="clientName" name="client_name" placeholder="Client Name">

            </div>

            <div class="form-group">

                <label for="companyLogo">Company Logo:</label>

                <input type="file" id="companyLogo" name="company_logo" accept="image/png, image/jpeg">

            </div>

            

            <div class="info-section">

                <h3>How to Use This Tool</h3>

                <ul>

                    <li>Use the provided Excel template with 15 POC columns</li>

                    <li>Add vulnerability details in each row</li>

                    <li>Insert POC screenshot filenames in the respective "POC X" columns</li>

                    <li>Ensure your step description and POC screenshot filename is separated by line break (Enter/Return key)</li>

                    <li>The system will automatically map screenshots to steps in the final report</li>

                    <li>For each vulnerability, the number of steps should match the number of screenshots</li>

                </ul>

            </div>

            

            <button type="button" onclick="generateReport()" id="generateReportBtn">Generate Report</button>

        </form>

        

        <div id="status" class="status"></div>

        <div id="validationErrors" class="error-list" style="display: none;">

            <h4>Validation Errors</h4>

            <ul id="errorList"></ul>

        </div>

    </div>

    

    <script>

        function setStatus(message, isError = false) {

            const statusDiv = document.getElementById('status');

            statusDiv.innerText = message;

            statusDiv.style.backgroundColor = isError ? 'rgba(255, 0, 0, 0.3)' : 'rgba(0, 255, 0, 0.1)';

        }

        

        function showValidationErrors(errors) {

            const errorDiv = document.getElementById('validationErrors');

            const errorList = document.getElementById('errorList');

            

            // Clear previous errors

            errorList.innerHTML = '';

            

            if (errors && errors.length > 0) {

                errors.forEach(error => {

                    const li = document.createElement('li');

                    li.textContent = `Row ${error.row}: ${error.message}`;

                    errorList.appendChild(li);

                });

                

                errorDiv.style.display = 'block';

            } else {

                errorDiv.style.display = 'none';

            }

        }

        

        function generateReport() {

            const fileInput = document.getElementById('excelFile');

            const assessmentFileInput = document.getElementById('assessmentFile');

            const clientNameInput = document.getElementById('clientName');

            const pocImagesInput = document.getElementById('pocImages');

            const companyLogoInput = document.getElementById('companyLogo');

            const clientName = clientNameInput.value;

            const file = fileInput.files[0];

            const assessmentFile = assessmentFileInput.files[0];

            const pocImagesFile = pocImagesInput.files[0];

            const companyLogoFile = companyLogoInput.files[0];

            

            if (!file) {

                setStatus('Please select a Vulnerabilities Excel file', true);

                return;

            }

            if (!assessmentFile) {

                setStatus('Please select an Assessment Detail Excel file', true);

                return;

            }

            

            setStatus('Generating report, please wait...');

            document.getElementById('generateReportBtn').disabled = true;

            document.getElementById('validationErrors').style.display = 'none';

            

            const formData = new FormData();

            formData.append('file', file);

            formData.append('assessment_file', assessmentFile);

            formData.append('client_name', clientName);

            if (pocImagesFile) {

                formData.append('poc_images', pocImagesFile);

            }

            if (companyLogoFile) {

                formData.append('company_logo', companyLogoFile);

            }

            

            fetch('/type4/generate-report/', {

                method: 'POST',

                body: formData

            })

            .then(response => {

                document.getElementById('generateReportBtn').disabled = false;

                

                // Check for validation errors

                const validationErrors = response.headers.get('X-Validation-Errors');

                if (validationErrors) {

                    try {

                        const errors = JSON.parse(validationErrors);

                        if (errors && errors.length > 0) {

                            showValidationErrors(errors);

                            setStatus('Report generated with validation warnings. Please review the errors below.', true);

                        }

                    } catch (e) {

                        console.error('Error parsing validation errors:', e);

                    }

                }

                

                if (response.ok) {

                    if (!validationErrors || validationErrors === '[]') {

                        setStatus('Report generated successfully! Downloading...');

                    }

                    return response.blob();

                }

                return response.json().then(error => {

                    throw new Error(error.detail || 'Error generating report');

                });

            })

            .then(blob => {

                const url = window.URL.createObjectURL(blob);

                const a = document.createElement('a');

                a.href = url;

                a.download = `${clientName || 'Client'} VAPT Report.docx`;

                document.body.appendChild(a);

                a.click();

                window.URL.revokeObjectURL(url);

            })

            .catch(error => {

                console.error('Error:', error);

                setStatus('Error: ' + error.message, true);

                document.getElementById('generateReportBtn').disabled = false;

            });

        }

    </script>

</body>

</html>"""
    return HTMLResponse(content=html_content)

# Add a new API endpoint to save the default logo
@app.post("/upload-default-logo/")
async def upload_default_logo(logo_data: dict):
    """
    Upload a default logo image to be used when the user doesn't provide one.
    Expects a JSON payload with a base64 encoded image.
    """
    try:
        if 'base64_image' not in logo_data:
            raise HTTPException(status_code=400, detail="No image data provided")
        
        # Extract the base64 string and strip any header data
        base64_data = logo_data['base64_image']
        if ',' in base64_data:
            base64_data = base64_data.split(',', 1)[1]
        
        # Decode the base64 string to binary
        image_data = base64.b64decode(base64_data)
        
        # Save to the backend directory
        logo_path = os.path.join(os.path.dirname(__file__), "default_logo.png")
        with open(logo_path, "wb") as f:
            f.write(image_data)
        
        return JSONResponse(content={"status": "success", "message": "Default logo saved"})
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e)) 